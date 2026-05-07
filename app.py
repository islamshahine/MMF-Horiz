"""
AQUASIGHT™ MMF  —  Horizontal Multi-Media Filter Calculator
============================================================
Block flow (matches calculation dependency order):
  1  Process basis       → flow / filter per scenario
  2  Water properties    → density, viscosity (feed + BW, 3 scenarios)
  3  Vessel geometry     → real ID, volumes, cross-sections
     Mechanical          → ASME thickness, weight: body + nozzles + plate + supports
  4  Media design        → volumes, EBCT, LV, ΔP, inventory
  5  Backwash design     → collector check, expansion, pump/blower, sequence
  6  Weight summary      → consolidated empty weight
"""

import io
import pandas as pd
import streamlit as st

try:
    from docx import Document as _DocxDocument

    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

from engine.geometry   import segment_area, dish_volume
from engine.process    import filter_loading
from engine.water      import water_properties, FEED_PRESETS, BW_PRESETS
from engine.mechanical import (
    thickness, apply_thickness_override, empty_weight,
    nozzle_plate_design, saddle_weight, internals_weight,
    operating_weight, saddle_design,
    MATERIALS, RADIOGRAPHY_OPTIONS, JOINT_EFFICIENCY,
    STEEL_DENSITY_KG_M3, SUPPORT_TYPES,
    NOZZLE_DENSITY_MIN, NOZZLE_DENSITY_MAX, NOZZLE_DENSITY_DEFAULT,
    STRAINER_WEIGHT_KG, MANHOLE_WEIGHT_KG,
)
from engine.nozzles import (
    estimate_nozzle_schedule,
    FLANGE_RATINGS, SCHEDULES, DN_SERIES,
)
from engine.backwash import (
    backwash_hydraulics, bed_expansion, pressure_drop,
    bw_sequence, filtration_cycle, bw_system_sizing,
)
from engine.collector_ext import collector_check_ext
from engine.coating import (
    internal_surface_areas, lining_cost,
    PROTECTION_TYPES, RUBBER_TYPES, EPOXY_TYPES, CERAMIC_TYPES,
    DEFAULT_LABOR_RUBBER_M2, DEFAULT_LABOR_EPOXY_M2, DEFAULT_LABOR_CERAMIC_M2,
)
from engine.cartridge import (
    cartridge_design, cartridge_optimise,
    ELEMENT_SIZE_LABELS, RATING_UM_OPTIONS,
    HOUSING_CAPACITY_OPTIONS, DEFAULT_ELEMENTS_PER_HOUSING,
    DP_REPLACEMENT_BAR,
    SAFETY_FACTOR_STD, SAFETY_FACTOR_CIP,
    COST_TABLE_POLYMER, COST_TABLE_SS316L,
)
from engine.energy import hydraulic_profile, energy_summary
from engine.drawing import vessel_section_elevation, LAYER_COLORS as DRAWING_LAYER_COLORS
from engine.economics import (
    capex_breakdown, opex_annual, carbon_footprint, global_benchmark_comparison,
)

try:
    import plotly.graph_objects as _go
    import plotly.express as _px
    _PLOTLY_OK = True
except ImportError:
    _PLOTLY_OK = False

# ══════════════════════════════════════════════════════════════════════════════
# PAGE CONFIG
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(page_title="AQUASIGHT™ MMF", layout="wide",
                   initial_sidebar_state="collapsed")

# ══════════════════════════════════════════════════════════════════════════════
# MEDIA PRESETS  (default 3-layer: Gravel / Sand / Anthracite)
# ══════════════════════════════════════════════════════════════════════════════
def _eps0_from_psi(psi: float) -> float:
    """Empirical estimate: ε₀ ≈ 0.4 + 0.1·(1−ψ)/ψ  (Kozeny-based, random packing)."""
    return round(0.4 + 0.1 * (1.0 - psi) / max(psi, 0.01), 3)

def _rho_eff_porous(rho_dry: float, eps_p: float, rho_water: float = 1025.0) -> float:
    """Water-saturated particle density for porous media.
    ρ_eff = ρ_dry + ρ_water × ε_p
    where ρ_dry is apparent dry-particle density and ε_p is internal particle porosity.
    """
    return rho_dry + rho_water * eps_p

# Fields: d10 (mm), cu (CU=d60/d10), epsilon0, rho_p_eff (kg/m³), d60 (mm),
#         psi (sphericity), is_porous, default_depth (m)
DEFAULT_MEDIA_PRESETS = {
    "Gravel":            {"d10": 6.0,  "cu": 1.0, "epsilon0": 0.46, "psi": 0.90,
                          "rho_p_eff": 2600, "d60": 6.00, "is_porous": False, "default_depth": 0.20},
    "Coarse sand":       {"d10": 1.35, "cu": 1.5, "epsilon0": 0.44, "psi": 0.85,
                          "rho_p_eff": 2650, "d60": 2.03, "is_porous": False, "default_depth": 0.60},
    "Fine sand":         {"d10": 0.80, "cu": 1.3, "epsilon0": 0.42, "psi": 0.80,
                          "rho_p_eff": 2650, "d60": 1.04, "is_porous": False, "default_depth": 0.80},
    "Fine sand (extra)": {"d10": 0.50, "cu": 1.3, "epsilon0": 0.41, "psi": 0.75,
                          "rho_p_eff": 2650, "d60": 0.65, "is_porous": False, "default_depth": 0.70},
    "Anthracite":        {"d10": 1.30, "cu": 1.5, "epsilon0": 0.48, "psi": 0.70,
                          "rho_p_eff": 1450, "d60": 2.25, "is_porous": False, "default_depth": 0.80},
    "MnO₂":             {"d10": 1.00, "cu": 2.4, "epsilon0": 0.50, "psi": 0.65,
                          "rho_p_eff": 4200, "d60": 2.40, "is_porous": False, "default_depth": 0.40},
    "Medium GAC":        {"d10": 1.00, "cu": 1.6, "epsilon0": 0.55, "psi": 0.65,
                          "rho_p_eff": 1000, "d60": 1.44, "is_porous": True,  "default_depth": 1.00},
    "Biodagene":         {"d10": 2.50, "cu": 1.4, "epsilon0": 0.42, "psi": 0.80,
                          "rho_p_eff": 1600, "d60": 3.50, "is_porous": False, "default_depth": 0.60},
    "Schist":            {"d10": 3.30, "cu": 1.5, "epsilon0": 0.47, "psi": 0.65,
                          "rho_p_eff": 1300, "d60": 4.95, "is_porous": False, "default_depth": 0.30},
    "Limestone":         {"d10": 3.00, "cu": 1.4, "epsilon0": 0.55, "psi": 0.60,
                          "rho_p_eff": 2700, "d60": 4.20, "is_porous": False, "default_depth": 0.50},
    "Pumice":            {"d10": 1.50, "cu": 1.3, "epsilon0": 0.55, "psi": 0.55,
                          "rho_p_eff":  900, "d60": 1.56, "is_porous": True,  "default_depth": 0.60},
    "FILTRALITE clay":   {"d10": 1.20, "cu": 1.5, "epsilon0": 0.48, "psi": 0.50,
                          "rho_p_eff": 1250, "d60": 1.80, "is_porous": True,  "default_depth": 0.80},
    "Custom":            {"d10": 0.0,  "cu": 1.5, "epsilon0": 0.42, "psi": 0.80,
                          "rho_p_eff": 2650, "d60": 0.0,  "is_porous": False, "default_depth": 0.50},
}

# Always sync: add any new catalogue entries, remove renamed ones
if "media_presets" not in st.session_state or set(
        st.session_state.media_presets.keys()) != set(DEFAULT_MEDIA_PRESETS.keys()):
    st.session_state.media_presets = DEFAULT_MEDIA_PRESETS.copy()


# ══════════════════════════════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════════════════════════════
h1, h2 = st.columns([3, 1])
with h1:
    st.markdown("## AQUASIGHT™ &nbsp; Horizontal Multi-Media Filter")
with h2:
    st.markdown(
        "<p style='text-align:right;color:grey;font-size:13px;margin-top:18px'>"
        "Islam Shahine · Process Expert</p>", unsafe_allow_html=True)
st.divider()

# ══════════════════════════════════════════════════════════════════════════════
# LAYOUT: context tabs (left) | output tabs (right)
# ══════════════════════════════════════════════════════════════════════════════
ctx, main = st.columns([1, 4])

with ctx:
    proc_tab, vessel_tab, media_tab, bw_tab, econ_tab = st.tabs([
        "⚙️ Process", "🏗️ Vessel", "🧱 Media", "🔄 BW", "💰 Econ"
    ])

    # ── Tab 1: Process ────────────────────────────────────────────────────
    with proc_tab:
        st.markdown("**Project**")
        project_name = st.text_input("Project",     value="NPC SWRO 60 000 m³/d")
        doc_number   = st.text_input("Doc. No.",    value="EXXXX-VWT-PCS-CAL-2001")
        revision     = st.text_input("Revision",    value="A1")
        client       = st.text_input("Client",      value="")
        engineer     = st.text_input("Prepared by", value="Islam Shahine")

        st.markdown("**Filter configuration**")
        total_flow = st.number_input("Total plant flow (m³/h)", value=21000.0, step=100.0)
        streams    = int(st.number_input("Streams", value=1, min_value=1))
        n_filters  = int(st.number_input("Filters / stream", value=16, min_value=1))
        redundancy = int(st.selectbox("Redundancy (per stream)", [0, 1, 2, 3, 4], index=1))
        q_n = total_flow / streams / n_filters
        st.caption(
            f"Flow / filter (N): **{q_n:.1f} m³/h**  \n"
            f"Redundancy = {redundancy} standby filter(s) per stream  \n"
            f"Total active filters (N scenario): **{streams * n_filters} plant-wide**"
        )

        st.markdown("**Water quality — feed**")
        feed_preset = st.selectbox("Feed preset", list(FEED_PRESETS.keys()), index=2, key="feed_pre")
        fp = FEED_PRESETS[feed_preset]
        feed_sal  = st.number_input("Feed salinity (ppt)",    value=fp["salinity_ppt"], step=0.5, key="f_sal")
        feed_temp = st.number_input("Feed temp — avg (°C)",   value=fp["temp_c"],        step=1.0, key="f_tmp")
        temp_low  = st.number_input("Feed temp — min (°C)",   value=15.0, step=1.0, key="t_low")
        temp_high = st.number_input("Feed temp — max (°C)",   value=35.0, step=1.0, key="t_high")
        tss_low   = st.number_input("Feed TSS — low (mg/L)",  value=5.0,  step=1.0)
        tss_avg   = st.number_input("Feed TSS — avg (mg/L)",  value=10.0, step=1.0)
        tss_high  = st.number_input("Feed TSS — high (mg/L)", value=20.0, step=1.0)

        st.markdown("**Water quality — backwash**")
        bw_preset = st.selectbox("BW preset", list(BW_PRESETS.keys()), index=0, key="bw_pre")
        bp = BW_PRESETS[bw_preset] or fp
        bw_sal  = st.number_input("BW salinity (ppt)", value=bp["salinity_ppt"], step=0.5, key="b_sal")
        bw_temp = st.number_input("BW temp (°C)",      value=bp["temp_c"],       step=1.0, key="b_tmp")

        st.markdown("**Performance thresholds**")
        velocity_threshold = st.number_input("Max LV (m/h)",   value=12.0)
        ebct_threshold     = st.number_input("Min EBCT (min)", value=5.0)

        st.markdown("**Cartridge filter**")
        cart_flow   = st.number_input(
            "Design flow (m³/h)", value=float(total_flow), step=100.0, key="cart_flow",
            help="Total flow to the cartridge station (usually = plant flow)")
        cart_size   = st.selectbox(
            "Element length", ELEMENT_SIZE_LABELS, index=2, key="cart_size",
            help="All elements are 2.5\" (63.5 mm) OD.")
        cart_rating = st.selectbox(
            "Rating (μm absolute)", RATING_UM_OPTIONS, index=1, key="cart_rating")
        cart_cip = st.toggle(
            "CIP system (SS 316L elements)", value=False, key="cart_cip",
            help="CIP: regenerable SS 316L elements. SF=1.2, DHC=45 g/TIE.")
        _hsg_options     = [str(r) for r in HOUSING_CAPACITY_OPTIONS] + ["Custom…"]
        _hsg_default_idx = HOUSING_CAPACITY_OPTIONS.index(DEFAULT_ELEMENTS_PER_HOUSING)
        cart_hsg_sel = st.selectbox(
            "Elements per housing", _hsg_options, index=_hsg_default_idx, key="cart_hsg_sel")
        if cart_hsg_sel == "Custom…":
            cart_housing = st.number_input(
                "Custom elements per housing", min_value=1, max_value=500,
                value=100, step=1, key="cart_hsg_custom")
        else:
            cart_housing = int(cart_hsg_sel)
        _cf_inlet_max  = float(tss_avg)
        _cf_outlet_max = round(0.15 * tss_avg, 2)
        cf_inlet_tss = st.number_input(
            "CF inlet TSS (mg/L)", min_value=0.0, max_value=_cf_inlet_max,
            value=min(2.0, _cf_inlet_max), step=0.1, format="%.2f", key="cf_inlet_tss",
            help=f"MMF effluent entering CF. Max = feed TSS ({tss_avg:.1f} mg/L).")
        cf_outlet_tss = st.number_input(
            "CF outlet TSS — target (mg/L)", min_value=0.0, max_value=_cf_outlet_max,
            value=min(0.5, _cf_outlet_max), step=0.05, format="%.2f", key="cf_outlet_tss",
            help=f"Max = 15 % of MMF feed TSS = {_cf_outlet_max:.2f} mg/L.")
        _sf_label = f"SF = {SAFETY_FACTOR_CIP}" if cart_cip else f"SF = {SAFETY_FACTOR_STD}"
        st.caption(
            f"{'🔩 SS 316L CIP — ' + _sf_label if cart_cip else '🔵 Polymer standard — ' + _sf_label}.  "
            "Replacement interval calculated from DHC ÷ TSS loading rate.")

    # ── Tab 2: Vessel ─────────────────────────────────────────────────────
    with vessel_tab:
        st.markdown("**Vessel geometry**")
        nominal_id   = st.number_input("Nominal internal diameter (m)", value=5.5, step=0.1,
                                       help="Lining reduces the hydraulic ID.")
        total_length = st.number_input("Total length T/T (m)", value=24.3, step=0.1)
        end_geometry = st.selectbox("End geometry", ["Elliptic 2:1", "Torispherical 10%"])

        st.markdown("**Mechanical (ASME)**")
        material_name   = st.selectbox("Material", list(MATERIALS.keys()), index=3)
        mat_info        = MATERIALS[material_name]
        st.caption(f"*{mat_info['description']}*")
        design_pressure = st.number_input("Design pressure (bar)", value=7.0, step=0.5)
        design_temp     = st.number_input("Design temperature (°C)", value=50.0, step=5.0)
        corrosion       = st.number_input("Corrosion allowance (mm)", value=1.5, step=0.5)
        st.markdown("*Radiography (ASME UW-11)*")
        rc1, rc2 = st.columns(2)
        with rc1:
            shell_radio = st.selectbox("Shell", RADIOGRAPHY_OPTIONS, index=2, key="sh_r")
            st.caption(f"E = {JOINT_EFFICIENCY[shell_radio]:.2f}")
        with rc2:
            head_radio  = st.selectbox("Head",  RADIOGRAPHY_OPTIONS, index=2, key="hd_r")
            st.caption(f"E = {JOINT_EFFICIENCY[head_radio]:.2f}")
        st.markdown("*Thickness overrides* (0 = use calculated)")
        ov_shell = st.number_input("Shell t override (mm)", value=0.0, step=1.0, key="ov_sh")
        ov_head  = st.number_input("Head t override (mm)",  value=0.0, step=1.0, key="ov_hd")
        steel_density = st.number_input("Steel density (kg/m³)", value=STEEL_DENSITY_KG_M3,
                                        help="7850 CS · 7900 SS 304/316")

        st.markdown("**Internal protection**")
        protection_type = st.selectbox(
            "Protection type", PROTECTION_TYPES, index=1, key="prot_type",
            help="Rubber lining: bonded sheet, reduces hydraulic ID by 2×thickness.  "
                 "Epoxy / Ceramic: applied in coats (DFT in µm), negligible ID impact.")

        if protection_type == "Rubber lining":
            rubber_type_sel = st.selectbox(
                "Rubber type", list(RUBBER_TYPES.keys()), index=1, key="rub_type")
            lining_mm = st.number_input(
                "Rubber thickness / layer (mm)", value=4.0, step=0.5,
                min_value=0.5, key="rub_t",
                help="Nominal ID reduced by 2 × thickness × layers.")
            rubber_layers = st.number_input(
                "Layers", value=2, min_value=1, max_value=6, key="rub_lay")
            _rub_def_cost = RUBBER_TYPES[rubber_type_sel]["default_cost_m2"]
            rubber_cost_m2 = st.number_input(
                "Rubber material cost (USD/m²)", value=float(_rub_def_cost),
                step=5.0, key="rub_cost")
            rubber_labor_m2 = st.number_input(
                "Application labour (USD/m²)", value=DEFAULT_LABOR_RUBBER_M2,
                step=5.0, key="rub_lab")
        else:
            lining_mm = 0.0
            rubber_type_sel = "EPDM"; rubber_layers = 2
            rubber_cost_m2 = 0.0; rubber_labor_m2 = DEFAULT_LABOR_RUBBER_M2

        if protection_type == "Epoxy coating":
            epoxy_type_sel = st.selectbox(
                "Epoxy type", list(EPOXY_TYPES.keys()), index=1, key="epx_type")
            _epx_cat = EPOXY_TYPES[epoxy_type_sel]
            epoxy_dft_um = st.number_input(
                "DFT per coat (µm)", value=float(_epx_cat["default_dft_um"]),
                step=25.0, min_value=50.0, key="epx_dft")
            epoxy_coats = st.number_input(
                "Number of coats", value=_epx_cat["default_coats"],
                min_value=1, max_value=6, key="epx_coats")
            epoxy_cost_m2 = st.number_input(
                "Epoxy material cost (USD/m²)", value=float(_epx_cat["default_cost_m2"]),
                step=2.0, key="epx_cost")
            epoxy_labor_m2 = st.number_input(
                "Application labour (USD/m²)", value=DEFAULT_LABOR_EPOXY_M2,
                step=2.0, key="epx_lab")
        else:
            epoxy_type_sel = "High-build epoxy"; epoxy_dft_um = 350.0
            epoxy_coats = 2; epoxy_cost_m2 = 0.0; epoxy_labor_m2 = DEFAULT_LABOR_EPOXY_M2

        if protection_type == "Ceramic coating":
            ceramic_type_sel = st.selectbox(
                "Ceramic type", list(CERAMIC_TYPES.keys()), index=0, key="cer_type")
            _cer_cat = CERAMIC_TYPES[ceramic_type_sel]
            cc1, cc2 = st.columns(2)
            ceramic_dft_um = cc1.number_input(
                "DFT / coat (µm)", value=float(_cer_cat["default_dft_um"]),
                step=50.0, min_value=100.0, key="cer_dft")
            ceramic_coats = int(cc2.number_input(
                "Coats", value=int(_cer_cat["default_coats"]),
                step=1, min_value=1, max_value=6, key="cer_coats"))
            ceramic_cost_m2 = st.number_input(
                "Ceramic material cost (USD/m²)", value=float(_cer_cat["default_cost_m2"]),
                step=10.0, key="cer_cost")
            ceramic_labor_m2 = st.number_input(
                "Application labour (USD/m²)", value=DEFAULT_LABOR_CERAMIC_M2,
                step=5.0, key="cer_lab")
        else:
            ceramic_type_sel = "Ceramic-filled epoxy"
            ceramic_dft_um = 500.0; ceramic_coats = 2
            ceramic_cost_m2 = 0.0; ceramic_labor_m2 = DEFAULT_LABOR_CERAMIC_M2

    # ── Tab 3: Media ──────────────────────────────────────────────────────
    with media_tab:
        st.markdown("**Nozzle plate**")
        nozzle_plate_h = st.number_input("Nozzle plate height (m)", value=1.0, step=0.05)
        np_bore_dia    = st.number_input("Bore diameter (mm)", value=50.0,
                                         step=5.0, min_value=10.0, key="np_bd")
        np_density     = st.number_input(
            "Nozzle density (/m²)", value=NOZZLE_DENSITY_DEFAULT,
            min_value=NOZZLE_DENSITY_MIN, max_value=NOZZLE_DENSITY_MAX,
            step=1.0, key="np_den",
            help=f"{NOZZLE_DENSITY_MIN:.0f}–{NOZZLE_DENSITY_MAX:.0f} nozzles/m²")
        np_beam_sp     = st.number_input("Beam spacing (mm)", value=500.0,
                                         step=50.0, key="np_bs",
                                         help="Stiffener beam spacing — effective bending span")
        np_override_t  = st.number_input("Override plate t (mm) — 0=calc",
                                         value=0.0, step=1.0, key="np_ov")

        st.markdown("**Media layers**")
        n_layers = int(st.selectbox("Layers", [1, 2, 3, 4, 5, 6], index=2))
        _rho_water_sidebar = water_properties(feed_temp, feed_sal)["density_kg_m3"]
        layers = []
        default_types = ["Gravel", "Fine sand", "Anthracite"]
        for i in range(n_layers):
            st.markdown(f"**Layer {i+1}** (bottom → top)")
            def_type = default_types[i] if i < 3 else "Custom"
            m_type = st.selectbox("Type",
                                  list(st.session_state.media_presets.keys()),
                                  index=list(st.session_state.media_presets.keys()
                                             ).index(def_type),
                                  key=f"lt_{i}")
            preset = st.session_state.media_presets[m_type]
            depth  = st.number_input("Depth (m)", value=preset["default_depth"],
                                     step=0.05, key=f"ld_{i}")
            default_sup = (m_type == "Gravel")
            is_sup = st.checkbox("Support media (no clogging)", value=default_sup,
                                 key=f"sup_{i}")
            if not is_sup:
                cap_raw = st.number_input(
                    "Capture weight", value=round(depth * 100, 0), step=5.0,
                    min_value=0.0, key=f"cap_{i}",
                    help="Relative TSS capture weight — auto-normalised across "
                         "non-support layers. E.g., 30 + 70 → 30 %/70 %.")
                cap_frac = cap_raw / 100.0
            else:
                cap_frac = 0.0
            data = preset.copy()
            if m_type == "Custom":
                _c1, _c2 = st.columns(2)
                _d10 = _c1.number_input("d10 (mm)", value=1.0, step=0.05,
                                        min_value=0.01, key=f"d10_{i}")
                _cu  = _c2.number_input("CU (d60/d10)", value=1.5, step=0.05,
                                        min_value=1.0, key=f"cu_{i}")
                _psi = _c1.number_input("Sphericity ψ", value=0.80,
                                        step=0.05, min_value=0.3, max_value=1.0,
                                        key=f"psi_{i}",
                                        help="1 = perfect sphere; 0.5–0.9 typical")
                _eps0_est = _eps0_from_psi(_psi)
                _eps0 = _c2.number_input(
                    "Voidage ε₀", value=_eps0_est, step=0.01,
                    min_value=0.25, max_value=0.70, key=f"ep_{i}",
                    help=f"Estimated from ψ: {_eps0_est:.3f} (Kozeny empirical)")
                _is_por = st.checkbox("Porous media (water fills particle pores)",
                                      value=False, key=f"por_{i}")
                if _is_por:
                    _p1, _p2 = st.columns(2)
                    _rho_dry = _p1.number_input(
                        "Dry apparent density (kg/m³)", value=500.0, step=50.0,
                        min_value=100.0, key=f"rhd_{i}")
                    _eps_p = _p2.number_input(
                        "Particle internal porosity εₚ", value=0.50,
                        step=0.05, min_value=0.0, max_value=0.95, key=f"epp_{i}")
                    _rho_eff = _rho_eff_porous(_rho_dry, _eps_p, _rho_water_sidebar)
                    st.caption(f"ρ_eff = {_rho_eff:.0f} kg/m³")
                else:
                    _rho_eff = st.number_input(
                        "Particle density (kg/m³)", value=2650.0, step=50.0,
                        min_value=100.0, key=f"rh_{i}")
                data["d10"]       = _d10
                data["cu"]        = _cu
                data["d60"]       = round(_d10 * _cu, 3)
                data["epsilon0"]  = _eps0
                data["psi"]       = _psi
                data["is_porous"] = _is_por
                data["rho_p_eff"] = round(_rho_eff, 1)
            layers.append({**data, "Type": m_type, "Depth": depth,
                           "is_support": is_sup, "capture_frac": cap_frac})

        st.markdown("**Filtration performance**")
        solid_loading = st.number_input("Solid loading before BW (kg/m²)", value=1.5, step=0.1)
        captured_solids_density = st.number_input(
            "Captured solids density (kg/m³)", value=1020.0, step=10.0,
            help="Density of TSS retained in media voids — typically 1010–1050 kg/m³")
        alpha_9 = st.number_input(
            "Specific cake resistance α (× 10⁹ m/kg)",
            value=0.0, step=5.0, min_value=0.0, key="alpha_res",
            help=(
                "Resistance of deposited TSS cake per unit mass (Ruth model). "
                "0 = auto-calibrate so that ΔP reaches the trigger at M_max. "
                "Typical ranges: coarse mineral / silt 0.1–10 · "
                "seawater mixed TSS 10–50 · organic-rich / algae 100–500 · "
                "clay / fine colloids 1 000–10 000  (all × 10⁹ m/kg)."
            ))
        alpha_specific = alpha_9 * 1e9

    # ── Tab 4: BW ─────────────────────────────────────────────────────────
    with bw_tab:
        st.markdown("**BW hydraulics**")
        collector_h    = st.number_input(
            "BW outlet collector height (m)", value=4.2, step=0.1,
            help="Height from vessel bottom to BW outlet collector / trough")
        freeboard_mm   = st.number_input(
            "Min. freeboard (mm)", value=200, step=50, min_value=50, key="fb_mm",
            help="Minimum clearance between expanded bed top and collector.")
        bw_velocity    = st.number_input("Proposed BW velocity (m/h)", value=30.0, step=5.0)
        air_scour_rate = st.number_input("Air scour rate (m/h)", value=55.0, step=5.0)

        st.markdown("**BW sequence**")
        bw_cycles_day  = int(st.number_input("BW cycles / filter / day", value=1, min_value=1))
        dp_trigger_bar = st.number_input(
            "BW initiation ΔP setpoint (bar)", value=1.0, step=0.1,
            min_value=0.01, key="dp_trig",
            help="Filter triggers BW when ΔP across media reaches this value")
        bw_s_drain  = st.number_input("① Gravity drain (min)",       value=10, step=1, min_value=0, key="bws1")
        bw_s_air    = st.number_input("② Air scour only (min)",       value=1,  step=1, min_value=0, key="bws2")
        bw_s_airw   = st.number_input("③ Air + low-rate water (min)", value=5,  step=1, min_value=0, key="bws3")
        bw_s_hw     = st.number_input("④ High-rate water flush (min)",value=10, step=1, min_value=0, key="bws4")
        bw_s_settle = st.number_input("⑤ Settling (min)",             value=2,  step=1, min_value=0, key="bws5")
        bw_s_fill   = st.number_input("⑥ Fill & rinse (min)",         value=10, step=1, min_value=0, key="bws6")
        bw_total_min = bw_s_drain + bw_s_air + bw_s_airw + bw_s_hw + bw_s_settle + bw_s_fill
        st.metric("Total BW duration", f"{bw_total_min} min")

        st.markdown("**Equipment sizing**")
        vessel_pressure_bar = st.number_input(
            "Vessel operating pressure (bar g)", value=2.0, step=0.5,
            min_value=0.0, key="ves_press",
            help="Gauge pressure inside the filter vessel during BW.")
        blower_eta = st.number_input(
            "Blower isentropic efficiency", value=0.70, step=0.01,
            min_value=0.30, max_value=0.95, key="blower_eta")
        blower_inlet_temp_c = st.number_input(
            "Blower inlet air temperature (°C)", value=30.0, step=5.0,
            min_value=-10.0, max_value=60.0, key="blower_t")
        tank_sf = st.number_input(
            "BW tank safety factor", value=1.5, step=0.1,
            min_value=1.0, max_value=3.0, key="tank_sf",
            help="Tank volume = BW vol/cycle × simultaneous systems × SF.")
        bw_head_mwc = st.number_input(
            "BW pump total head (mWC)", value=15.0, step=1.0,
            min_value=1.0, key="bw_hd",
            help="Typical 12–20 mWC; includes bed + nozzle plate + BW piping losses.")

        st.markdown("**Nozzles & supports**")
        default_rating  = st.selectbox("Flange rating", FLANGE_RATINGS, index=1)
        nozzle_stub_len = st.number_input("Nozzle stub length (mm)", value=350, step=50)
        strainer_mat    = st.selectbox("Strainer material",
                                       list(STRAINER_WEIGHT_KG.keys()), index=0,
                                       help="SS316 seawater · HDPE/PP fresh/brackish")
        air_header_dn   = st.number_input("Air scour header DN (mm)", value=200, step=50, key="ah_dn")
        manhole_dn      = st.selectbox("Manhole size", list(MANHOLE_WEIGHT_KG.keys()), index=0)
        n_manholes      = int(st.number_input("No. of manholes", value=1, min_value=0, step=1))
        support_type    = st.selectbox("Support type", SUPPORT_TYPES, key="sup_t")
        if "Saddle" in support_type:
            saddle_h      = st.number_input("Saddle height (m)", value=0.8, step=0.05, key="sad_h")
            base_plate_t  = st.number_input("Base plate t (mm)", value=20.0, step=2.0, key="sad_bp")
            gusset_t      = st.number_input("Gusset t (mm)",     value=12.0, step=2.0, key="sad_gt")
            saddle_contact_angle = st.number_input(
                "Saddle contact angle (°)", value=120.0, step=15.0,
                min_value=90.0, max_value=180.0, key="sad_ang",
                help="120° is standard; 150° for heavy/thin-walled vessels.")
            leg_h = 1.2; leg_section = 150.0
        else:
            leg_h         = st.number_input("Leg height (m)",   value=1.2, step=0.1, key="leg_h")
            leg_section   = st.number_input("Leg section (mm)", value=150.0, step=25.0, key="leg_s")
            base_plate_t  = st.number_input("Base plate t (mm)", value=20.0, step=2.0, key="leg_bp")
            gusset_t      = st.number_input("Gusset t (mm)",     value=12.0, step=2.0, key="leg_gt")
            saddle_h = 0.8
            saddle_contact_angle = 120.0

    # ── Tab 5: Econ ───────────────────────────────────────────────────────
    with econ_tab:
        st.markdown("**Pump hydraulics**")
        np_slot_dp    = st.number_input(
            "Strainer nozzle plate ΔP at design LV (bar)", value=0.02,
            step=0.005, min_value=0.0, format="%.3f", key="np_slot",
            help="ΔP through strainer nozzle slots. Typical 0.01–0.05 bar.")
        p_residual    = st.number_input(
            "Required downstream pressure (barg)", value=2.50, step=0.25,
            min_value=0.0, key="p_res",
            help="Residual pressure at downstream tie-in. Typically 2–4 barg.")
        dp_inlet_pipe = st.number_input(
            "Inlet piping losses (bar)", value=0.30, step=0.05, min_value=0.0, key="dp_in")
        dp_dist       = st.number_input(
            "Inlet distributor ΔP (bar)", value=0.02, step=0.01, min_value=0.0, key="dp_dist")
        dp_outlet_pipe = st.number_input(
            "Outlet piping losses (bar)", value=0.20, step=0.05, min_value=0.0, key="dp_out")
        static_head   = st.number_input(
            "Static elevation head (m)", value=0.0, step=0.5, key="stat_h")

        st.markdown("**Efficiencies**")
        pump_eta    = st.number_input("Filtration pump η", value=0.75, step=0.01,
                                      min_value=0.30, max_value=0.95, key="pump_e")
        bw_pump_eta = st.number_input("BW pump η",         value=0.72, step=0.01,
                                      min_value=0.30, max_value=0.95, key="bwp_e")
        motor_eta   = st.number_input("Motor η (all motors)", value=0.95, step=0.01,
                                      min_value=0.70, max_value=0.99, key="mot_e")

        st.markdown("**Energy economics**")
        elec_tariff = st.number_input("Electricity tariff (USD/kWh)", value=0.10,
                                      step=0.01, min_value=0.01, key="elec_t")
        op_hours_yr = st.number_input("Operating hours / year", value=8400,
                                      step=100, min_value=1000, key="op_hr")

        st.markdown("**CAPEX inputs**")
        design_life_years        = st.number_input("Design life (years)", value=20, step=1, min_value=5, key="des_life")
        discount_rate            = st.number_input("Discount rate (%)", value=5.0, step=0.5, min_value=0.0, key="disc_rate")
        currency                 = st.selectbox("Currency", ["USD", "EUR", "GBP", "SAR", "AED"], key="currency")
        steel_cost_usd_kg        = st.number_input("Steel cost (USD/kg)", value=3.5, step=0.1, key="st_cost")
        erection_usd_vessel      = st.number_input("Erection cost (USD/vessel)", value=50000.0, step=5000.0, key="erect_usd")
        piping_usd_vessel        = st.number_input("Piping cost (USD/vessel)",   value=80000.0, step=5000.0, key="pip_usd")
        instrumentation_usd_vessel = st.number_input("Instrumentation (USD/vessel)", value=30000.0, step=5000.0, key="instr_usd")
        civil_usd_vessel         = st.number_input("Civil works (USD/vessel)",   value=40000.0, step=5000.0, key="civil_usd")
        engineering_pct          = st.number_input("Engineering (%)", value=12.0, step=1.0, min_value=0.0, key="eng_pct")
        contingency_pct          = st.number_input("Contingency (%)", value=10.0, step=1.0, min_value=0.0, key="cont_pct")

        st.markdown("**OPEX inputs**")
        media_replace_years   = st.number_input("Media replacement interval (years)", value=7.0, step=1.0, key="med_int")
        econ_media_gravel     = st.number_input("Gravel cost (USD/t)",     value=80.0,  step=10.0, key="mc_gr")
        econ_media_sand       = st.number_input("Sand cost (USD/t)",       value=150.0, step=10.0, key="mc_sd")
        econ_media_anthracite = st.number_input("Anthracite cost (USD/t)", value=400.0, step=25.0, key="mc_an")
        nozzle_replace_years  = st.number_input("Nozzle replacement interval (years)", value=10.0, step=1.0, key="noz_int")
        nozzle_unit_cost      = st.number_input("Nozzle unit cost (USD/nozzle)", value=15.0, step=1.0, key="noz_cost")
        labour_usd_filter_yr  = st.number_input("Labour (USD/filter/year)", value=5000.0, step=500.0, key="lab_usd")
        chemical_cost_m3      = st.number_input("Chemical cost (USD/m³ treated)", value=0.005,
                                                step=0.001, format="%.3f", key="chem_m3")

        st.markdown("**Carbon footprint**")
        grid_intensity       = st.number_input("Grid intensity (kgCO₂/kWh)", value=0.45, step=0.01, key="grid_co2")
        steel_carbon_kg      = st.number_input("Steel embodied carbon (kgCO₂/kg)", value=1.85, step=0.05, key="st_co2")
        concrete_carbon_kg   = st.number_input("Concrete embodied carbon (kgCO₂/kg)", value=0.13, step=0.01, key="con_co2")
        media_co2_gravel     = st.number_input("Gravel carbon (kgCO₂/kg)", value=0.004, step=0.001, format="%.3f", key="mco_gr")
        media_co2_sand       = st.number_input("Sand carbon (kgCO₂/kg)", value=0.006, step=0.001, format="%.3f", key="mco_sd")
        media_co2_anthracite = st.number_input("Anthracite carbon (kgCO₂/kg)", value=0.15, step=0.01, key="mco_an")

# ══════════════════════════════════════════════════════════════════════════════
# PRE-COMPUTE — all calculations in dependency order
# ══════════════════════════════════════════════════════════════════════════════

# ── Block 2: Water properties ──────────────────────────────────────────────
feed_wp = water_properties(feed_temp, feed_sal)
bw_wp   = water_properties(bw_temp,  bw_sal)

rho_feed = feed_wp["density_kg_m3"]
mu_feed  = feed_wp["viscosity_pa_s"]
rho_bw   = bw_wp["density_kg_m3"]
mu_bw    = bw_wp["viscosity_pa_s"]

# ── Block 3: Vessel geometry ───────────────────────────────────────────────
h_dish  = (nominal_id / 4) if end_geometry == "Elliptic 2:1" \
          else (0.2 * nominal_id)
cyl_len = total_length - 2 * h_dish

# Real hydraulic ID (nominal minus lining)
real_id = nominal_id - 2.0 * lining_mm / 1000.0

# ── Block 3: Mechanical ────────────────────────────────────────────────────
mech_base = thickness(
    diameter_m=nominal_id,
    design_pressure_bar=design_pressure,
    material_name=material_name,
    shell_radio=shell_radio,
    head_radio=head_radio,
    corrosion_mm=corrosion,
    internal_lining_mm=lining_mm,
)
mech = apply_thickness_override(
    mech_base,
    override_shell_mm=ov_shell,
    override_head_mm=ov_head,
    internal_lining_mm=lining_mm,
    nominal_id_m=nominal_id,
)
# Attach corrosion for override floor check
mech["corrosion_mm"] = corrosion

wt_body = empty_weight(
    diameter_m=real_id,
    straight_length_m=cyl_len,
    end_geometry=end_geometry,
    t_shell_mm=mech["t_shell_design_mm"],
    t_head_mm=mech["t_head_design_mm"],
    density_kg_m3=steel_density,
)

# ── Block 4: Media geometry (uses real_id) ─────────────────────────────────
geo_rows, base = [], []
curr_h = nozzle_plate_h

if nozzle_plate_h > 0:
    a0  = segment_area(0, real_id)
    a1  = segment_area(nozzle_plate_h, real_id)
    v_c = (a1 - a0) * cyl_len
    v_e = (dish_volume(nozzle_plate_h, real_id, h_dish, end_geometry) -
           dish_volume(0,              real_id, h_dish, end_geometry)) * 2
    tot = v_c + v_e
    geo_rows.append(["Nozzle Plate", nozzle_plate_h,
                     tot / nozzle_plate_h if nozzle_plate_h else 0,
                     v_c, v_e, tot])

for L in layers:
    h1, h2 = curr_h, curr_h + L["Depth"]
    v_c = (segment_area(h2, real_id) - segment_area(h1, real_id)) * cyl_len
    v_e = (dish_volume(h2, real_id, h_dish, end_geometry) -
           dish_volume(h1, real_id, h_dish, end_geometry)) * 2
    vol  = v_c + v_e
    area = vol / L["Depth"] if L["Depth"] > 0 else 0
    base.append({**L, "Vol": vol, "Area": area})
    geo_rows.append([L["Type"], L["Depth"], area, v_c, v_e, vol])
    curr_h = h2

avg_area     = sum(b["Area"] for b in base) / len(base) if base else 1.0
q_per_filter = (total_flow / streams) / n_filters

# ── Block 4: Pressure drop (Ergun) uses BW water properties ───────────────
bw_dp = pressure_drop(
    layers=layers,
    q_filter_m3h=q_per_filter,
    avg_area_m2=avg_area,
    solid_loading_kg_m2=solid_loading,
    captured_density_kg_m3=captured_solids_density,
    water_temp_c=feed_temp,
    rho_water=rho_feed,
    alpha_m_kg=alpha_specific,
    dp_trigger_bar=dp_trigger_bar,
)
# Use dirty ΔP as the nozzle plate design ΔP (auto-wired, no manual input needed)
np_dp_auto = bw_dp["dp_dirty_bar"]

# ── Block 3: Nozzle plate (uses dirty ΔP and real_id) ─────────────────────
wt_np = nozzle_plate_design(
    vessel_id_m=real_id,
    cyl_len_m=cyl_len,
    h_dish_m=h_dish,
    h_plate_m=nozzle_plate_h,
    design_dp_bar=np_dp_auto,
    media_layers=layers,
    water_density_kg_m3=rho_feed,
    nozzle_density_per_m2=np_density,
    bore_diameter_mm=np_bore_dia,
    beam_spacing_mm=np_beam_sp,
    allowable_stress_kgf_cm2=float(mech["allowable_stress"]),
    corrosion_allowance_mm=corrosion,
    density_kg_m3=steel_density,
    override_thickness_mm=np_override_t,
)

# ── Block 6: Nozzle schedule ───────────────────────────────────────────────
nozzle_sched = estimate_nozzle_schedule(
    q_filter_m3h=q_per_filter,
    bw_velocity_ms=bw_velocity,
    area_filter_m2=avg_area,
    default_rating=default_rating,
    stub_length_mm=float(nozzle_stub_len),
)

# ── Block 6: Supports ──────────────────────────────────────────────────────
wt_sup = saddle_weight(
    vessel_od_m=mech["od_m"],
    support_type=support_type,
    saddle_height_m=saddle_h,
    leg_height_m=leg_h,
    leg_section_mm=leg_section,
    base_plate_thickness_mm=base_plate_t,
    gusset_thickness_mm=gusset_t,
    density_kg_m3=steel_density,
)

# ── Block 5: Backwash ──────────────────────────────────────────────────────
bw_hyd = backwash_hydraulics(
    filter_area_m2=avg_area,
    bw_rate_m_h=bw_velocity,
    air_scour_rate_m_h=air_scour_rate,
    filtration_flow_m3h=q_per_filter,
)

bw_col = collector_check_ext(
    layers=layers,
    nozzle_plate_h_m=nozzle_plate_h,
    collector_h_m=collector_h,
    bw_velocity_m_h=bw_velocity,
    water_temp_c=bw_temp,
    rho_water=rho_bw,
    min_freeboard_m=freeboard_mm / 1000.0,
)

bw_exp = bed_expansion(
    layers=layers,
    bw_velocity_m_h=bw_velocity,
    water_temp_c=bw_temp,
    rho_water=rho_bw,
)

bw_seq = bw_sequence(
    filter_area_m2=avg_area,
    tss_scenarios=[tss_low, tss_avg, tss_high],
    n_filters_total=streams * n_filters,
    bw_per_day_per_filter=bw_cycles_day,
)

# ── Block 6: Internals weight ─────────────────────────────────────────────
wt_int = internals_weight(
    n_strainer_nozzles=wt_np.get("n_bores", 0),
    strainer_material=strainer_mat,
    air_header_dn_mm=int(air_header_dn),
    cyl_len_m=cyl_len,
    manhole_dn=manhole_dn,
    n_manholes=n_manholes,
)

# ── TSS mass balance ────────────────────────────────────────────────────────
run_time_h = bw_seq.get("run_time_h", 24.0)
waste_vol   = bw_seq.get("waste_vol_avg_m3", 1.0)

def _tss_bal(tss_mg_l):
    m_sol = tss_mg_l * q_per_filter * run_time_h / 1000.0
    w_tss = (m_sol * 1e3) / waste_vol if waste_vol > 0 else 0.0
    m_day = m_sol * (streams * n_filters) * bw_cycles_day
    return round(m_sol, 1), round(w_tss, 0), round(m_day, 0)

m_sol_low,  w_tss_low,  m_daily_low  = _tss_bal(tss_low)
m_sol_avg,  w_tss_avg,  m_daily_avg  = _tss_bal(tss_avg)
m_sol_high, w_tss_high, m_daily_high = _tss_bal(tss_high)

# ── Filtration cycle (DP-trigger based, design temperature) ──────────────
_load_data_cyc = filter_loading(total_flow, streams, n_filters, redundancy)
filt_cycles: dict = {}
for _x, _nact, _q in _load_data_cyc:
    _sc = "N" if _x == 0 else f"N-{_x}"
    filt_cycles[_sc] = filtration_cycle(
        layers=layers,
        q_filter_m3h=_q,
        avg_area_m2=avg_area,
        solid_loading_kg_m2=solid_loading,
        captured_density_kg_m3=captured_solids_density,
        water_temp_c=feed_temp,
        rho_water=rho_feed,
        dp_trigger_bar=dp_trigger_bar,
        alpha_m_kg=alpha_specific,
        tss_mg_l_list=[tss_low, tss_avg, tss_high],
    )

# ── Filtration cycle matrix: TSS × temperature, α fixed at design value ──
# α is locked to the design-temperature calibration so that temperature
# variation reflects real viscosity effect on cycle duration.
_alpha_fixed  = filt_cycles["N"]["alpha_used_m_kg"] if filt_cycles else 0.0
_tss_labels   = [f"Low ({tss_low:.0f} mg/L)",
                 f"Avg ({tss_avg:.0f} mg/L)",
                 f"High ({tss_high:.0f} mg/L)"]
_tss_vals     = [tss_low, tss_avg, tss_high]
_temp_vals    = [temp_low, feed_temp, temp_high]
_temp_labels  = [f"Min ({temp_low:.0f}°C)",
                 f"Design ({feed_temp:.0f}°C)",
                 f"Max ({temp_high:.0f}°C)"]
# cycle_matrix[sc_label][temp_label] = filtration_cycle result
cycle_matrix: dict = {}
for _x, _nact, _q in _load_data_cyc:
    _sc = "N" if _x == 0 else f"N-{_x}"
    cycle_matrix[_sc] = {}
    for _tv, _tl in zip(_temp_vals, _temp_labels):
        cycle_matrix[_sc][_tl] = filtration_cycle(
            layers=layers,
            q_filter_m3h=_q,
            avg_area_m2=avg_area,
            solid_loading_kg_m2=solid_loading,
            captured_density_kg_m3=captured_solids_density,
            water_temp_c=_tv,
            rho_water=rho_feed,
            dp_trigger_bar=dp_trigger_bar,
            alpha_m_kg=_alpha_fixed,
            tss_mg_l_list=_tss_vals,
        )

# ── BW scheduling & feasibility ──────────────────────────────────────────
import math as _math
_bw_dur_h = bw_total_min / 60.0   # BW duration in hours

# feasibility_matrix[sc_label][temp_label][tss_label] = dict of KPIs
def _feas_kpis(t_cycle_h, bw_dur_h, n_active_per_stream, n_streams):
    """Return operational KPIs for one (cycle_time, BW_duration, n_filters) set.

    BW trains are plant-wide; redundancy (standby filters) is per stream.
    """
    t_total   = t_cycle_h + bw_dur_h          # full filtration + BW period
    avail_pct = t_cycle_h / t_total * 100 if t_total > 0 else 0.0
    bw_per_day = 24.0 / t_total if t_total > 0 else 0.0
    # Plant-wide simultaneous BW demand
    n_active_total = n_active_per_stream * n_streams
    sim_demand = n_active_total * bw_dur_h / t_total if t_total > 0 else 0.0
    bw_trains  = max(1, _math.ceil(sim_demand))

    # Feasibility score
    if avail_pct >= 90 and bw_trains <= 1 and t_cycle_h >= 6:
        score, flag = "🟢 Good", "OK"
    elif avail_pct >= 80 and bw_trains <= 2 and t_cycle_h >= 3:
        score, flag = "🟡 Caution", "Review"
    else:
        score, flag = "🔴 Critical", "Redesign"

    return {
        "t_cycle_h":        round(t_cycle_h,        2),
        "avail_pct":        round(avail_pct,         1),
        "bw_per_day":       round(bw_per_day,        1),
        "sim_demand":       round(sim_demand,         2),
        "n_active_total":   n_active_total,
        "n_active_stream":  n_active_per_stream,
        "bw_trains":        bw_trains,
        "score":            score,
        "flag":             flag,
    }

feasibility_matrix: dict = {}
for _x, _nact, _q in _load_data_cyc:
    _sc = "N" if _x == 0 else f"N-{_x}"
    feasibility_matrix[_sc] = {}
    for _t_lbl in _temp_labels:
        feasibility_matrix[_sc][_t_lbl] = {}
        for _tss_lbl, _tss_v in zip(_tss_labels, _tss_vals):
            _cyc_t = cycle_matrix[_sc][_t_lbl]
            _tr    = next((r for r in _cyc_t["tss_results"]
                           if r["TSS (mg/L)"] == _tss_v), None)
            _t_cyc = _tr["Cycle duration (h)"] if _tr else 0.0
            feasibility_matrix[_sc][_t_lbl][_tss_lbl] = _feas_kpis(
                _t_cyc, _bw_dur_h, _nact, streams
            )

# ── Cartridge design ──────────────────────────────────────────────────────
_cart_mu_cP = mu_feed * 1000.0   # Pa·s → cP

cart_result = cartridge_design(
    design_flow_m3h=cart_flow,
    element_size=cart_size,
    rating_um=cart_rating,
    mu_cP=_cart_mu_cP,
    n_elem_per_housing=cart_housing,
    is_CIP_system=cart_cip,
    cf_inlet_tss_mg_l=cf_inlet_tss,
    cf_outlet_tss_mg_l=cf_outlet_tss,
)

cart_optim = cartridge_optimise(
    design_flow_m3h=cart_flow,
    rating_um=cart_rating,
    mu_cP=_cart_mu_cP,
    is_CIP_system=cart_cip,
)

# ── Hydraulic profile & energy ────────────────────────────────────────────
# Strainer nozzle plate ΔP: user-specified at design LV (vendor data).
# The plate bore (50 mm) is the nozzle body hole — actual ΔP is set by
# the fine slots on the nozzle head, which vary by manufacturer.
_np_dp_bar = np_slot_dp
hyd_prof = hydraulic_profile(
    dp_media_clean_bar  = bw_dp["dp_clean_bar"],
    dp_media_dirty_bar  = bw_dp["dp_dirty_bar"],
    np_dp_filt_bar      = _np_dp_bar,
    distributor_dp_bar  = dp_dist,
    dp_inlet_pipe_bar   = dp_inlet_pipe,
    dp_outlet_pipe_bar  = dp_outlet_pipe,
    p_residual_bar      = p_residual,
    static_head_m       = static_head,
    rho_feed_kg_m3      = rho_feed,
)

# Design scenario N at design temp / avg TSS for energy basis
_n_feas = feasibility_matrix.get("N", {}).get(
    f"Design ({feed_temp:.0f}°C)", {}).get(
    f"Avg ({tss_avg:.0f} mg/L)", {})
_bw_per_day_design = _n_feas.get("bw_per_day", 24.0 / (bw_total_min/60.0 + 1.0))
_avail_design      = _n_feas.get("avail_pct",  90.0)
_n_total_filters   = streams * n_filters

energy = energy_summary(
    q_filter_m3h        = q_per_filter,
    n_filters_total     = _n_total_filters,
    filt_head_dirty_mwc = hyd_prof["dirty"]["total_mwc"],
    filt_head_clean_mwc = hyd_prof["clean"]["total_mwc"],
    pump_eta            = pump_eta,
    motor_eta           = motor_eta,
    rho_feed_kg_m3      = rho_feed,
    q_bw_m3h            = bw_hyd["q_bw_design_m3h"],
    bw_head_mwc         = bw_head_mwc,
    bw_pump_eta         = bw_pump_eta,
    bw_motor_eta        = motor_eta,
    rho_bw_kg_m3        = rho_bw,
    p_blower_kw         = bw_hyd["p_blower_est_kw"],
    blower_motor_eta    = motor_eta,
    bw_duration_h       = _bw_dur_h,
    bw_per_day_design   = _bw_per_day_design,
    availability_pct    = _avail_design,
    elec_tariff_usd_kwh = elec_tariff,
    op_hours_per_year   = float(op_hours_yr),
)

# ── BW system equipment sizing ────────────────────────────────────────────
_n_bw_systems = _n_feas.get("bw_trains", 1)
bw_sizing = bw_system_sizing(
    q_bw_design_m3h     = bw_hyd["q_bw_design_m3h"],
    bw_head_mwc         = bw_head_mwc,
    bw_pump_eta         = bw_pump_eta,
    motor_eta           = motor_eta,
    q_air_design_m3h    = bw_hyd["q_air_design_m3h"],
    vessel_pressure_bar = vessel_pressure_bar,
    filter_id_m         = real_id,
    blower_inlet_temp_c = blower_inlet_temp_c,
    blower_eta          = blower_eta,
    bw_vol_per_cycle_m3 = bw_seq["total_vol_avg_m3"],
    n_bw_systems        = _n_bw_systems,
    tank_sf             = tank_sf,
    rho_bw_kg_m3        = rho_bw,
)

# ── Consolidated weight ────────────────────────────────────────────────────
nozzle_wt_total = sum(r.get("Total wt (kg)", 0) for r in nozzle_sched)
w_body  = wt_body["weight_body_kg"]
w_np    = wt_np["weight_total_kg"]
w_sup   = wt_sup["weight_all_supports_kg"]
w_noz   = nozzle_wt_total
w_int   = wt_int["weight_internals_kg"]
w_total = w_body + w_np + w_sup + w_noz + w_int

# ── Internal surface areas & lining / coating (needed before operating weight) ─
vessel_areas = internal_surface_areas(
    vessel_id_m          = real_id,
    cyl_len_m            = cyl_len,
    h_dish_m             = h_dish,
    end_type             = end_geometry,
    nozzle_plate_area_m2 = wt_np.get("area_total_m2", 0.0),
)

lining_result = lining_cost(
    protection_type      = protection_type,
    areas                = vessel_areas,
    rubber_type          = rubber_type_sel,
    rubber_thickness_mm  = lining_mm if lining_mm > 0 else 4.0,
    rubber_layers        = rubber_layers,
    rubber_cost_m2       = rubber_cost_m2,
    rubber_labor_m2      = rubber_labor_m2,
    epoxy_type           = epoxy_type_sel,
    epoxy_dft_um         = epoxy_dft_um,
    epoxy_coats          = epoxy_coats,
    epoxy_cost_m2        = epoxy_cost_m2,
    epoxy_labor_m2       = epoxy_labor_m2,
    ceramic_type         = ceramic_type_sel,
    ceramic_dft_um       = ceramic_dft_um,
    ceramic_coats        = ceramic_coats,
    ceramic_cost_m2      = ceramic_cost_m2,
    ceramic_labor_m2     = ceramic_labor_m2,
)

wt_oper = operating_weight(
    layers           = layers,
    avg_area_m2      = avg_area,
    vessel_id_m      = real_id,
    cyl_len_m        = cyl_len,
    h_dish_m         = h_dish,
    end_type         = end_geometry,
    w_empty_kg       = w_total,
    n_supports       = wt_sup["n_supports"],
    rho_water_kg_m3  = rho_feed,
    w_lining_kg      = lining_result["weight_kg"],
)

wt_saddle = saddle_design(
    total_length_m    = total_length,
    vessel_od_m       = mech["od_m"],
    vessel_id_m       = real_id,
    w_operating_kg    = wt_oper["w_operating_kg"],
    n_saddles         = wt_sup["n_supports"],
    contact_angle_deg = saddle_contact_angle,
)

# ── Economics ────────────────────────────────────────────────────────────────
_n_total_vessels = streams * n_filters

# Build per-type media inventory (kg, all vessels) and unit costs (USD/kg)
_media_inventory: dict = {}
_media_usd_kg:   dict = {}
_media_co2_kg:   dict = {}
for _b in base:
    _mt   = _b["Type"]
    _mkg  = _b["Vol"] * _b["rho_p_eff"] * _n_total_vessels
    _media_inventory[_mt] = _media_inventory.get(_mt, 0.0) + _mkg
    _is_grav = "Gravel" in _mt
    _is_anth = "Anthracite" in _mt
    _media_usd_kg[_mt]  = (econ_media_gravel if _is_grav
                           else econ_media_anthracite if _is_anth
                           else econ_media_sand) / 1000.0
    _media_co2_kg[_mt]  = (media_co2_gravel if _is_grav
                           else media_co2_anthracite if _is_anth
                           else media_co2_sand)

econ_capex = capex_breakdown(
    weight_total_kg        = w_total,
    n_vessels              = _n_total_vessels,
    steel_cost_usd_kg      = steel_cost_usd_kg,
    erection_usd           = erection_usd_vessel,
    piping_usd             = piping_usd_vessel,
    instrumentation_usd    = instrumentation_usd_vessel,
    civil_usd              = civil_usd_vessel,
    engineering_pct        = engineering_pct,
    contingency_pct        = contingency_pct,
)

econ_opex = opex_annual(
    filtration_power_kw        = energy["p_filt_avg_kw"],
    bw_power_kw                = energy["p_bw_kw"],
    blower_power_kw            = energy["p_blower_elec_kw"],
    n_vessels                  = _n_total_vessels,
    electricity_tariff         = elec_tariff,
    operating_hours            = float(op_hours_yr),
    media_inventory_kg_by_type = _media_inventory,
    media_costs_by_type        = _media_usd_kg,
    media_interval_years       = media_replace_years,
    n_strainer_nozzles         = wt_np.get("n_bores", 0) * _n_total_vessels,
    nozzle_cost_usd            = nozzle_unit_cost,
    nozzle_interval_years      = nozzle_replace_years,
    labour_usd_per_filter_year = labour_usd_filter_yr,
    n_filters_total            = _n_total_vessels,
    chemical_cost_usd_m3       = chemical_cost_m3,
    total_flow_m3h             = total_flow,
)

econ_carbon = carbon_footprint(
    filtration_power_kw    = energy["p_filt_avg_kw"],
    bw_power_kw            = energy["p_bw_kw"],
    blower_power_kw        = energy["p_blower_elec_kw"],
    operating_hours        = float(op_hours_yr),
    grid_intensity_kg_kwh  = grid_intensity,
    weight_steel_kg        = w_total * _n_total_vessels,
    steel_carbon_kg_kg     = steel_carbon_kg,
    weight_concrete_kg     = 0.0,
    concrete_carbon_kg_kg  = concrete_carbon_kg,
    media_mass_by_type_kg  = _media_inventory,
    media_carbon_by_type   = _media_co2_kg,
    design_life_years      = int(design_life_years),
    total_flow_m3h         = total_flow,
)

econ_bench = global_benchmark_comparison(
    capex_total_usd    = econ_capex["total_capex_usd"],
    opex_usd_year      = econ_opex["total_opex_usd_yr"],
    total_flow_m3h     = total_flow,
    n_filters          = _n_total_vessels,
    design_life_years  = int(design_life_years),
    co2_per_m3         = econ_carbon["co2_per_m3_operational"],
    electricity_tariff = elec_tariff,
    operating_hours    = float(op_hours_yr),
)

# ══════════════════════════════════════════════════════════════════════════════
# STATUS BADGES  (appended below the tab group in the context column)
# ══════════════════════════════════════════════════════════════════════════════
with ctx:
    st.divider()
    _status_items = {
        "Project":    bool(project_name),
        "Process":    total_flow > 0 and n_filters > 0,
        "Water":      feed_sal >= 0 and feed_temp > 0,
        "Geometry":   nominal_id > 0 and total_length > 0,
        "Mechanical": design_pressure > 0,
        "Media":      len(layers) > 0 and all(L["Depth"] > 0 for L in layers),
        "Backwash":   not bw_col["media_loss_risk"],
        "Weight":     w_total > 0,
    }
    _cols_status = st.columns(2)
    for i, (label, done) in enumerate(_status_items.items()):
        icon = "🟢" if done else ("🔴" if label == "Backwash"
                                   and bw_col["media_loss_risk"] else "⚪")
        _cols_status[i % 2].markdown(f"{icon} {label}")

    if bw_col["media_loss_risk"]:
        st.warning(f"⚠️ Media carryover risk — max safe BW rate: "
                   f"{bw_col['max_safe_bw_m_h']:.1f} m/h")
    st.caption("AQUASIGHT™ | Proprietary")

# ── Engineering alert helper ───────────────────────────────────────────────
def show_alert(level: str, title: str, message: str) -> None:
    _icons  = {"info": "🔵", "advisory": "🟡", "warning": "🟠", "critical": "🔴"}
    _colors = {
        "info":     "#0d1f35",
        "advisory": "#2a2000",
        "warning":  "#2a1200",
        "critical": "#2a0000",
    }
    _borders = {
        "info":     "#1a3a5c",
        "advisory": "#b8860b",
        "warning":  "#cc5500",
        "critical": "#cc0000",
    }
    st.markdown(
        f"""<div style='padding:10px 14px;border-radius:6px;
        background:{_colors[level]};border-left:4px solid {_borders[level]};
        margin:6px 0;'>
        {_icons[level]} <b>{title}</b><br>
        <span style='font-size:0.9em;opacity:0.92'>{message}</span>
        </div>""",
        unsafe_allow_html=True,
    )

# ══════════════════════════════════════════════════════════════════════════════
# CONTENT TABS
# ══════════════════════════════════════════════════════════════════════════════
with main:
    (tab_proj, tab_proc, tab_water, tab_vessel,
     tab_media, tab_bw, tab_weight, tab_cart, tab_energy,
     tab_section, tab_datasheet, tab_report, tab_econ) = st.tabs([
        "📋 Project",
        "💧 Process",
        "🌊 Water",
        "🏗️ Vessel",
        "🧱 Media",
        "🔄 Backwash",
        "⚖️ Weight",
        "🔷 Cartridge",
        "⚡ Energy",
        "🖼️ Section",
        "📋 Datasheet",
        "📄 Report",
        "💰 Economics",
    ])

    # ─────────────────────────────────────────────────────────────────────
    # TAB 1 · PROJECT
    # ─────────────────────────────────────────────────────────────────────
    with tab_proj:
        st.subheader("Project information")
        c1, c2 = st.columns(2)
        with c1:
            st.table(pd.DataFrame([
                ["Project",     project_name],
                ["Document",    doc_number],
                ["Revision",    revision],
                ["Client",      client or "—"],
                ["Prepared by", engineer],
            ], columns=["Field", "Value"]))
        with c2:
            st.markdown("**Scope**")
            st.info(
                "1. Horizontal MMF process sizing\n"
                "2. Water properties — feed and BW\n"
                "3. Vessel geometry and ASME mechanical\n"
                "4. Media volumes, EBCT, LV, ΔP, inventory\n"
                "5. Backwash: collector check, expansion,\n"
                "   pump/blower, sequence, waste volumes\n"
                "6. Empty weight — complete vessel\n"
                "7. Report export *(pending)*"
            )

    # ─────────────────────────────────────────────────────────────────────
    # TAB 2 · PROCESS
    # ─────────────────────────────────────────────────────────────────────
    with tab_proc:
        st.subheader("Process — filter loading")

        load_data = filter_loading(total_flow, streams, n_filters, redundancy)

        # ── Pre-compute all scenario / layer KPIs for assessment ──────────
        def _lv_severity(vel, threshold):
            if vel <= threshold:
                return None
            ovr = (vel - threshold) / threshold
            if ovr <= 0.05:  return "advisory"
            if ovr <= 0.15:  return "warning"
            return "critical"

        def _ebct_severity(ebct, threshold):
            if ebct >= threshold:
                return None
            short = (threshold - ebct) / threshold
            if short <= 0.10: return "advisory"
            if short <= 0.25: return "warning"
            return "critical"

        _all_lv_issues   = []   # (scenario_label, layer, severity, vel)
        _all_ebct_issues = []   # (scenario_label, layer, severity, ebct)
        _n_criticals = 0
        _n_warnings  = 0
        _n_advisories = 0

        for _x, _a, _q in load_data:
            _sc = "N" if _x == 0 else f"N-{_x}"
            for _b in base:
                _vel  = _q / _b["Area"] if _b["Area"] > 0 else 0
                _ebct = (_b["Vol"] / _q) * 60 if _q > 0 else 0
                _lv_sev  = _lv_severity(_vel, velocity_threshold)
                _eb_sev  = _ebct_severity(_ebct, ebct_threshold)
                if _lv_sev:
                    _all_lv_issues.append((_sc, _b["Type"], _lv_sev, _vel))
                    if _lv_sev == "critical":   _n_criticals += 1
                    elif _lv_sev == "warning":  _n_warnings  += 1
                    else:                        _n_advisories += 1
                if _eb_sev:
                    _all_ebct_issues.append((_sc, _b["Type"], _eb_sev, _ebct))
                    if _eb_sev == "critical":   _n_criticals += 1
                    elif _eb_sev == "warning":  _n_warnings  += 1
                    else:                        _n_advisories += 1

        # ── Overall risk level ────────────────────────────────────────────
        _n_scenario_criticals = sum(
            1 for s, _, sv, __ in (_all_lv_issues + _all_ebct_issues)
            if sv == "critical" and s == "N"
        )
        _n_scenario_warnings = sum(
            1 for s, _, sv, __ in (_all_lv_issues + _all_ebct_issues)
            if sv == "warning" and s == "N"
        )
        if _n_criticals == 0 and _n_warnings == 0 and _n_advisories <= 1:
            _overall_risk = "STABLE"
            _risk_color   = "#0a2a0a"
            _risk_border  = "#1a7a1a"
            _risk_icon    = "🟢"
        elif _n_scenario_criticals > 1 or (_n_criticals > 0 and _n_warnings > 2):
            _overall_risk = "CRITICAL"
            _risk_color   = "#2a0000"
            _risk_border  = "#cc0000"
            _risk_icon    = "🔴"
        elif _n_scenario_criticals > 0 or _n_warnings >= 3:
            _overall_risk = "ELEVATED"
            _risk_color   = "#2a1200"
            _risk_border  = "#cc5500"
            _risk_icon    = "🟠"
        elif _n_scenario_warnings > 0 or _n_criticals > 0:
            _overall_risk = "MARGINAL"
            _risk_color   = "#2a2000"
            _risk_border  = "#b8860b"
            _risk_icon    = "🟡"
        else:
            _overall_risk = "STABLE"
            _risk_color   = "#0a2a0a"
            _risk_border  = "#1a7a1a"
            _risk_icon    = "🟢"

        # Key drivers for the assessment panel
        _drivers = []
        if _all_lv_issues:
            _worst_lv = max(_all_lv_issues, key=lambda t: t[3])
            _drivers.append(
                f"Filtration velocity reaches {_worst_lv[3]:.2f} m/h "
                f"(recommended envelope upper limit {velocity_threshold:.1f} m/h) "
                f"in scenario {_worst_lv[0]}, layer {_worst_lv[1]}."
            )
        if _all_ebct_issues:
            _worst_eb = min(_all_ebct_issues, key=lambda t: t[3])
            _drivers.append(
                f"Contact time reduces to {_worst_eb[3]:.2f} min "
                f"(recommended lower limit {ebct_threshold:.1f} min) "
                f"in scenario {_worst_eb[0]}, layer {_worst_eb[1]}."
            )
        if not _drivers:
            _drivers.append("All hydraulic parameters remain within the recommended operating envelope across all evaluated scenarios.")

        _impacts = {
            "STABLE":   [
                "Filter performance expected to be consistent across the full redundancy range.",
                "Particulate capture efficiency maintains design margin under peak loading.",
                "No hydraulic adjustments indicated at current configuration.",
            ],
            "MARGINAL": [
                "Performance remains acceptable under normal N-scenario operation.",
                "One or more standby scenarios approach the hydraulic envelope boundary — review BW cycle frequency.",
                "Minor adjustments to filter area or media depth may improve N-1 resilience.",
            ],
            "ELEVATED": [
                "Elevated velocity or reduced contact time may compromise particulate capture under peak hydraulic loading.",
                "Run time between backwash cycles likely shortened by 15–30% compared to design basis.",
                "Consider increasing number of operating filters or filter area to restore operating margin.",
            ],
            "CRITICAL": [
                "Hydraulic loading significantly exceeds the recommended operating envelope.",
                "Risk of particulate breakthrough and accelerated media fouling under sustained operation.",
                "System redesign recommended — increase filter area, add filter units, or reduce total flow per vessel.",
            ],
        }
        _recommendations = {
            "STABLE":   "Maintain current configuration. Review again if flow demand increases or media condition degrades.",
            "MARGINAL": "Review N-1 scenario performance with the client and confirm acceptance of reduced margin during filter outage.",
            "ELEVATED": "Increase filter area or reduce per-filter hydraulic loading. Adding one filter unit per stream typically resolves elevated ratings.",
            "CRITICAL": "Redesign required. The current number of filters or vessel area is insufficient for the specified flow. Consult process design basis before proceeding.",
        }

        # ── Part 4: Combined Process Assessment Panel ─────────────────────
        st.markdown(
            f"""<div style='padding:16px 20px;border-radius:8px;
            background:{_risk_color};border-left:5px solid {_risk_border};
            margin-bottom:16px;'>
            <div style='font-size:1.15em;font-weight:bold;margin-bottom:8px'>
            Overall Process Assessment &nbsp; {_risk_icon} {_overall_risk}</div>
            <div style='margin-bottom:8px'><b>Key drivers:</b><br>
            {"<br>".join("· " + d for d in _drivers)}</div>
            <div style='margin-bottom:8px'><b>Expected operational impact:</b><br>
            {"<br>".join("· " + i for i in _impacts[_overall_risk])}</div>
            <div><b>Recommendation:</b><br>
            {_recommendations[_overall_risk]}</div>
            </div>""",
            unsafe_allow_html=True,
        )

        # ── Flow distribution table ───────────────────────────────────────
        st.markdown("#### Flow distribution by scenario")
        comp = []
        for x, a, q in load_data:
            lv = q / avg_area if avg_area > 0 else 0
            _lv_flag = "Within envelope" if lv <= velocity_threshold else (
                "Approaching limit" if lv <= velocity_threshold * 1.05 else "Outside envelope")
            comp.append({
                "Scenario":             "N" if x == 0 else f"N-{x}",
                "Active filters":       a,
                "Flow / filter (m³/h)": round(q, 2),
                "LV (m/h)":             round(lv, 2),
                "Hydraulic status":     _lv_flag,
            })
        st.dataframe(pd.DataFrame(comp), use_container_width=True, hide_index=True)

        # ── Part 6: Grouped warnings by cause ────────────────────────────
        st.markdown("#### Operating envelope review by scenario")

        for x, a, q in load_data:
            label = "N (normal)" if x == 0 else f"N-{x}"
            _sc_lbl = "N" if x == 0 else f"N-{x}"
            with st.expander(f"Scenario {label} — {q:.1f} m³/h / filter",
                             expanded=(x == 0)):
                rows = []
                _sc_lv_issues   = []
                _sc_ebct_issues = []
                for b in base:
                    vel  = q / b["Area"] if b["Area"] > 0 else 0
                    ebct = (b["Vol"] / q) * 60 if q > 0 else 0
                    _lv_sev = _lv_severity(vel, velocity_threshold)
                    _eb_sev = _ebct_severity(ebct, ebct_threshold)
                    _lv_env = ("Within envelope" if not _lv_sev
                               else "Approaching limit" if _lv_sev == "advisory"
                               else "Outside envelope")
                    _eb_env = ("Within envelope" if not _eb_sev
                               else "Approaching limit" if _eb_sev == "advisory"
                               else "Outside envelope")
                    rows.append({
                        "Layer":        b["Type"],
                        "Area (m²)":    round(b["Area"], 3),
                        "LV (m/h)":     round(vel, 2),
                        "LV envelope":  _lv_env,
                        "EBCT (min)":   round(ebct, 2),
                        "EBCT envelope": _eb_env,
                    })
                    if _lv_sev:
                        _sc_lv_issues.append((b["Type"], _lv_sev, vel))
                    if _eb_sev:
                        _sc_ebct_issues.append((b["Type"], _eb_sev, ebct))

                st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

                # Group 1: Hydraulic Loading
                if _sc_lv_issues:
                    with st.expander(f"🟠 Hydraulic Loading — {len(_sc_lv_issues)} layer(s) outside envelope"):
                        for _layer, _sev, _vel in _sc_lv_issues:
                            show_alert(
                                _sev,
                                f"{_layer}: filtration velocity {_vel:.2f} m/h",
                                "Elevated filtration velocity increases risk of media disturbance "
                                "and localised particulate breakthrough. "
                                "Consider increasing filter area or reducing per-filter loading.",
                            )

                # Group 2: Contact Time
                if _sc_ebct_issues:
                    with st.expander(f"🟡 Contact Time — {len(_sc_ebct_issues)} layer(s) below design target"):
                        for _layer, _sev, _ebct in _sc_ebct_issues:
                            show_alert(
                                _sev,
                                f"{_layer}: contact time {_ebct:.2f} min",
                                "Reduced contact time may compromise particulate capture "
                                "stability under peak hydraulic loading. "
                                "Extended run times or deeper media beds recommended.",
                            )

                if not _sc_lv_issues and not _sc_ebct_issues:
                    st.success("All layers operate within the recommended hydraulic envelope for this scenario.")

        st.divider()
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("LV — N scenario",            f"{q_per_filter/avg_area:.2f} m/h")
        m2.metric("Flow / filter (N)",           f"{q_per_filter:.1f} m³/h")
        m3.metric("Total filters",               f"{streams * n_filters}")
        m4.metric("Recommended LV envelope",     f"≤ {velocity_threshold:.1f} m/h")

        # ── Part 5: Design Robustness Index ──────────────────────────────
        st.markdown("#### Design Robustness Index")
        st.caption(
            "Engineering summary across all evaluated redundancy scenarios. "
            "Intended as a one-page design review reference."
        )
        _rob_rows = []
        _all_scenarios = [("N", 0)] + [(f"N-{i}", i) for i in range(1, redundancy + 1)]
        _eval_set = {("N" if x == 0 else f"N-{x}"): q
                     for x, _, q in load_data}

        for _sc_name, _x in _all_scenarios:
            if _sc_name not in _eval_set:
                _rob_rows.append({
                    "Scenario":          _sc_name,
                    "Filtration rate":   "—",
                    "Hydraulic status":  "Not evaluated",
                    "EBCT status":       "Not evaluated",
                    "Overall":           "Not evaluated",
                })
                continue
            _q = _eval_set[_sc_name]
            _lv_n = _q / avg_area if avg_area > 0 else 0
            _worst_lv_sev = None
            _worst_eb_sev = None
            for _b in base:
                _v = _q / _b["Area"] if _b["Area"] > 0 else 0
                _e = (_b["Vol"] / _q) * 60 if _q > 0 else 0
                _sv = _lv_severity(_v, velocity_threshold)
                _se = _ebct_severity(_e, ebct_threshold)
                _sev_rank = {"critical": 3, "warning": 2, "advisory": 1, None: 0}
                if _sev_rank.get(_sv, 0) > _sev_rank.get(_worst_lv_sev, 0):
                    _worst_lv_sev = _sv
                if _sev_rank.get(_se, 0) > _sev_rank.get(_worst_eb_sev, 0):
                    _worst_eb_sev = _se

            def _sev_to_label(s):
                return ("Within envelope" if s is None
                        else "Approaching limit" if s == "advisory"
                        else "Outside envelope")

            _lv_label = _sev_to_label(_worst_lv_sev)
            _eb_label = _sev_to_label(_worst_eb_sev)

            _worst_overall = max(
                _worst_lv_sev or "", _worst_eb_sev or "",
                key=lambda s: {"critical": 3, "warning": 2, "advisory": 1, "": 0}.get(s, 0)
            )
            _overall_label = (
                "Stable" if not _worst_overall
                else "Marginal" if _worst_overall == "advisory"
                else "Sensitive" if _worst_overall == "warning"
                else "Critical"
            )

            _rob_rows.append({
                "Scenario":          _sc_name,
                "Filtration rate":   f"{_lv_n:.2f} m/h",
                "Hydraulic status":  _lv_label,
                "EBCT status":       _eb_label,
                "Overall":           _overall_label,
            })

        st.dataframe(pd.DataFrame(_rob_rows), use_container_width=True, hide_index=True)
        st.caption(
            f"Recommended operating envelope: LV ≤ {velocity_threshold:.1f} m/h  ·  "
            f"EBCT ≥ {ebct_threshold:.1f} min  ·  "
            "Stable = both within envelope · Marginal = approaching limit · "
            "Sensitive = one parameter outside envelope · Critical = both outside."
        )

    # ─────────────────────────────────────────────────────────────────────
    # TAB 3 · WATER PROPERTIES
    # ─────────────────────────────────────────────────────────────────────
    with tab_water:
        st.subheader("Water properties — feed & backwash")

        w1, w2 = st.columns(2)
        with w1:
            st.markdown("**Feed water**")
            st.table(pd.DataFrame([
                ["Temperature",   f"{feed_wp['temp_c']:.1f} °C"],
                ["Salinity",      f"{feed_wp['salinity_ppt']:.2f} ppt"],
                ["Density",       f"{feed_wp['density_kg_m3']:.3f} kg/m³"],
                ["Viscosity",     f"{feed_wp['viscosity_cp']:.4f} cP"],
                ["TDS (approx.)", f"{feed_wp['tds_mg_l']:,.0f} mg/L"],
            ], columns=["Property", "Value"]))

        with w2:
            st.markdown("**Backwash water**")
            st.table(pd.DataFrame([
                ["Temperature",   f"{bw_wp['temp_c']:.1f} °C"],
                ["Salinity",      f"{bw_wp['salinity_ppt']:.2f} ppt"],
                ["Density",       f"{bw_wp['density_kg_m3']:.3f} kg/m³"],
                ["Viscosity",     f"{bw_wp['viscosity_cp']:.4f} cP"],
                ["TDS (approx.)", f"{bw_wp['tds_mg_l']:,.0f} mg/L"],
            ], columns=["Property", "Value"]))

        st.info(
            "Water properties feed directly into: terminal velocity (u_t), "
            "minimum fluidisation velocity (u_mf), Ergun pressure drop (ΔP), "
            "and nozzle velocity checks. "
            "BW water properties govern the expansion and collector check calculations."
        )

    # ─────────────────────────────────────────────────────────────────────
    # TAB 4 · VESSEL
    # ─────────────────────────────────────────────────────────────────────
    with tab_vessel:
        st.subheader("Vessel geometry & mechanical")

        # ── Geometry summary ──
        with st.expander("1 · Geometry", expanded=True):
            g1, g2, g3, g4, g5 = st.columns(5)
            g1.metric("Nominal ID",   f"{nominal_id:.3f} m")
            g2.metric("Lining",       f"{lining_mm:.1f} mm")
            g3.metric("Real hyd. ID", f"{real_id:.4f} m",
                      delta=f"−{lining_mm*2:.1f} mm", delta_color="off")
            g4.metric("Cyl. length",  f"{cyl_len:.3f} m")
            g5.metric("Dish depth",   f"{h_dish:.3f} m")
            st.caption(
                f"Real hydraulic ID = {nominal_id:.3f} m − "
                f"2 × {lining_mm:.1f} mm = **{real_id:.4f} m**. "
                "This ID is used in all hydraulic calculations below."
            )

        # ── Mechanical ──
        with st.expander("2 · ASME wall thickness", expanded=True):
            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**Material & radiography**")
                st.table(pd.DataFrame([
                    ["Material",            material_name],
                    ["Standard",            mat_info["standard"]],
                    ["Allowable stress (S)", f"{mech['allowable_stress']} kg/cm²"],
                    ["Shell radiography",   f"{shell_radio}  →  E={mech['shell_E']:.2f}"],
                    ["Head radiography",    f"{head_radio}  →  E={mech['head_E']:.2f}"],
                ], columns=["Parameter", "Value"]))
            with c2:
                st.markdown("**Thickness results**")
                def fmt_t(t_min, t_des, overridden):
                    flag = " ✏️ overridden" if overridden else ""
                    return f"{t_des} mm{flag}  (t_min={t_min:.2f} mm)"
                st.table(pd.DataFrame([
                    ["Design pressure",
                     f"{design_pressure:.2f} bar ({mech['p_kgf_cm2']:.3f} kg/cm²)"],
                    ["Corrosion allowance",  f"{corrosion:.1f} mm"],
                    ["Shell t_min",          f"{mech['t_shell_min_mm']:.2f} mm"],
                    ["Shell t_design",
                     fmt_t(mech['t_shell_min_mm'], mech['t_shell_design_mm'],
                           mech.get('shell_overridden', False))],
                    ["Head t_min",           f"{mech['t_head_min_mm']:.2f} mm"],
                    ["Head t_design",
                     fmt_t(mech['t_head_min_mm'], mech['t_head_design_mm'],
                           mech.get('head_overridden', False))],
                    ["Nominal ID",           f"{mech['nominal_id_m']:.4f} m"],
                    ["Real hydraulic ID",    f"{real_id:.4f} m"],
                    ["Outside diameter",     f"{mech['od_m']:.4f} m"],
                ], columns=["Parameter", "Value"]))

        # ── Nozzle plate ──
        with st.expander("3 · Nozzle plate design", expanded=True):
            st.info(f"Design ΔP auto-wired from Ergun dirty-bed result: "
                    f"**{np_dp_auto:.5f} bar** "
                    f"({np_dp_auto*1e5/1000:.3f} kPa)")
            l1, l2, l3 = st.columns(3)
            l1.metric("ΔP hydraulic", f"{wt_np['q_dp_kpa']:.3f} kPa")
            l2.metric("Media load",   f"{wt_np['q_media_kpa']:.3f} kPa")
            l3.metric("Total load",   f"{wt_np['q_total_kpa']:.3f} kPa")

            c1, c2 = st.columns(2)
            with c1:
                st.markdown("**Geometry & bore layout**")
                st.table(pd.DataFrame([
                    ["Plate height",      f"{wt_np['h_plate_m']:.3f} m"],
                    ["Chord at plate",    f"{wt_np['chord_m']:.4f} m"],
                    ["Angle θ",          f"{wt_np['theta_deg']:.2f}°"],
                    ["Cyl. plate area",   f"{wt_np['area_cyl_m2']:.4f} m²"],
                    ["Dish ends area",    f"{wt_np['area_both_dish_m2']:.4f} m²"],
                    ["Total plate area",  f"{wt_np['area_total_m2']:.4f} m²"],
                    ["Number of bores",   str(wt_np["n_bores"])],
                    ["Bore diameter",     f"{wt_np['bore_diameter_mm']:.0f} mm"],
                    ["Nozzle density",    f"{wt_np['actual_density_per_m2']:.1f} /m²"],
                    ["Open area ratio",   f"{wt_np['open_ratio_pct']:.1f} %"],
                ], columns=["Parameter", "Value"]))
            with c2:
                st.markdown("**Thickness & support beams**")
                st.table(pd.DataFrame([
                    ["Beam spacing",      f"{wt_np['beam_spacing_mm']:.0f} mm"],
                    ["t_min (Roark)",     f"{wt_np['t_min_mm']:.2f} mm"],
                    ["t_design",         f"{wt_np['t_design_mm']} mm"],
                    ["t used",           f"{wt_np['t_used_mm']} mm  "
                                        f"({wt_np['thickness_source']})"],
                    ["Beam M_max",        f"{wt_np['M_max_kNm']:.1f} kN·m"],
                    ["Required Z",        f"{wt_np['beam_Z_req_cm3']:.0f} cm³"],
                    ["Selected section",  wt_np["beam_section"]],
                    ["No. of beams",      str(wt_np["n_beams"])],
                    ["Plate weight",      f"{wt_np['weight_plate_kg']:,.1f} kg"],
                    ["Beams weight",      f"{wt_np['weight_beams_kg']:,.1f} kg"],
                    ["Total plate assy.", f"{wt_np['weight_total_kg']:,.1f} kg"],
                ], columns=["Parameter", "Value"]))

    # ─────────────────────────────────────────────────────────────────────
    # TAB 5 · MEDIA
    # ─────────────────────────────────────────────────────────────────────
    with tab_media:
        st.subheader("Media design")

        with st.expander("1 · Geometric volumes", expanded=True):
            df_geo = pd.DataFrame(geo_rows, columns=[
                "Item", "Depth (m)", "Avg area (m²)",
                "V_cyl (m³)", "V_ends (m³)", "Total vol (m³)"])
            st.dataframe(df_geo.style.format({
                "Depth (m)":      "{:.3f}",
                "Avg area (m²)":  "{:.4f}",
                "V_cyl (m³)":     "{:.4f}",
                "V_ends (m³)":    "{:.4f}",
                "Total vol (m³)": "{:.4f}",
            }), use_container_width=True, hide_index=True)

        with st.expander("2 · Media properties", expanded=True):
            df_med = pd.DataFrame(base)[
                ["Type","Depth","Vol","Area","rho_p_eff","epsilon0","d10","cu"]
            ].rename(columns={
                "Type":"Media","Depth":"Depth (m)","Vol":"Vol (m³)",
                "Area":"Avg area (m²)","rho_p_eff":"ρ (kg/m³)",
                "epsilon0":"ε₀","d10":"d10 (mm)","cu":"CU"})
            st.dataframe(df_med, use_container_width=True, hide_index=True)

        with st.expander("3 · Pressure drop — clean/moderate/dirty (all scenarios)", expanded=True):
            st.caption(
                f"Clean ΔP: Ergun equation on virgin bed.  "
                f"Moderate = 50% loaded · Dirty = 100% loaded — cake model (Ruth): "
                f"ΔP_cake = α × μ × LV × M.  "
                f"α ({bw_dp['alpha_source']}) = "
                f"{bw_dp['alpha_used_m_kg']/1e9:.1f} × 10⁹ m/kg  |  "
                f"M_max = {solid_loading:.2f} kg/m²  |  "
                f"Feed: ρ={rho_feed:.1f} kg/m³, μ={mu_feed*1000:.4f} cP"
            )

            # Compute ΔP for all redundancy scenarios
            load_data_dp = filter_loading(total_flow, streams, n_filters, redundancy)
            dp_summary = []
            for x, n_act, q in load_data_dp:
                sc_label = "N" if x == 0 else f"N-{x}"
                sc_dp = pressure_drop(
                    layers=layers,
                    q_filter_m3h=q,
                    avg_area_m2=avg_area,
                    solid_loading_kg_m2=solid_loading,
                    captured_density_kg_m3=captured_solids_density,
                    water_temp_c=feed_temp,
                    rho_water=rho_feed,
                    alpha_m_kg=alpha_specific,
                    dp_trigger_bar=dp_trigger_bar,
                )
                dp_summary.append({
                    "Scenario":        sc_label,
                    "LV (m/h)":        sc_dp["u_m_h"],
                    "ΔP clean (bar)":  sc_dp["dp_clean_bar"],
                    "ΔP clean (mWC)":  sc_dp["dp_clean_mwc"],
                    "ΔP moderate (bar)": sc_dp["dp_moderate_bar"],
                    "ΔP dirty (bar)":  sc_dp["dp_dirty_bar"],
                    "ΔP dirty (mWC)":  sc_dp["dp_dirty_mwc"],
                })

            st.markdown("**Summary — all scenarios**")
            st.dataframe(pd.DataFrame(dp_summary),
                         use_container_width=True, hide_index=True)

            st.markdown("**Per-layer breakdown — N scenario**")
            st.dataframe(pd.DataFrame(bw_dp["layers"]),
                         use_container_width=True, hide_index=True)

            p1, p2, p3 = st.columns(3)
            p1.metric("ΔP clean (N)",
                      f"{bw_dp['dp_clean_bar']:.5f} bar",
                      delta=f"{bw_dp['dp_clean_mwc']:.3f} mWC", delta_color="off")
            p2.metric("ΔP moderate (N)",
                      f"{bw_dp['dp_moderate_bar']:.5f} bar",
                      delta=f"{bw_dp['dp_moderate_mwc']:.3f} mWC", delta_color="off")
            p3.metric("ΔP dirty → nozzle plate ΔP",
                      f"{bw_dp['dp_dirty_bar']:.5f} bar",
                      delta=f"{bw_dp['dp_dirty_mwc']:.3f} mWC", delta_color="off")

        with st.expander("4 · Media inventory", expanded=True):
            total_vessels = streams * n_filters
            inv_rows = []; total_mass = 0
            for b in base:
                mf = b["Vol"] * b["rho_p_eff"]
                mt = mf * total_vessels
                total_mass += mt
                inv_rows.append({
                    "Media":             b["Type"],
                    "d10/CU":            f"{b['d10']}/{b['cu']}",
                    "Vol/filter (m³)":   round(b["Vol"], 4),
                    "Mass/filter (kg)":  round(mf),
                    "Total mass (kg)":   round(mt),
                })
            st.dataframe(pd.DataFrame(inv_rows),
                         use_container_width=True, hide_index=True)
            i1, i2, i3 = st.columns(3)
            i1.metric("Total filters",  total_vessels)
            i2.metric("Total media",    f"{total_mass/1000:.2f} t")
            i3.metric("Per filter",
                      f"{total_mass/total_vessels/1000:.2f} t"
                      if total_vessels else "—")

        with st.expander("5 · Clogging analysis — N scenario", expanded=True):
            st.caption(
                f"Captured solids density: **{captured_solids_density:.0f} kg/m³**  |  "
                f"Total solid loading: **{solid_loading:.2f} kg/m²**"
            )
            clog_cols = ["Media", "Support", "Capture (%)",
                         "Solid load (kg/m²)", "Solid vol (m³/m²)",
                         "ΔεF", "Clogging (%)", "ε clean",
                         "Cake ΔP mod (bar)", "Cake ΔP dirty (bar)"]
            clog_df = pd.DataFrame(bw_dp["layers"])[clog_cols]
            st.dataframe(clog_df, use_container_width=True, hide_index=True)
            st.caption(
                "Support layers (e.g., Gravel) retain no solids. "
                "Cake ΔP = α × μ × LV × M, distributed by capture fraction. "
                "ΔεF shown for reference only — cake model, not voidage reduction, "
                "drives moderate/dirty ΔP."
            )

    # ─────────────────────────────────────────────────────────────────────
    # TAB 6 · BACKWASH
    # ─────────────────────────────────────────────────────────────────────
    with tab_bw:
        st.subheader("Backwash design")

        # Collector check — most important, shown first
        with st.expander("1 · Collector height check — media loss guard",
                         expanded=True):
            status_color = ("🔴" if bw_col["media_loss_risk"]
                            else "🟡" if "WARNING" in bw_col["status"]
                            else "🟢")
            st.markdown(f"### {status_color} {bw_col['status']}")

            cc1, cc2, cc3, cc4 = st.columns(4)
            cc1.metric("Settled bed top",  f"{bw_col['settled_top_m']:.3f} m")
            cc2.metric("Expanded bed top", f"{bw_col['expanded_top_m']:.3f} m")
            cc3.metric("Collector height", f"{bw_col['collector_h_m']:.3f} m")
            cc4.metric("Freeboard",
                       f"{bw_col['freeboard_m']:.3f} m",
                       delta=f"{bw_col['freeboard_pct']:.1f}% of bed",
                       delta_color="normal"
                       if bw_col["freeboard_m"] >= bw_col["min_freeboard_m"]
                       else "inverse")

            st.info(
                f"**Max safe BW velocity: {bw_col['max_safe_bw_m_h']:.1f} m/h** "
                f"(maintains ≥ {freeboard_mm:.0f} mm freeboard below collector).  "
                f"Proposed BW: **{bw_col['proposed_bw_m_h']:.1f} m/h**."
            )

            # Expansion — two scenarios
            # A: water-only BW (pump rate) — governs collector check
            # B: air+water combined (air scour rate as equivalent) — governs cleaning
            exp_rows = []
            _bw_integrity_alerts = []
            for L in bw_col["per_layer"]:
                if L.get("elutriation_risk"):
                    status = "Approaching terminal velocity"
                    _bw_integrity_alerts.append((
                        "critical",
                        f"{L['media_type']}: backwash velocity approaches terminal settling velocity",
                        "Risk of progressive media loss over repeated backwash cycles. "
                        "Reduce backwash rate or raise the outlet collector to maintain media inventory."
                    ))
                elif L["fluidised"]:
                    status = f"Fluidised — {L['expansion_pct']}% bed expansion"
                else:
                    status = f"Below fluidisation threshold (u_mf = {L['u_mf_m_h']} m/h)"
                    _bw_integrity_alerts.append((
                        "warning" if freeboard_mm >= 150 else "advisory",
                        f"{L['media_type']}: hydraulic bed lift not achieved at current water rate",
                        "Backwash velocity is below the minimum fluidisation threshold for this media fraction. "
                        "Air scour provides primary mechanical cleaning action. "
                        "Hydraulic lift of settled solids is not achieved by water flow alone at this rate."
                    ))
                exp_rows.append({
                    "Media":          L["media_type"],
                    "d10 (mm)":       L["d10_mm"],
                    "d50 (mm)":       L.get("d50_mm", "—"),
                    "Ar":             L.get("Ar", "—"),
                    "Re_mf":          L.get("Re_mf", "—"),
                    "u_mf (m/h)":     L["u_mf_m_h"],
                    "n_rz":           L.get("n_rz", "—"),
                    "u_t (m/h)":      L["u_t_m_h"],
                    "ε₀":             L["epsilon0"],
                    "ε_f":            L["eps_f"],
                    "Settled (m)":    L["depth_settled_m"],
                    "Expanded (m)":   L["depth_expanded_m"],
                    "Expansion (%)":  L["expansion_pct"],
                    "Status":         status,
                })
            st.dataframe(pd.DataFrame(exp_rows),
                         use_container_width=True, hide_index=True)

            # Air+water combined scenario
            from engine.backwash import bed_expansion as _bed_exp
            exp_combined = _bed_exp(
                layers=layers,
                bw_velocity_m_h=air_scour_rate,   # use air scour rate as combined equiv.
                water_temp_c=bw_temp,
                rho_water=rho_bw,
            )
            st.markdown(f"**Air + water combined phase** "
                        f"(equivalent velocity = air scour rate = {air_scour_rate:.0f} m/h):")
            comb_rows = []
            for L in exp_combined["layers"]:
                comb_rows.append({
                    "Media":          L["media_type"],
                    "u_mf (m/h)":     L["u_mf_m_h"],
                    "Fluidised":      "Yes ✅" if L["fluidised"] else "No",
                    "ε_f":            L["eps_f"],
                    "Settled (m)":    L["depth_settled_m"],
                    "Expanded (m)":   L["depth_expanded_m"],
                    "Expansion (%)":  L["expansion_pct"],
                    "Note":           L["warning"] if L["warning"] else "OK",
                })
            st.dataframe(pd.DataFrame(comb_rows),
                         use_container_width=True, hide_index=True)

            if _bw_integrity_alerts:
                with st.expander(f"🔴 Backwash Integrity — {len(_bw_integrity_alerts)} concern(s) identified"):
                    for _lvl, _ttl, _msg in _bw_integrity_alerts:
                        show_alert(_lvl, _ttl, _msg)

            ec1, ec2, ec3 = st.columns(3)
            ec1.metric("Settled bed",   f"{exp_combined['total_settled_m']:.3f} m")
            ec2.metric("Expanded bed",  f"{exp_combined['total_expanded_m']:.3f} m")
            ec3.metric("Net expansion", f"{exp_combined['total_expansion_pct']:.1f} %")

            # Dynamic note: built from actual per-layer u_mf results
            _wbw_lines = []
            for _L in bw_col["per_layer"]:
                _umf  = _L["u_mf_m_h"]
                _name = _L["media_type"]
                if _L["fluidised"]:
                    _wbw_lines.append(
                        f"**{_name}** (d10={_L['d10_mm']} mm, u_mf={_umf} m/h): "
                        f"✅ fluidised at {bw_velocity:.0f} m/h — "
                        f"{_L['expansion_pct']:.0f} % bed expansion."
                    )
                else:
                    _wbw_lines.append(
                        f"**{_name}** (d10={_L['d10_mm']} mm, u_mf={_umf} m/h): "
                        f"⚪ NOT fluidised at {bw_velocity:.0f} m/h "
                        f"(need >{_umf:.0f} m/h water-only) — "
                        f"bed rests on nozzle plate during water-only phase."
                    )
            _wbw_note = "  \n".join(_wbw_lines)

            _any_not_fluidised = any(not _L["fluidised"] for _L in bw_col["per_layer"])
            _all_fluidised     = all(_L["fluidised"]     for _L in bw_col["per_layer"])

            if _all_fluidised:
                _cleaning_note = (
                    f"All layers are hydraulically fluidised at the water-only BW rate "
                    f"({bw_velocity:.0f} m/h). Air scour ({air_scour_rate:.0f} m/h) "
                    "provides additional mechanical agitation to release compacted cake."
                )
            elif _any_not_fluidised:
                _cleaning_note = (
                    f"One or more layers rest on the nozzle plate at the water-only BW rate "
                    f"({bw_velocity:.0f} m/h). In horizontal MMF, gravity cannot transport "
                    "dislodged solids upward — **air scour is essential** for agitation and "
                    "gas-lift transport of solids to the outlet regardless of fluidisation state."
                )

            st.caption(
                f"**Water-only BW at {bw_velocity:.0f} m/h** (per-layer status):  \n"
                + _wbw_note + "  \n  \n"
                + _cleaning_note + "  \n"
                "The combined phase table above shows expansion when air+water act simultaneously."
            )

            if bw_col["media_loss_risk"]:
                show_alert(
                    "critical",
                    "Collector height insufficient — media carryover risk",
                    f"The expanded bed top ({bw_col['expanded_top_m']:.3f} m) reaches or exceeds "
                    f"the BW outlet collector ({bw_col['collector_h_m']:.3f} m). "
                    f"Media loss will occur progressively over repeated backwash cycles. "
                    f"Maximum safe backwash velocity with current geometry: "
                    f"{bw_col['max_safe_bw_m_h']:.1f} m/h. "
                    "Raise the collector or reduce backwash rate to restore freeboard margin."
                )

        with st.expander("2 · BW pump & air blower capacity", expanded=True):
            bh1, bh2, bh3, bh4 = st.columns(4)
            bh1.metric("BW flow",        f"{bw_hyd['q_bw_m3h']:,.0f} m³/h",
                       help=bw_hyd["bw_governs"])
            bh2.metric("BW LV actual",   f"{bw_hyd['bw_lv_actual_m_h']:.1f} m/h")
            bh3.metric("Air scour flow", f"{bw_hyd['q_air_m3h']:,.0f} m³/h")
            bh4.metric("Blower est.",    f"{bw_hyd['p_blower_est_kw']:.1f} kW")
            st.table(pd.DataFrame([
                ["Governing BW flow",
                 f"{bw_hyd['q_bw_m3h']:,.1f} m³/h ({bw_hyd['bw_governs']})"],
                ["BW design capacity (×1.10)",
                 f"{bw_hyd['q_bw_design_m3h']:,.1f} m³/h"],
                ["Air design capacity (×1.10)",
                 f"{bw_hyd['q_air_design_m3h']:,.1f} m³/h"],
                ["Blower power (est., η=0.65)",
                 f"{bw_hyd['p_blower_est_kw']:.1f} kW"],
                ["BW water: ρ",
                 f"{rho_bw:.2f} kg/m³  |  μ={mu_bw*1000:.4f} cP"],
            ], columns=["Parameter", "Value"]))

        with st.expander("3 · BW sequence & waste volumes", expanded=True):
            st.dataframe(pd.DataFrame(bw_seq["steps"]),
                         use_container_width=True, hide_index=True)
            st.divider()
            w1, w2, w3, w4 = st.columns(4)
            w1.metric("BW duration (avg)",  f"{bw_seq['dur_total_avg_min']} min")
            w2.metric("Total vol / filter", f"{bw_seq['total_vol_avg_m3']:.0f} m³")
            w3.metric("Waste / filter",     f"{bw_seq['waste_vol_avg_m3']:.0f} m³")
            w4.metric("Plant waste / day",  f"{bw_seq['waste_vol_daily_m3']:.0f} m³/d")

            st.markdown("**Waste volume & TSS mass balance**")
            st.table(pd.DataFrame([
                ["Low TSS",
                 f"{tss_low:.0f} mg/L",
                 f"{bw_seq['total_vol_low_m3']:.0f} m³",
                 f"{m_sol_low:.0f} kg",
                 f"{w_tss_low:.0f} mg/L",
                 f"{m_daily_low:,.0f} kg/d"],
                ["Avg TSS",
                 f"{tss_avg:.0f} mg/L",
                 f"{bw_seq['total_vol_avg_m3']:.0f} m³",
                 f"{m_sol_avg:.0f} kg",
                 f"{w_tss_avg:.0f} mg/L",
                 f"{m_daily_avg:,.0f} kg/d"],
                ["High TSS",
                 f"{tss_high:.0f} mg/L",
                 f"{bw_seq['total_vol_high_m3']:.0f} m³",
                 f"{m_sol_high:.0f} kg",
                 f"{w_tss_high:.0f} mg/L",
                 f"{m_daily_high:,.0f} kg/d"],
            ], columns=["Scenario", "Feed TSS",
                        "BW vol / filter", "Solids captured / filter",
                        "Waste TSS conc.", "Plant solids / day"]))
            st.caption(
                f"Run time between BW cycles: {run_time_h:.1f} h  "
                f"({bw_cycles_day} cycle/d per filter).  "
                "Solids captured = TSS × Q_filter × run_time.  "
                "Waste TSS = solids mass / waste volume (excl. rinse)."
            )

        with st.expander("4 · Filtration cycle matrix — TSS × temperature", expanded=True):
            if filt_cycles and cycle_matrix:
                first_cyc   = next(iter(filt_cycles.values()))
                _alpha_src  = first_cyc["alpha_source"]
                _alpha_used = first_cyc["alpha_used_m_kg"]
                _alpha_cal  = first_cyc["alpha_calibrated_m_kg"]

                st.info(
                    f"**Ruth cake model** · BW setpoint {dp_trigger_bar:.2f} bar · "
                    f"M_max {solid_loading:.2f} kg/m² · "
                    f"α ({_alpha_src}) = {_alpha_used/1e9:.1f} × 10⁹ m/kg "
                    f"(auto-cal at design temp = {_alpha_cal/1e9:.1f} × 10⁹ m/kg) · "
                    f"Temperature range {temp_low:.0f} – {feed_temp:.0f} – {temp_high:.0f} °C"
                )

                # ── Matrix tables: one per redundancy scenario ────────────
                for sc_lbl, sc_temps in cycle_matrix.items():
                    _lv = filt_cycles[sc_lbl]["lv_m_h"]
                    st.markdown(
                        f"**Scenario {sc_lbl} · LV = {_lv:.1f} m/h**  "
                        f"— Cycle duration (h) to reach {dp_trigger_bar:.2f} bar ΔP"
                    )
                    # Build matrix: rows = TSS label, cols = temp label
                    mat_rows = []
                    for tss_lbl, tss_v in zip(_tss_labels, _tss_vals):
                        row = {"Feed TSS": tss_lbl}
                        for t_lbl in _temp_labels:
                            cyc_t = sc_temps[t_lbl]
                            # find the matching TSS row in tss_results
                            tr = next(
                                (r for r in cyc_t["tss_results"]
                                 if r["TSS (mg/L)"] == tss_v), None)
                            row[t_lbl] = f"{tr['Cycle duration (h)']:.1f} h" if tr else "—"
                        mat_rows.append(row)
                    mat_df = pd.DataFrame(mat_rows).set_index("Feed TSS")
                    st.dataframe(mat_df, use_container_width=True)

                    # Sub-caption: ΔP clean at each temperature for this scenario
                    dp_clean_str = "  ·  ".join(
                        f"{t_lbl}: ΔP_clean = {sc_temps[t_lbl]['dp_clean_bar']:.4f} bar"
                        for t_lbl in _temp_labels
                    )
                    st.caption(
                        f"M* (load at trigger): "
                        + "  ·  ".join(
                            f"{t_lbl}: {sc_temps[t_lbl]['loading_at_trigger_kg_m2']:.2f} kg/m²"
                            for t_lbl in _temp_labels
                        )
                        + f"  |  {dp_clean_str}"
                    )

                # ── ΔP vs solid load curve (N, design temp) ───────────────
                with st.expander("ΔP vs M curve — N scenario, design temperature", expanded=False):
                    st.dataframe(pd.DataFrame(first_cyc["dp_curve"]),
                                 use_container_width=True, hide_index=True)

                st.caption(
                    "ΔP_total = ΔP_clean (Ergun) + α × μ(T) × LV × M.  "
                    "α fixed at design-temperature calibration; temperature effect enters "
                    "through μ(T) — colder water → higher viscosity → shorter cycle.  "
                    "α = 0 → auto-calibrated so dirty ΔP = trigger at M_max and design temp."
                )

                with st.expander("Reference: typical α by TSS type", expanded=False):
                    st.dataframe(pd.DataFrame([
                        {"TSS type": "Coarse mineral / silt",        "α (× 10⁹ m/kg)": "0.1 – 10"},
                        {"TSS type": "Seawater mixed TSS (typical)", "α (× 10⁹ m/kg)": "10 – 50"},
                        {"TSS type": "Organic-rich / algae-laden",   "α (× 10⁹ m/kg)": "100 – 500"},
                        {"TSS type": "Clay / fine colloids",          "α (× 10⁹ m/kg)": "1 000 – 10 000"},
                    ]), use_container_width=True, hide_index=True)
            else:
                st.info("No filtration cycle data available.")

        with st.expander("5 · BW scheduling & system feasibility", expanded=True):
            _bw_steps = [
                ("① Gravity drain",        bw_s_drain),
                ("② Air scour only",       bw_s_air),
                ("③ Air + low-rate water", bw_s_airw),
                ("④ High-rate water",      bw_s_hw),
                ("⑤ Settling",             bw_s_settle),
                ("⑥ Fill & rinse",         bw_s_fill),
            ]
            # BW step summary
            cum = 0
            step_rows = []
            for nm, dur in _bw_steps:
                cum += dur
                step_rows.append({"Step": nm, "Duration (min)": dur,
                                   "Cumulative (min)": cum})
            step_rows.append({"Step": "TOTAL", "Duration (min)": bw_total_min,
                               "Cumulative (min)": bw_total_min})
            cA, cB = st.columns([1, 2])
            with cA:
                st.markdown("**BW step breakdown**")
                st.dataframe(pd.DataFrame(step_rows),
                             use_container_width=True, hide_index=True)

            with cB:
                st.markdown(
                    f"**BW duration: {bw_total_min} min ({_bw_dur_h*60:.0f} min)**  \n"
                    f"Filter cycle + BW = cycle time + {bw_total_min} min.  \n"
                    f"Availability = cycle time / (cycle + BW time) × 100 %."
                )

            if feasibility_matrix:
                for sc_lbl, sc_temps in feasibility_matrix.items():
                    _lv  = filt_cycles[sc_lbl]["lv_m_h"]
                    _nact_f = next(
                        n for x, n, _ in _load_data_cyc
                        if ("N" if x == 0 else f"N-{x}") == sc_lbl)
                    _nact_total = _nact_f * streams
                    st.markdown(
                        f"---\n**Scenario {sc_lbl} · "
                        f"{_nact_f} active / stream × {streams} stream(s) "
                        f"= {_nact_total} active filters plant-wide · "
                        f"LV = {_lv:.1f} m/h**"
                    )

                    # ── Availability % matrix ─────────────────────────────
                    st.markdown("*Availability (%) = filtration time / total period*")
                    avail_rows = []
                    for tss_lbl in _tss_labels:
                        row = {"Feed TSS": tss_lbl}
                        for t_lbl in _temp_labels:
                            kpi = sc_temps[t_lbl][tss_lbl]
                            row[t_lbl] = f"{kpi['avail_pct']:.1f} %"
                        avail_rows.append(row)
                    st.dataframe(
                        pd.DataFrame(avail_rows).set_index("Feed TSS"),
                        use_container_width=True)

                    # ── BW per day matrix ─────────────────────────────────
                    st.markdown("*BW cycles per filter per day*")
                    bwday_rows = []
                    for tss_lbl in _tss_labels:
                        row = {"Feed TSS": tss_lbl}
                        for t_lbl in _temp_labels:
                            kpi = sc_temps[t_lbl][tss_lbl]
                            row[t_lbl] = f"{kpi['bw_per_day']:.1f}"
                        bwday_rows.append(row)
                    st.dataframe(
                        pd.DataFrame(bwday_rows).set_index("Feed TSS"),
                        use_container_width=True)

                    # ── Simultaneity matrix ───────────────────────────────
                    st.markdown(
                        "*BW systems required (plant-wide) — "
                        "simultaneous demand = active filters (plant-wide) × BW dur / total period*"
                    )
                    sim_rows = []
                    for tss_lbl in _tss_labels:
                        row = {"Feed TSS": tss_lbl}
                        for t_lbl in _temp_labels:
                            kpi = sc_temps[t_lbl][tss_lbl]
                            row[t_lbl] = (
                                f"{kpi['sim_demand']:.2f} "
                                f"→ {kpi['bw_trains']} BW system(s)"
                            )
                        sim_rows.append(row)
                    st.dataframe(
                        pd.DataFrame(sim_rows).set_index("Feed TSS"),
                        use_container_width=True)

                    # ── Feasibility scorecard ─────────────────────────────
                    st.markdown("*Feasibility score*")
                    score_rows = []
                    for tss_lbl in _tss_labels:
                        row = {"Feed TSS": tss_lbl}
                        for t_lbl in _temp_labels:
                            kpi = sc_temps[t_lbl][tss_lbl]
                            row[t_lbl] = kpi["score"]
                        score_rows.append(row)
                    st.dataframe(
                        pd.DataFrame(score_rows).set_index("Feed TSS"),
                        use_container_width=True)

                # ── Design guidance ───────────────────────────────────────
                st.markdown("---")
                st.markdown("**Design guidance**")

                # Find worst-case cell (design scenario N, design temp, high TSS)
                _n_kpi = feasibility_matrix.get("N", {}).get(
                    f"Design ({feed_temp:.0f}°C)", {}).get(
                    f"High ({tss_high:.0f} mg/L)", {})
                if _n_kpi:
                    _flag = _n_kpi["flag"]
                    _trains = _n_kpi["bw_trains"]
                    _avail  = _n_kpi["avail_pct"]
                    _bwday  = _n_kpi["bw_per_day"]
                    _tcyc   = _n_kpi["t_cycle_h"]

                    guidance = []
                    if _flag == "OK":
                        guidance.append(
                            f"🟢 Design is feasible at high TSS / design temp: "
                            f"availability {_avail:.1f} %, "
                            f"{_bwday:.1f} BW cycles/filter/day, "
                            f"1 BW system (plant-wide) is sufficient."
                        )
                    else:
                        if _trains > 1:
                            guidance.append(
                                f"🔴 **{_trains} simultaneous BW systems required plant-wide** "
                                f"at high TSS ({tss_high:.0f} mg/L) / design temp — "
                                f"consider: (a) add a dedicated BW pump/blower per stream, "
                                f"(b) staggered BW scheduling with time-lock logic, or "
                                f"(c) increase filter area to extend cycle > 6 h."
                            )
                        if _tcyc < 4:
                            guidance.append(
                                f"⚠️ Cycle time {_tcyc:.1f} h at worst case — "
                                f"BW occupies {bw_total_min} min every "
                                f"{_tcyc*60:.0f} min: consider increasing "
                                f"filter area or adding redundancy to extend cycle."
                            )
                        if _avail < 80:
                            guidance.append(
                                f"🔴 Availability {_avail:.1f} % — "
                                f"the system spends > 20 % of time in BW; "
                                f"net production capacity is significantly reduced."
                            )

                    for g in guidance:
                        st.markdown(g)

                st.caption(
                    "Scoring: 🟢 Good = avail ≥ 90 %, BW systems ≤ 1, cycle ≥ 6 h  |  "
                    "🟡 Caution = avail 80–90 %, BW systems ≤ 2, cycle ≥ 3 h  |  "
                    "🔴 Critical = avail < 80 %, BW systems > 2, or cycle < 3 h.  "
                    "BW systems are plant-wide (baseline: 1–2 for the whole plant).  "
                    "Simultaneity = n_active_plant × BW_dur / (cycle + BW_dur);  "
                    "BW systems needed = ⌈simultaneity⌉.  "
                    "Redundancy (N-1, N-2 …) = standby filters per stream."
                )

        with st.expander("6 · BW system equipment data sheet", expanded=True):
            st.markdown("### BW pump")
            p1, p2, p3, p4 = st.columns(4)
            p1.metric("Design flow",  f"{bw_sizing['q_bw_design_m3h']:,.0f} m³/h")
            p2.metric("Total head",   f"{bw_sizing['bw_head_mwc']:.1f} mWC  "
                                       f"({bw_sizing['bw_head_bar']:.2f} bar)")
            p3.metric("Shaft power",  f"{bw_sizing['p_pump_shaft_kw']:.0f} kW")
            p4.metric("Motor power",  f"{bw_sizing['p_pump_motor_kw']:.0f} kW")
            st.table(pd.DataFrame([
                ["Design flow (duty)",        f"{bw_sizing['q_bw_design_m3h']:,.1f} m³/h"],
                ["Total dynamic head",        f"{bw_sizing['bw_head_mwc']:.1f} mWC  ({bw_sizing['bw_head_bar']:.3f} bar)"],
                ["Pump hydraulic efficiency", f"{bw_sizing['bw_pump_eta']*100:.0f} %"],
                ["Shaft power",               f"{bw_sizing['p_pump_shaft_kw']:.1f} kW"],
                ["Motor power (absorbed)",    f"{bw_sizing['p_pump_motor_kw']:.1f} kW"],
                ["Duty / standby",            f"{_n_bw_systems}D / 1S  (plant-wide)"],
                ["Fluid",                     f"BW water  ρ={rho_bw:.1f} kg/m³"],
            ], columns=["Parameter", "Value"]))

            st.markdown("### Air blower")
            b1, b2, b3, b4 = st.columns(4)
            b1.metric("Design flow",     f"{bw_sizing['q_air_design_m3h']:,.0f} m³/h  "
                                          f"({bw_sizing['q_air_design_m3min']:.1f} m³/min)")
            b2.metric("ΔP (total)",      f"{bw_sizing['dp_total_bar']:.3f} bar")
            b3.metric("Shaft power",     f"{bw_sizing['p_blower_shaft_kw']:.0f} kW")
            b4.metric("Motor power",     f"{bw_sizing['p_blower_motor_kw']:.0f} kW")
            st.table(pd.DataFrame([
                ["Inlet volume flow",          f"{bw_sizing['q_air_design_m3h']:,.1f} m³/h  ({bw_sizing['q_air_design_m3min']:.1f} m³/min)"],
                ["Inlet conditions",           f"T={blower_inlet_temp_c:.0f} °C  ρ={bw_sizing['rho_air_kg_m3']:.4f} kg/m³  P={bw_sizing['P1_pa']/1000:.1f} kPa (a)"],
                ["Vessel back-pressure",       f"{vessel_pressure_bar:.2f} bar g"],
                ["Water submergence (≈ ID/2)", f"{bw_sizing['h_submergence_m']:.2f} m  →  {bw_sizing['dp_sub_bar']:.3f} bar"],
                ["Discharge pressure (abs)",   f"{bw_sizing['P2_pa']/1000:.1f} kPa  (ratio {bw_sizing['pressure_ratio']:.3f})"],
                ["Total ΔP",                   f"{bw_sizing['dp_total_bar']:.3f} bar"],
                ["Isentropic efficiency",      f"{bw_sizing['blower_eta']*100:.0f} %"],
                ["Ideal (adiabatic) power",    f"{bw_sizing['p_blower_ideal_kw']:.1f} kW"],
                ["Shaft power",               f"{bw_sizing['p_blower_shaft_kw']:.1f} kW"],
                ["Motor power (absorbed)",    f"{bw_sizing['p_blower_motor_kw']:.1f} kW"],
                ["Basis",                      "Adiabatic compression  γ = 1.4 (dry air)"],
            ], columns=["Parameter", "Value"]))

            st.markdown("### BW water storage tank")
            t1, t2, t3 = st.columns(3)
            t1.metric("Vol/cycle/system",  f"{bw_sizing['bw_vol_per_cycle_m3']:.0f} m³")
            t2.metric("Simultaneous syst.", f"{bw_sizing['n_bw_systems']}")
            t3.metric("Recommended tank",  f"{bw_sizing['v_tank_m3']:.0f} m³",
                      help=f"Governs: {bw_sizing['tank_governs']}")
            st.table(pd.DataFrame([
                ["BW vol / filter / cycle (avg)",  f"{bw_sizing['bw_vol_per_cycle_m3']:.1f} m³"],
                ["Simultaneous BW systems",         f"{bw_sizing['n_bw_systems']}  (design-point N scenario, avg TSS)"],
                ["Safety factor",                   f"{bw_sizing['tank_sf']:.1f}×"],
                ["Volume — cycle-based",            f"{bw_sizing['v_cycle_m3']:.0f} m³  (vol × systems × SF)"],
                ["Volume — 10-min rule",            f"{bw_sizing['v_10min_m3']:.0f} m³  (Q_bw_design / 6)"],
                ["Recommended tank volume",         f"{bw_sizing['v_tank_m3']:.0f} m³  (governs: {bw_sizing['tank_governs']})"],
            ], columns=["Parameter", "Value"]))
            st.caption(
                "BW pump head is a user input (sidebar → Energy block).  "
                "Breakdown: vessel operating head + bed/nozzle ΔP + pipe losses.  "
                "Blower submergence = filter ID / 2 (centroid of water column above nozzle plate).  "
                "Tank volume covers simultaneous BW systems at design point; "
                "operator may add live-well or equalisation margin on top."
            )

    # ─────────────────────────────────────────────────────────────────────
    # TAB 7 · WEIGHT
    # ─────────────────────────────────────────────────────────────────────
    with tab_weight:
        st.subheader("Empty weight — consolidated")

        with st.expander("1 · Vessel body (shell + 2 heads)", expanded=True):
            wa, wb = st.columns(2)
            with wa:
                st.markdown("**Cylindrical shell**")
                st.table(pd.DataFrame([
                    ["Mean diameter",    f"{wt_body['d_mean_shell_m']:.4f} m"],
                    ["Wall thickness",   f"{mech['t_shell_design_mm']} mm"],
                    ["Surface area",     f"{wt_body['area_shell_m2']:.3f} m²"],
                    ["Metal volume",     f"{wt_body['vol_shell_m3']:.4f} m³"],
                    ["Shell weight",     f"{wt_body['weight_shell_kg']:,.1f} kg"],
                ], columns=["Item", "Value"]))
            with wb:
                st.markdown(f"**Dish ends × 2  ({end_geometry})**")
                st.table(pd.DataFrame([
                    ["Mean diameter",       f"{wt_body['d_mean_head_m']:.4f} m"],
                    ["Wall thickness",      f"{mech['t_head_design_mm']} mm"],
                    ["Surface area (one)", f"{wt_body['area_one_head_m2']:.3f} m²"],
                    ["Both heads weight",  f"{wt_body['weight_two_heads_kg']:,.1f} kg"],
                ], columns=["Item", "Value"]))

        with st.expander("2 · Nozzles (stubs + flanges)", expanded=True):
            df_nozzle = pd.DataFrame(nozzle_sched)
            edited = st.data_editor(
                df_nozzle, use_container_width=True,
                hide_index=True, num_rows="dynamic",
                column_config={
                    "Service":        st.column_config.TextColumn(width=160),
                    "DN (mm)":        st.column_config.SelectboxColumn(
                                          options=DN_SERIES, width=85),
                    "Schedule":       st.column_config.SelectboxColumn(
                                          options=SCHEDULES, width=90),
                    "Rating":         st.column_config.SelectboxColumn(
                                          options=FLANGE_RATINGS, width=80),
                    "Qty":            st.column_config.NumberColumn(width=55),
                    "Wt/nozzle (kg)": st.column_config.NumberColumn(format="%.1f"),
                    "Total wt (kg)":  st.column_config.NumberColumn(format="%.1f"),
                })
            nozzle_wt_edited = (edited["Total wt (kg)"].sum()
                                if "Total wt (kg)" in edited.columns else w_noz)

        with st.expander("3 · Nozzle plate assembly", expanded=True):
            n1, n2, n3 = st.columns(3)
            n1.metric("Plate",        f"{wt_np['weight_plate_kg']:,.0f} kg")
            n2.metric("Support beams",f"{wt_np['weight_beams_kg']:,.0f} kg")
            n3.metric("Total assy.",  f"{wt_np['weight_total_kg']:,.0f} kg")

        with st.expander("4 · Supports", expanded=True):
            q1, q2, q3 = st.columns(3)
            q1.metric("Type",          wt_sup["support_type"])
            q2.metric("Qty",           wt_sup["n_supports"])
            q3.metric("Total weight",  f"{wt_sup['weight_all_supports_kg']:,.0f} kg")

        with st.expander("5 · Vessel internals", expanded=True):
            ii1, ii2, ii3 = st.columns(3)
            ii1.metric("Strainer nozzles",
                       f"{wt_int['weight_strainers_kg']:,.0f} kg",
                       delta=f"{wt_int['n_strainer_nozzles']} × "
                             f"{wt_int['strainer_material']} "
                             f"@ {wt_int['weight_per_strainer_kg']:.3f} kg",
                       delta_color="off")
            ii2.metric("Air scour header",
                       f"{wt_int['weight_air_header_kg']:,.0f} kg",
                       delta=f"DN{wt_int['air_header_dn_mm']} × "
                             f"{wt_int['air_header_length_m']:.1f} m @ "
                             f"{wt_int['air_header_kg_per_m']} kg/m",
                       delta_color="off")
            ii3.metric("Manholes",
                       f"{wt_int['weight_manholes_kg']:,.0f} kg",
                       delta=f"{wt_int['n_manholes']} × {wt_int['manhole_dn']}",
                       delta_color="off")
            st.caption(
                f"Total internals: **{wt_int['weight_internals_kg']:,.0f} kg = "
                f"{wt_int['weight_internals_t']:.3f} t**"
            )

        with st.expander("6 · Consolidated summary", expanded=True):
            w_total_final = (wt_body["weight_body_kg"] + nozzle_wt_edited
                             + wt_np["weight_total_kg"]
                             + wt_sup["weight_all_supports_kg"])
            w_total_final = (wt_body["weight_body_kg"] + nozzle_wt_edited
                             + wt_np["weight_total_kg"]
                             + wt_sup["weight_all_supports_kg"]
                             + wt_int["weight_internals_kg"])
            st.table(pd.DataFrame([
                ["Shell (cylindrical)",
                 f"{wt_body['weight_shell_kg']:>12,.1f} kg"],
                ["2 × Dish ends",
                 f"{wt_body['weight_two_heads_kg']:>12,.1f} kg"],
                ["Nozzles (stubs + flanges)",
                 f"{nozzle_wt_edited:>12,.1f} kg"],
                ["Nozzle plate + IPE beams",
                 f"{wt_np['weight_total_kg']:>12,.1f} kg"],
                [f"Supports ({wt_sup['support_type']})",
                 f"{wt_sup['weight_all_supports_kg']:>12,.1f} kg"],
                ["Strainer nozzles",
                 f"{wt_int['weight_strainers_kg']:>12,.1f} kg"],
                ["Air scour header",
                 f"{wt_int['weight_air_header_kg']:>12,.1f} kg"],
                ["Manholes",
                 f"{wt_int['weight_manholes_kg']:>12,.1f} kg"],
                ["─" * 30, "─" * 16],
                ["TOTAL EMPTY WEIGHT",
                 f"{w_total_final:>12,.1f} kg"],
                ["",
                 f"= {w_total_final/1000:.3f} t"],
            ], columns=["Component", "Weight"]))

            e1, e2, e3, e4, e5, e6 = st.columns(6)
            e1.metric("Shell+heads", f"{wt_body['weight_body_kg']:,.0f} kg")
            e2.metric("Nozzles",     f"{nozzle_wt_edited:,.0f} kg")
            e3.metric("Plate+beams", f"{wt_np['weight_total_kg']:,.0f} kg")
            e4.metric("Supports",    f"{wt_sup['weight_all_supports_kg']:,.0f} kg")
            e5.metric("Internals",   f"{wt_int['weight_internals_kg']:,.0f} kg")
            e6.metric("TOTAL",       f"{w_total_final/1000:.3f} t",
                      delta=f"{w_total_final:,.0f} kg", delta_color="off")
            st.caption(
                "⚠️  Underdrains, platform/walkway, piping manifolds, "
                "and instrument connections not included."
            )

        with st.expander("7 · Operating (working) weight & support loads", expanded=True):
            st.markdown(
                "Working weight = empty vessel + dry media + process water "
                "(vessel running full at design conditions). "
                "This is the load the saddles or legs must carry."
            )

            o1, o2, o3, o4, o5 = st.columns(5)
            o1.metric("Empty vessel",    f"{wt_oper['w_empty_kg']:,.0f} kg")
            o2.metric("Lining / coating", f"{wt_oper['w_lining_kg']:,.0f} kg")
            o3.metric("Media (dry)",     f"{wt_oper['w_media_kg']:,.0f} kg")
            o4.metric("Process water",   f"{wt_oper['w_water_kg']:,.0f} kg")
            o5.metric("Operating total", f"{wt_oper['w_operating_t']:.3f} t",
                      delta=f"{wt_oper['w_operating_kg']:,.0f} kg",
                      delta_color="off")

            st.markdown("**Weight breakdown**")
            lining_label = (f"Internal {lining_result['protection_type'].lower()}"
                            if lining_result['protection_type'] != "None"
                            else "Internal lining / coating")
            st.table(pd.DataFrame([
                ["Empty vessel (steel structure)",
                 f"{wt_oper['w_empty_kg']:>12,.1f} kg"],
                [lining_label,
                 f"{wt_oper['w_lining_kg']:>12,.1f} kg"],
                ["Media — dry solid mass",
                 f"{wt_oper['w_media_kg']:>12,.1f} kg"],
                ["Process water (vessel full)",
                 f"{wt_oper['w_water_kg']:>12,.1f} kg"],
                ["─" * 30, "─" * 16],
                ["OPERATING WEIGHT",
                 f"{wt_oper['w_operating_kg']:>12,.1f} kg  =  {wt_oper['w_operating_t']:.3f} t"],
            ], columns=["Component", "Weight"]))

            st.markdown("**Internal volume breakdown**")
            st.table(pd.DataFrame([
                ["Cylindrical shell (internal)",
                 f"{wt_oper['v_cylinder_m3']:.3f} m³"],
                ["Two dish ends (internal)",
                 f"{wt_oper['v_heads_m3']:.3f} m³"],
                ["Total internal volume",
                 f"{wt_oper['v_total_internal_m3']:.3f} m³"],
                ["Media solid volume  Σ(depth × area × (1−ε₀))",
                 f"{wt_oper['v_solid_media_m3']:.4f} m³"],
                ["Water volume  (total − solid media)",
                 f"{wt_oper['v_water_m3']:.3f} m³"],
            ], columns=["Item", "Value"]))

            st.markdown("**Media layer detail**")
            st.dataframe(pd.DataFrame(wt_oper["media_rows"]),
                         use_container_width=True, hide_index=True)

            st.markdown("**Support / saddle loads**")
            s1, s2, s3 = st.columns(3)
            s1.metric("Support type",     wt_sup["support_type"])
            s2.metric("Number of supports", str(wt_oper["n_supports"]))
            s3.metric("Load per support",
                      f"{wt_oper['load_per_support_t']:.3f} t",
                      delta=f"{wt_oper['load_per_support_kN']:.1f} kN",
                      delta_color="off")
            st.table(pd.DataFrame([
                ["Number of supports",
                 str(wt_oper["n_supports"])],
                ["Operating weight",
                 f"{wt_oper['w_operating_kg']:,.1f} kg  =  {wt_oper['w_operating_t']:.3f} t"],
                ["Vertical load per support",
                 f"{wt_oper['load_per_support_kg']:,.1f} kg  =  "
                 f"{wt_oper['load_per_support_t']:.3f} t  =  "
                 f"{wt_oper['load_per_support_kN']:.1f} kN"],
            ], columns=["Parameter", "Value"]))
            st.caption(
                "Load per support assumes uniform distribution (equal reaction at each support). "
                "For saddle design, this is the vertical reaction used to check web shear, "
                "bending stress at the horn, and base plate bearing pressure (Zick analysis). "
                "Wind/seismic and friction loads are not included here — add as project-specific "
                "factors per local code (ASCE 7, EN 1991, etc.)."
            )

        with st.expander("8 · Saddle positioning & section selection", expanded=True):

            # ── Positioning ───────────────────────────────────────────────
            st.markdown("### Positioning (Zick method)")
            _aov_color = "🟢" if wt_saddle["a_over_R_ok"] else "🔴"
            sp1, sp2, sp3, sp4 = st.columns(4)
            sp1.metric("L / D ratio",          f"{wt_saddle['ld_ratio']:.2f}")
            sp2.metric("Spacing factor α",     f"{wt_saddle['alpha_pct']} %")
            sp3.metric("Saddle 1 from head",   f"{wt_saddle['saddle_1_from_left_m']:.3f} m")
            sp4.metric("Saddle 2 from head",   f"{wt_saddle['saddle_2_from_left_m']:.3f} m",
                       help="Measured from right head tangent = same distance as saddle 1")

            st.table(pd.DataFrame([
                ["Total vessel length (T/T)",      f"{total_length:.3f} m"],
                ["Vessel OD",                       f"{mech['od_m']:.4f} m"],
                ["L / D ratio",                     f"{wt_saddle['ld_ratio']:.2f}"],
                ["Spacing factor α",               f"{wt_saddle['alpha_pct']} %  "
                                                    f"({'L/D<3 →25%' if wt_saddle['ld_ratio']<3 else 'L/D<5 →22%' if wt_saddle['ld_ratio']<5 else 'L/D≥5 →20%'})"],
                ["Saddle 1 — from left head",      f"{wt_saddle['saddle_1_from_left_m']:.3f} m"],
                ["Saddle 2 — from left head",      f"{wt_saddle['saddle_2_from_left_m']:.3f} m"],
                ["Span between saddles",           f"{wt_saddle['saddle_gap_m']:.3f} m"],
                ["a / R  (Zick parameter)",        f"{wt_saddle['a_over_R']:.3f}  {_aov_color} {'OK — ≤ 0.5' if wt_saddle['a_over_R_ok'] else 'REVIEW — > 0.5'}"],
                ["Contact arc length (120°)",      f"{wt_saddle['arc_m']:.3f} m"],
                ["Simplified saddle moment",       f"{wt_saddle['m_saddle_kNm']:.0f} kN·m"],
            ], columns=["Parameter", "Value"]))
            st.caption(
                "α = 0.25 (L/D<3) / 0.22 (L/D<5) / 0.20 (L/D≥5).  "
                "Zick parameter a/R ≤ 0.5 ensures the dish head can stiffen "
                "the shell at the saddle horn — per Zick (1951), Welded Pressure Vessels."
            )

            # ── Load & section ────────────────────────────────────────────
            st.markdown("### Vertical reaction & section selection")
            _over = wt_saddle["overstressed"]
            sl1, sl2, sl3, sl4 = st.columns(4)
            sl1.metric("Operating weight",  f"{wt_oper['w_operating_t']:.2f} t")
            sl2.metric("Reaction / saddle", f"{wt_saddle['reaction_t']:.2f} t",
                       delta=f"{wt_saddle['reaction_kN']:.0f} kN", delta_color="off")
            sl3.metric("Catalogue capacity",f"{wt_saddle['capacity_t']} t",
                       delta="⚠️ exceeds max" if _over else "✅ adequate",
                       delta_color="inverse" if _over else "normal")
            sl4.metric("Selected section",  wt_saddle["section"])

            # ── Recommendations panel ─────────────────────────────────────
            _alts = wt_saddle["alternatives"]
            _min_n = wt_saddle["min_n_needed"]

            if _over:
                st.error(
                    f"⚠️ **Load exceeds catalogue maximum** — reaction "
                    f"{wt_saddle['reaction_t']:.1f} t/saddle > {wt_saddle['capacity_t']} t max.  \n"
                    f"Minimum **{_min_n} supports** required to stay within catalogue. "
                    f"Change the **Support type** in the sidebar to a "
                    f"{'3-saddle' if _min_n == 3 else str(_min_n)+'-support'} arrangement."
                )
            else:
                # Optimisation: could fewer saddles still fit within catalogue?
                if wt_saddle["n_saddles"] > 2:
                    _fewer_n = wt_saddle["n_saddles"] - 1
                    _fewer_alt = next(
                        (a for a in _alts if a["n_saddles"] == _fewer_n), None)
                    if _fewer_alt and _fewer_alt["fits_catalogue"]:
                        _save_kg = (wt_saddle["w_one_saddle_kg"]
                                    - _fewer_alt["struct_wt_ea_kg"]) * _fewer_n
                        st.info(
                            f"💡 **Optimisation** — {_fewer_n} supports would also work: "
                            f"{_fewer_alt['reaction_t']:.1f} t/saddle → "
                            f"**{_fewer_alt['section']}** (cap {_fewer_alt['capacity_t']} t). "
                            f"Removing one support saves ~{abs(_save_kg):,.0f} kg of structural steel. "
                            f"Change **Support type** in the sidebar to "
                            f"{'Saddle (2-support)' if _fewer_n == 2 else str(_fewer_n)+'-support'}."
                        )

            # ── Alternatives comparison table ─────────────────────────────
            st.markdown("**Support arrangement alternatives**")
            alt_rows = []
            for a in _alts:
                status = "▶ current" if a["is_current"] else (
                    "✅ fits" if a["fits_catalogue"] else "❌ exceeds max")
                alt_rows.append({
                    "Supports": a["n_saddles"],
                    "Reaction/saddle (t)": a["reaction_t"],
                    "Section": a["section"],
                    "Capacity (t)": a["capacity_t"],
                    "Status": status,
                    "Struct. wt/saddle (kg)": a["struct_wt_ea_kg"],
                    "Total struct. wt (kg)": a["struct_wt_total_kg"],
                    "Saddle positions from left (m)": "  /  ".join(str(p) for p in a["positions_m"]),
                })
            _alt_df = pd.DataFrame(alt_rows)
            st.dataframe(_alt_df, use_container_width=True, hide_index=True)
            st.caption(
                "Positions = distance from left head tangent line. "
                "For 2 saddles: Zick α×L from each end. "
                "For 3+ saddles: outer two at α×L, inner ones evenly distributed. "
                "To change the arrangement, update **Support type** in the sidebar "
                "(🔩 Nozzles & supports expander)."
            )

            st.table(pd.DataFrame([
                ["Vertical reaction per saddle",  f"{wt_saddle['reaction_t']:.3f} t  =  {wt_saddle['reaction_kN']:.1f} kN"],
                ["Selected section",              wt_saddle["section"]],
                ["Section unit weight",           f"{wt_saddle['kg_per_m']:.2f} kg/m"],
                ["Piece length (saddle width)",   f"{wt_saddle['piece_length_m']:.1f} m"],
                ["Piece weight",                  f"{wt_saddle['piece_weight_kg']:.1f} kg"],
                ["Paint area per piece",          f"{wt_saddle['piece_paint_m2']:.2f} m²"],
                ["Ribs per saddle assembly",      str(wt_saddle["n_ribs"])],
                ["Rib mass per saddle",           f"{wt_saddle['w_ribs_kg']:.1f} kg"],
                ["Saddle structure wt (incl. plates +10%)", f"{wt_saddle['w_one_saddle_kg']:.1f} kg"],
                [f"Total {wt_saddle['n_saddles']} saddles structure wt",
                 f"{wt_saddle['w_two_saddles_kg']:.1f} kg"],
                ["Base plate width (est.)",      f"{wt_saddle['base_width_m']:.3f} m"],
                ["Base plate area (est.)",       f"{wt_saddle['base_area_m2']:.3f} m²"],
                ["Foundation bearing pressure",  f"{wt_saddle['bearing_kPa']:.1f} kPa"],
            ], columns=["Parameter", "Value"]))

            # ── Full catalogue ────────────────────────────────────────────
            st.markdown("### Full catalogue — all section types")
            st.dataframe(pd.DataFrame(wt_saddle["catalogue_rows"]),
                         use_container_width=True, hide_index=True)
            st.caption(
                "Ribs/saddle = number of catalogue-section pieces in one saddle assembly "
                f"(for vessel ID = {real_id:.2f} m).  "
                "Structure weight includes +10% for wear plate and base plate.  "
                "Foundation bearing pressure = reaction / (0.8×OD × piece length).  "
                "Full Zick stress check (shell longitudinal bending, circumferential "
                "compression at horn, and web shear) should be performed per "
                "ASME VIII-1 App. L or EN 13445-3 §16 for detailed design."
            )

        with st.expander("9 · Internal lining / coating", expanded=True):
            st.markdown(f"### Internal surface areas")
            ia1, ia2, ia3, ia4 = st.columns(4)
            ia1.metric("Cylinder",       f"{vessel_areas['a_cylinder_m2']:.1f} m²")
            ia2.metric("Two dish ends",  f"{vessel_areas['a_two_heads_m2']:.1f} m²")
            ia3.metric("Nozzle plate",   f"{vessel_areas['a_nozzle_plate_m2']:.1f} m²")
            ia4.metric("Total to coat",  f"{vessel_areas['a_total_m2']:.1f} m²")
            st.table(pd.DataFrame([
                ["Cylinder (shell)",       f"{vessel_areas['a_cylinder_m2']:.2f} m²",
                 f"π × {real_id:.3f} × {cyl_len:.3f}"],
                ["One dish end",           f"{vessel_areas['a_one_head_m2']:.2f} m²",
                 f"{end_geometry}"],
                ["Two dish ends",          f"{vessel_areas['a_two_heads_m2']:.2f} m²", ""],
                ["Shell total",            f"{vessel_areas['a_shell_m2']:.2f} m²",  ""],
                ["Nozzle plate (internal)",f"{vessel_areas['a_nozzle_plate_m2']:.2f} m²", ""],
                ["Total internal area",    f"{vessel_areas['a_total_m2']:.2f} m²",  ""],
            ], columns=["Surface", "Area", "Basis"]))
            st.caption(
                "Elliptic 2:1 head area: exact oblate-spheroid integral. "
                "Nozzle plate: flat plate area from nozzle plate design module. "
                "Nozzle stubs and manhole necks not included."
            )

            st.markdown(f"### Protection system — {protection_type}")
            if protection_type == "None":
                st.info("No internal lining or coating selected. Vessel relies on "
                        "corrosion allowance only.")
            else:
                lc1, lc2, lc3, lc4 = st.columns(4)
                lc1.metric("Area protected",  f"{lining_result['a_total_m2']:.1f} m²")
                lc2.metric("Lining weight",   f"{lining_result['weight_kg']:,.0f} kg")
                lc3.metric("Material cost",   f"USD {lining_result['material_cost_usd']:,.0f}")
                lc4.metric("Total cost",      f"USD {lining_result['total_cost_usd']:,.0f}",
                           delta=f"Labour: USD {lining_result['labor_cost_usd']:,.0f}",
                           delta_color="off")

                if lining_result["id_deduction_mm"] > 0:
                    st.info(
                        f"**Hydraulic ID deduction:** "
                        f"2 × {lining_mm:.1f} mm × {rubber_layers} layers "
                        f"= **{lining_result['id_deduction_mm'] * rubber_layers:.1f} mm** "
                        f"→ real ID = **{real_id:.4f} m**"
                    )
                else:
                    st.success(
                        f"**{protection_type}** — coating thickness "
                        f"{lining_result['thickness_mm']:.2f} mm is negligible; "
                        "no hydraulic ID deduction applied."
                    )

                st.markdown("**Specification & cost breakdown**")
                detail_rows = [[k, v] for k, v in lining_result["detail"].items()]
                st.table(pd.DataFrame(detail_rows, columns=["Parameter", "Value"]))

            # ── Reference tables ──────────────────────────────────────────
            with st.expander("Reference — available types & default costs", expanded=False):
                st.markdown("**Rubber lining types**")
                st.dataframe(pd.DataFrame([
                    {"Type": k, "Density (kg/m³)": v["density_kg_m3"],
                     "Default cost (USD/m²)": v["default_cost_m2"],
                     "Hardness": v["hardness"]}
                    for k, v in RUBBER_TYPES.items()
                ]), use_container_width=True, hide_index=True)

                st.markdown("**Epoxy coating types**")
                st.dataframe(pd.DataFrame([
                    {"Type": k, "Default DFT (µm)": v["default_dft_um"],
                     "Default coats": v["default_coats"],
                     "Cost (USD/m²)": v["default_cost_m2"],
                     "Note": v["note"]}
                    for k, v in EPOXY_TYPES.items()
                ]), use_container_width=True, hide_index=True)

                st.markdown("**Ceramic coating types**")
                st.dataframe(pd.DataFrame([
                    {"Type": k,
                     "Default DFT/coat (µm)": v["default_dft_um"],
                     "Default coats": v["default_coats"],
                     "Density (kg/m³)": v["density_kg_m3"],
                     "Cost (USD/m²)": v["default_cost_m2"],
                     "Note": v["note"]}
                    for k, v in CERAMIC_TYPES.items()
                ]), use_container_width=True, hide_index=True)

                st.caption(
                    "All default costs are indicative supply-only rates for Middle East / SE Asia markets "
                    "(2024). Add mobilisation, surface preparation (Sa 2.5 / SSPC-SP10), "
                    "QA/QC inspection, and cure/vulcanisation costs as project-specific items."
                )

    # ─────────────────────────────────────────────────────────────────────
    # TAB 8 · CARTRIDGE
    # ─────────────────────────────────────────────────────────────────────
    with tab_cart:
        st.subheader("Cartridge (polishing) filter sizing")

        # ── 1. Key metrics ────────────────────────────────────────────────
        with st.expander("1 · Sizing — selected configuration", expanded=True):
            ca1, ca2, ca3, ca4 = st.columns(4)
            ca1.metric("Elements required",  str(cart_result["n_elements"]))
            ca2.metric("Housings required",
                       str(cart_result["n_housings"]),
                       delta=f"{cart_result['n_elem_per_housing']} elem./housing",
                       delta_color="off")
            ca3.metric("Flow / element",
                       f"{cart_result['actual_flow_m3h_element']:.3f} m³/h",
                       delta=f"{cart_result['q_lpm_element']:.1f} lpm",
                       delta_color="off")
            ca4.metric("Dirt hold / element",
                       f"{cart_result['dhc_g_element']:.0f} g",
                       delta=f"{cart_result['element_ties']} TIE",
                       delta_color="off")

            _cip_badge = "🔩 SS 316L — CIP" if cart_result["is_CIP_system"] else "🔵 Polymer — standard"
            st.caption(
                f"{_cip_badge}  |  "
                f"Element: **{cart_size}** ({cart_result['element_ties']} TIE)  |  "
                f"Rating: **{cart_rating} µm absolute**  |  "
                f"Area: {cart_result['element_area_m2']} m²/element  |  "
                f"Feed viscosity: {_cart_mu_cP:.2f} cP  |  "
                f"Safety factor: {cart_result['safety_factor']}×"
            )
            st.table(pd.DataFrame([
                ["Design flow",                f"{cart_result['design_flow_m3h']:,.1f} m³/h"],
                ["Element size",               cart_result["element_size"]],
                ["Element area",               f"{cart_result['element_area_m2']} m²"],
                ["Rating",                     f"{cart_result['rating_um']} µm absolute"],
                ["Base capacity (1 cP)",       f"{cart_result['cap_m3h_element_base']:.3f} m³/h/element"],
                ["Capacity (derated @ μ)",     f"{cart_result['cap_m3h_element_visc']:.3f} m³/h/element"],
                ["Capacity (rated w/ 1.5 SF)", f"{cart_result['cap_m3h_element_rated']:.3f} m³/h/element"],
                ["Elements required",          str(cart_result["n_elements"])],
                ["Elements per housing",       str(cart_result["n_elem_per_housing"])],
                ["Housings required",          str(cart_result["n_housings"])],
                ["Actual flow / element",      f"{cart_result['actual_flow_m3h_element']:.3f} m³/h"],
                ["Actual flux density",        f"{cart_result['actual_flow_m3h_m2']:.3f} m³/h/m²"],
                ["Flow / element",             f"{cart_result['q_lpm_element']:.1f} lpm"],
            ], columns=["Parameter", "Value"]))

        # ── 2. ΔP & performance ───────────────────────────────────────────
        with st.expander("2 · ΔP & performance", expanded=True):
            cb1, cb2, cb3, cb4 = st.columns(4)
            cb1.metric("ΔP clean (BOL)",  f"{cart_result['dp_clean_bar']*1000:.1f} mbar")
            cb2.metric("ΔP average",      f"{cart_result['dp_avg_bar']*1000:.0f} mbar")
            cb3.metric("ΔP EOL",          f"{cart_result['dp_eol_bar']:.2f} bar")
            cb4.metric("DHC / element",   f"{cart_result['dhc_g_element']:.0f} g")

            _dp_pct = cart_result['dp_clean_bar'] / DP_REPLACEMENT_BAR * 100
            if cart_result['dp_overloaded']:
                st.error(
                    f"⚠️ BOL ΔP ({cart_result['dp_clean_bar']:.3f} bar) already exceeds the "
                    f"replacement trigger ({DP_REPLACEMENT_BAR:.2f} bar). "
                    "Increase number of elements or choose a longer element."
                )
            else:
                st.caption(
                    f"ΔP model: vendor quadratic (40\" reference + TIE-ratio scaling).  "
                    f"BOL = {cart_result['dp_clean_bar']*1000:.1f} mbar at "
                    f"{cart_result['q_lpm_element']:.1f} lpm/element "
                    f"({_dp_pct:.0f} % of EOL trigger).  "
                    f"ΔP grows linearly with cake mass → {DP_REPLACEMENT_BAR:.2f} bar at DHC."
                )

            # ΔP vs cake-mass curve
            _dp_df = pd.DataFrame(cart_result["dp_curve"]).set_index("mass_g")
            _dp_df["ΔP (mbar)"]       = _dp_df["dp_bar"] * 1000
            _dp_df["EOL trigger (mbar)"] = DP_REPLACEMENT_BAR * 1000
            st.caption("ΔP progression — accumulated cake mass per element (g) vs differential pressure (mbar)")
            st.line_chart(_dp_df[["ΔP (mbar)", "EOL trigger (mbar)"]])

        # ── 3. TSS loading & replacement interval ─────────────────────────
        with st.expander("3 · TSS loading & replacement interval", expanded=True):
            cd1, cd2, cd3, cd4 = st.columns(4)
            cd1.metric("CF inlet TSS",   f"{cart_result['cf_inlet_tss_mg_l']:.2f} mg/L")
            cd2.metric("CF outlet TSS",  f"{cart_result['cf_outlet_tss_mg_l']:.2f} mg/L",
                       delta=f"{cart_result['tss_removal_pct']:.0f} % removal",
                       delta_color="off")
            cd3.metric("Loading rate",   f"{cart_result['loading_g_h_element']:.4f} g/h/elem")
            cd4.metric("Replacement",    f"{cart_result['replacement_freq_days']:.0f} days",
                       delta=f"{cart_result['replacements_per_year']:.2f}/yr",
                       delta_color="off")

            _zero_load = cart_result['loading_g_h_element'] < 1e-9
            st.caption(
                f"Mass balance: ({cart_result['cf_inlet_tss_mg_l']:.2f} − "
                f"{cart_result['cf_outlet_tss_mg_l']:.2f}) mg/L × "
                f"{cart_result['actual_flow_m3h_element']:.3f} m³/h/element = "
                f"{cart_result['loading_g_h_element']:.4f} g/h per element.  "
                f"DHC {cart_result['dhc_g_element']:.0f} g ÷ loading → "
                + ("fallback (zero TSS removed — rating-table used)."
                   if _zero_load else
                   f"{cart_result['replacement_freq_days']:.0f} days replacement interval.")
            )
            st.table(pd.DataFrame([
                ["CF inlet TSS",         f"{cart_result['cf_inlet_tss_mg_l']:.2f} mg/L",
                 "MMF effluent entering CF"],
                ["CF outlet TSS",        f"{cart_result['cf_outlet_tss_mg_l']:.2f} mg/L",
                 "Quality target"],
                ["TSS removed",          f"{cart_result['tss_removed_mg_l']:.2f} mg/L",
                 f"{cart_result['tss_removal_pct']:.0f} %"],
                ["Loading rate",         f"{cart_result['loading_g_h_element']:.4f} g/h/element",
                 "Per element at design flow"],
                ["DHC per element",      f"{cart_result['dhc_g_element']:.0f} g",
                 f"{'45 g/TIE (SS 316L)' if cart_result['is_CIP_system'] else '30 g/TIE (polymer)'}"],
                ["Replacement interval", f"{cart_result['replacement_freq_days']:.0f} days",
                 f"= {cart_result['interval_h']:.0f} h"],
                ["Replacements/year",    f"{cart_result['replacements_per_year']:.2f}",
                 "All elements"],
            ], columns=["Parameter", "Value", "Note"]))

        # ── 4. Length optimisation table ──────────────────────────────────
        with st.expander("3 · Element length optimisation", expanded=True):
            st.caption(
                f"Comparing all standard lengths at **{cart_rating} µm** rating, "
                f"μ = **{_cart_mu_cP:.2f} cP**, SF = {cart_result['safety_factor']}×.  "
                "Market round = smallest standard housing size ≥ n_elements "
                "(fits in 1 housing where possible).  "
                "🏆 = recommended (fewest housings)."
            )
            _opt_rows = []
            for r in cart_optim:
                flag = "🏆" if r["is_recommended"] else ""
                sel  = " ◀" if r["size"] == cart_size else ""
                _opt_rows.append({
                    "":              flag,
                    "Length":        r["size"] + sel,
                    "TIE":           r["ties"],
                    "Cap. (m³/h/el)": f"{r['cap_m3h_elem']:.3f}",
                    "Elements":      r["n_elements"],
                    "Market round":  r["market_round"],
                    "Housings":      r["n_housings"],
                    "Fill %":        f"{r['fill_pct']:.0f} %",
                    "lpm/element":   f"{r['q_lpm_element']:.1f}",
                    "ΔP clean (bar)": f"{r['dp_clean_bar']:.4f}",
                    "ΔP EOL (bar)":  f"{r['dp_eol_bar']:.4f}",
                    "DHC/el (g)":    f"{r['dhc_g']:.0f}",
                })
            st.dataframe(
                pd.DataFrame(_opt_rows),
                use_container_width=True,
                hide_index=True,
            )

        # ── 5. Economics ──────────────────────────────────────────────────
        with st.expander("5 · Economics", expanded=True):
            # Editable unit-price table — all element sizes × ratings
            st.caption(
                "Unit prices below are mid-market indicative (2024).  "
                "**Edit any cell** to override — the active selection "
                f"(**{cart_size} · {cart_rating} µm**) is highlighted."
            )
            _base_table = COST_TABLE_SS316L if cart_cip else COST_TABLE_POLYMER
            _price_rows = [
                {
                    "Element":     size,
                    "1 µm (USD)":  int(_base_table.get((size, 1),  0)),
                    "5 µm (USD)":  int(_base_table.get((size, 5),  0)),
                    "10 µm (USD)": int(_base_table.get((size, 10), 0)),
                }
                for size in ELEMENT_SIZE_LABELS
            ]
            _edited_prices = st.data_editor(
                pd.DataFrame(_price_rows),
                key=f"cart_prices_{'cip' if cart_cip else 'std'}",
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Element":     st.column_config.TextColumn("Element", disabled=True),
                    "1 µm (USD)":  st.column_config.NumberColumn(
                                       "1 µm (USD)",  min_value=0, step=1, format="$ %d"),
                    "5 µm (USD)":  st.column_config.NumberColumn(
                                       "5 µm (USD)",  min_value=0, step=1, format="$ %d"),
                    "10 µm (USD)": st.column_config.NumberColumn(
                                       "10 µm (USD)", min_value=0, step=1, format="$ %d"),
                },
            )

            # Pull out price for the active (element_size, rating_um) selection
            _rating_col  = f"{cart_rating} µm (USD)"
            _sel_row     = _edited_prices[_edited_prices["Element"] == cart_size]
            _unit_price  = (int(_sel_row[_rating_col].values[0])
                            if not _sel_row.empty else cart_result["cost_per_element_usd"])
            _annual_cost = cart_result["n_elements"] * cart_result["replacements_per_year"] * _unit_price

            ce1, ce2, ce3 = st.columns(3)
            ce1.metric("Cost / element",  f"USD {_unit_price:,.0f}",
                       delta="active selection", delta_color="off")
            ce2.metric("Changes / year",  f"{cart_result['replacements_per_year']:.2f}")
            ce3.metric("Annual cost",     f"USD {_annual_cost:,.0f}")

            st.table(pd.DataFrame([
                ["Element",              f"{cart_size} · {cart_rating} µm"],
                ["Cost / element",       f"USD {_unit_price:,.0f}"],
                ["Replacement interval", f"{cart_result['replacement_freq_days']:.0f} days"],
                ["Changes / year",       f"{cart_result['replacements_per_year']:.2f}"],
                ["Total elements",       str(cart_result["n_elements"])],
                ["Annual cost",          f"USD {_annual_cost:,.0f}"],
            ], columns=["Item", "Value"]))

    # ─────────────────────────────────────────────────────────────────────
    # TAB 9 · ENERGY
    # ─────────────────────────────────────────────────────────────────────
    with tab_energy:
        st.subheader("Energy & hydraulic profile")

        # ── 1. Hydraulic head budget ──────────────────────────────────────
        with st.expander("1 · Filtration pump — head budget", expanded=True):
            st.caption(
                "Flow path: Pump → inlet piping → distributor → media → "
                "strainer nozzle plate → outlet piping → downstream pressure.  "
                f"Media: clean (Ergun) / dirty (cake at M_max = {solid_loading:.2f} kg/m²).  "
                f"Downstream residual = {p_residual:.2f} barg."
            )
            hA, hB = st.columns(2)
            for col, state, bud in [
                (hA, "Clean bed", hyd_prof["clean"]),
                (hB, f"Dirty bed (M_max = {solid_loading:.2f} kg/m²)", hyd_prof["dirty"]),
            ]:
                with col:
                    st.markdown(f"**{state}**")
                    rows = []
                    for component, bar_val in bud["items_bar"].items():
                        mwc_val = bud["items_mwc"][component]
                        rows.append([component,
                                     f"{bar_val:.4f} bar",
                                     f"{mwc_val:.2f} mWC"])
                    rows.append(["**TOTAL pump duty**",
                                 f"**{bud['total_bar']:.4f} bar**",
                                 f"**{bud['total_mwc']:.2f} mWC**"])
                    st.dataframe(pd.DataFrame(rows, columns=["Component", "bar", "mWC"]),
                                 use_container_width=True, hide_index=True)

        # ── 2. Power summary ─────────────────────────────────────────────
        with st.expander("2 · Power — per consumer", expanded=True):
            e1, e2, e3, e4 = st.columns(4)
            e1.metric("Filtration pump\n(dirty, per filter)",
                      f"{energy['p_filt_dirty_kw']:.1f} kW")
            e2.metric("Filtration pump\n(clean, per filter)",
                      f"{energy['p_filt_clean_kw']:.1f} kW")
            e3.metric("BW pump\n(per event)",
                      f"{energy['p_bw_kw']:.1f} kW")
            e4.metric("Air blower\n(electrical, per event)",
                      f"{energy['p_blower_elec_kw']:.1f} kW")

            st.markdown("**Power basis**")
            st.dataframe(pd.DataFrame([
                ["Filtration pump", f"{q_per_filter:.1f} m³/h × {_n_total_filters} filters",
                 f"η_pump={pump_eta:.2f} · η_motor={motor_eta:.2f}",
                 f"H_clean={hyd_prof['clean']['total_mwc']:.1f} mWC",
                 f"H_dirty={hyd_prof['dirty']['total_mwc']:.1f} mWC",
                 f"{energy['p_filt_avg_kw']:.1f} kW avg"],
                ["BW pump",         f"{bw_hyd['q_bw_design_m3h']:.0f} m³/h",
                 f"η_pump={bw_pump_eta:.2f} · η_motor={motor_eta:.2f}",
                 f"H={bw_head_mwc:.1f} mWC", "—",
                 f"{energy['p_bw_kw']:.1f} kW"],
                ["Air blower",      f"{bw_hyd['q_air_design_m3h']:.0f} m³/h air",
                 f"η_motor={motor_eta:.2f}",
                 "ΔP=0.5 bar", "—",
                 f"{energy['p_blower_elec_kw']:.1f} kW elec."],
            ], columns=["Consumer", "Flow", "Efficiency",
                        "Head / ΔP (clean)", "Head / ΔP (dirty)", "Power"]),
            use_container_width=True, hide_index=True)

        # ── 3. Annual energy & cost ───────────────────────────────────────
        with st.expander("3 · Annual energy & operating cost", expanded=True):
            st.caption(
                f"Basis: {op_hours_yr:,.0f} operating hours/year · "
                f"electricity = USD {elec_tariff:.3f}/kWh · "
                f"BW: {energy['bw_per_day_design']:.1f} cycles/filter/day × "
                f"{bw_total_min} min · {_n_total_filters} filters total."
            )
            fa, fb, fc, fd = st.columns(4)
            fa.metric("Filtration pump", f"{energy['e_filt_kwh_yr']/1e3:,.0f} MWh/yr")
            fb.metric("BW pump",         f"{energy['e_bw_pump_kwh_yr']/1e3:,.1f} MWh/yr")
            fc.metric("Air blower",      f"{energy['e_blower_kwh_yr']/1e3:,.1f} MWh/yr")
            fd.metric("TOTAL",           f"{energy['e_total_kwh_yr']/1e3:,.0f} MWh/yr")

            g1, g2, g3 = st.columns(3)
            g1.metric("Specific energy",  f"{energy['kwh_per_m3']:.4f} kWh/m³")
            g2.metric("Annual OPEX (energy)",
                      f"USD {energy['cost_usd_yr']:,.0f}")
            g3.metric("Water treated / yr",
                      f"{energy['total_flow_m3_yr']/1e6:.2f} Mm³")

            st.dataframe(pd.DataFrame([
                ["Filtration pump",   f"{energy['e_filt_kwh_yr']:,.0f}",
                 f"{energy['e_filt_kwh_yr']*elec_tariff:,.0f}"],
                ["BW pump",           f"{energy['e_bw_pump_kwh_yr']:,.0f}",
                 f"{energy['e_bw_pump_kwh_yr']*elec_tariff:,.0f}"],
                ["Air blower",        f"{energy['e_blower_kwh_yr']:,.0f}",
                 f"{energy['e_blower_kwh_yr']*elec_tariff:,.0f}"],
                ["TOTAL",             f"{energy['e_total_kwh_yr']:,.0f}",
                 f"{energy['cost_usd_yr']:,.0f}"],
            ], columns=["Consumer", "Energy (kWh/yr)", "Cost (USD/yr)"]),
            use_container_width=True, hide_index=True)

    # ─────────────────────────────────────────────────────────────────────
    # TAB 10 · SECTION DRAWING
    # ─────────────────────────────────────────────────────────────────────
    with tab_section:
        st.subheader("Theoretical elevation section")

        _show_exp = st.toggle("Show expanded bed (BW condition)", value=True,
                              key="sec_exp")

        _sec_fig = vessel_section_elevation(
            vessel_id_m      = nominal_id,
            total_length_m   = total_length,
            h_dish_m         = h_dish,
            nozzle_plate_h_m = nozzle_plate_h,
            layers           = layers,
            collector_h_m    = collector_h,
            bw_exp           = bw_exp,
            show_expansion   = _show_exp,
        )
        st.pyplot(_sec_fig, use_container_width=True)

        # Download as PNG
        _sec_buf = io.BytesIO()
        _sec_fig.savefig(_sec_buf, format="png", dpi=180, bbox_inches="tight",
                         facecolor=_sec_fig.get_facecolor())
        _sec_buf.seek(0)
        st.download_button(
            "⬇️ Download PNG",
            data=_sec_buf,
            file_name=f"{project_name or 'MMF'}_section.png",
            mime="image/png",
        )

        st.caption(
            "Elevation view — not to scale in the thickness direction.  "
            "Media layers shown as uniform-height horizontal bands.  "
            "Expanded bed shown for the water-only BW velocity input."
        )

    # ─────────────────────────────────────────────────────────────────────
    # TAB 11 · DATASHEET
    # ─────────────────────────────────────────────────────────────────────
    with tab_datasheet:
        st.subheader("Vessel fabrication data sheet")

        import datetime as _dt
        _today = _dt.date.today().strftime("%d-%b-%Y")

        # ── Title block ───────────────────────────────────────────────────
        tb1, tb2, tb3 = st.columns([2, 2, 1])
        with tb1:
            st.markdown("**Document**")
            st.table(pd.DataFrame([
                ["Project",     project_name],
                ["Doc. No.",    doc_number],
                ["Description", "Horizontal Multi-Media Filter"],
                ["Vessel/Tag",  f"MMF-001"],
            ], columns=["Field", "Value"]))
        with tb2:
            st.markdown("**Approval**")
            st.table(pd.DataFrame([
                ["Date",        _today],
                ["Revision",    revision],
                ["Client",      client or "—"],
                ["Prepared by", engineer],
            ], columns=["Field", "Value"]))
        with tb3:
            st.metric("Filters", streams * n_filters)
            st.metric("Active (N)", streams * n_filters - redundancy * streams)

        st.divider()

        ds_left, ds_right = st.columns([1, 1.1])

        # ── LEFT: Design data ─────────────────────────────────────────────
        with ds_left:
            st.markdown("### Design Data")
            _shell_E = mech["shell_E"]
            _head_E  = mech["head_E"]
            st.table(pd.DataFrame([
                ["Design Code",           "ASME VIII Div. 1"],
                ["Operating pressure",    f"{design_pressure * 0.7:.2f} barg  (≈ 70 % design)"],
                ["Design pressure",       f"{design_pressure:.2f} barg"],
                ["Test pressure (hydro)", f"{design_pressure * 1.5:.2f} barg"],
                ["Design temperature",    f"{design_temp:.0f} °C"],
                ["Min. design temp.",     "5 °C (ambient min.)"],
                ["Corrosion allowance",   f"{corrosion:.1f} mm"],
                ["Shell joint eff. E",    f"{_shell_E:.2f}  ({shell_radio})"],
                ["Head joint eff. E",     f"{_head_E:.2f}  ({head_radio})"],
                ["Head shape",            end_geometry],
            ], columns=["Parameter", "Value"]))

            st.markdown("### Material")
            st.table(pd.DataFrame([
                ["Shell & heads",         material_name],
                ["Allowable stress S",    f"{mech['allowable_stress']} kg/cm²"],
                ["Standard",              mat_info["standard"]],
                ["Nozzle plate",          material_name],
                ["Internal protection",   protection_type],
            ], columns=["Component", "Specification"]))

            st.markdown("### Dimensional Data")
            st.table(pd.DataFrame([
                ["Internal diameter (ID)",   f"{nominal_id * 1000:.0f} mm"],
                ["Outside diameter (OD)",    f"{mech['od_m'] * 1000:.0f} mm"],
                ["Shell thickness (t_des)",  f"{mech['t_shell_design_mm']} mm"],
                ["Head thickness (t_des)",   f"{mech['t_head_design_mm']} mm"],
                ["T/T length",               f"{total_length:.3f} m"],
                ["Cylindrical length",       f"{cyl_len:.3f} m"],
                ["Dish depth",               f"{h_dish * 1000:.0f} mm"],
                ["Nozzle plate height",      f"{nozzle_plate_h * 1000:.0f} mm"],
                ["Nozzle plate thickness",   f"{wt_np['t_design_mm']} mm"],
            ], columns=["Parameter", "Value"]))

            st.markdown("### Capacity & Weight")
            _lining_label = (protection_type if protection_type != "None"
                             else "None (bare metal)")
            st.table(pd.DataFrame([
                ["Fluid",                    "Seawater / filtered feed"],
                ["Internal volume",          f"{wt_oper['v_total_internal_m3']:.2f} m³"],
                ["Empty weight",             f"{w_total:,.0f} kg  =  {w_total/1000:.3f} t"],
                ["Internal lining/coating",  _lining_label],
                ["Lining weight",            f"{lining_result['weight_kg']:,.0f} kg"],
                ["Operating weight (full)",  f"{wt_oper['w_operating_kg']:,.0f} kg  "
                                             f"=  {wt_oper['w_operating_t']:.3f} t"],
                ["Load / support",           f"{wt_oper['load_per_support_kN']:.1f} kN  "
                                             f"({wt_oper['load_per_support_t']:.3f} t)"],
                ["Head thickness",           f"{mech['t_head_design_mm']} mm"],
                ["Shell thickness",          f"{mech['t_shell_design_mm']} mm"],
                ["Floor (nozzle plate)",     f"{wt_np['t_design_mm']} mm"],
                ["Est. head weight (×2)",    f"{wt_body['weight_two_heads_kg']:,.0f} kg"],
                ["Est. shell weight",        f"{wt_body['weight_shell_kg']:,.0f} kg"],
                ["Est. floor/plate weight",  f"{wt_np['weight_total_kg']:,.0f} kg"],
            ], columns=["Item", "Value"]))

        # ── RIGHT: section drawing + media + nozzles ───────────────────────
        with ds_right:
            st.markdown("### Elevation View — Theoretical Section")
            _ds_fig = vessel_section_elevation(
                vessel_id_m      = nominal_id,
                total_length_m   = total_length,
                h_dish_m         = h_dish,
                nozzle_plate_h_m = nozzle_plate_h,
                layers           = layers,
                collector_h_m    = collector_h,
                bw_exp           = bw_exp,
                show_expansion   = True,
                figsize          = (10, 4.5),
            )
            st.pyplot(_ds_fig, use_container_width=True)

            st.markdown("### Media Bed")
            _media_rows = []
            _cum = nozzle_plate_h
            for lyr in layers:
                _media_rows.append({
                    "Layer":        lyr["Type"],
                    "Depth (mm)":   int(lyr["Depth"] * 1000),
                    "Bottom (mm)":  int(_cum * 1000),
                    "Top (mm)":     int((_cum + lyr["Depth"]) * 1000),
                    "d10 (mm)":     lyr.get("d10", "—"),
                    "CU":           lyr.get("cu", "—"),
                    "ε₀":           lyr.get("epsilon0", "—"),
                    "ρ (kg/m³)":    lyr.get("rho_p_eff", "—"),
                })
                _cum += lyr["Depth"]
            st.dataframe(pd.DataFrame(_media_rows),
                         use_container_width=True, hide_index=True)

        st.divider()
        st.markdown("### Nomenclature of Nozzles")
        st.dataframe(pd.DataFrame(nozzle_sched),
                     use_container_width=True, hide_index=True)

        st.divider()
        st.markdown("### Vessel Standards")
        st.table(pd.DataFrame([
            ["Design Code",        "ASME Boiler & Pressure Vessel Code Sect. VIII Div. 1"],
            ["Material (shell)",   mat_info["standard"]],
            ["Flanges",            "ASME B16.5 / EN 1092-1"],
            ["Nozzles",            "ASME B36.10 / B36.19"],
            ["Welding",            "ASME IX — Welding & Brazing Qualifications"],
            ["NDE",               f"Shell: {shell_radio}   Head: {head_radio}"],
            ["Corrosion allow.",   f"{corrosion:.1f} mm"],
            ["Internal coating",   protection_type],
        ], columns=["Standard / Item", "Reference / Value"]))

        # ── DOCX export ───────────────────────────────────────────────────
        if _DOCX_OK:
            st.divider()
            if st.button("📄 Generate Datasheet DOCX", key="ds_docx_btn"):
                _ds_doc = _DocxDocument()
                _ds_doc.add_heading("VESSEL FABRICATION DATA SHEET", 0)

                # Title block
                _ds_doc.add_heading("1. Document Information", 1)
                _t = _ds_doc.add_table(rows=1, cols=2)
                _t.style = "Table Grid"
                _t.rows[0].cells[0].text = "Project"
                _t.rows[0].cells[1].text = project_name
                for _label, _val in [
                    ("Doc. No.", doc_number), ("Revision", revision),
                    ("Date", _today), ("Prepared by", engineer),
                    ("Client", client or "—"),
                ]:
                    _r = _t.add_row()
                    _r.cells[0].text = _label
                    _r.cells[1].text = _val

                # Design data
                _ds_doc.add_heading("2. Design Data", 1)
                _t2 = _ds_doc.add_table(rows=1, cols=2)
                _t2.style = "Table Grid"
                _t2.rows[0].cells[0].text = "Parameter"
                _t2.rows[0].cells[1].text = "Value"
                for _p, _v in [
                    ("Design Code", "ASME VIII Div. 1"),
                    ("Design pressure", f"{design_pressure:.2f} barg"),
                    ("Design temperature", f"{design_temp:.0f} °C"),
                    ("Corrosion allowance", f"{corrosion:.1f} mm"),
                    ("Material", material_name),
                    ("Shell thickness", f"{mech['t_shell_design_mm']} mm"),
                    ("Head thickness", f"{mech['t_head_design_mm']} mm"),
                    ("ID", f"{nominal_id * 1000:.0f} mm"),
                    ("OD", f"{mech['od_m'] * 1000:.0f} mm"),
                    ("T/T length", f"{total_length:.3f} m"),
                    ("Empty weight", f"{w_total:,.0f} kg"),
                    ("Operating weight", f"{wt_oper['w_operating_kg']:,.0f} kg"),
                    ("Internal protection", protection_type),
                ]:
                    _r2 = _t2.add_row()
                    _r2.cells[0].text = _p
                    _r2.cells[1].text = _v

                # Nozzle schedule
                _ds_doc.add_heading("3. Nozzle Schedule", 1)
                _cols_noz = list(pd.DataFrame(nozzle_sched).columns)
                _t3 = _ds_doc.add_table(rows=1, cols=len(_cols_noz))
                _t3.style = "Table Grid"
                for ci, cn in enumerate(_cols_noz):
                    _t3.rows[0].cells[ci].text = cn
                for _, _noz_row in pd.DataFrame(nozzle_sched).iterrows():
                    _r3 = _t3.add_row()
                    for ci, cn in enumerate(_cols_noz):
                        _r3.cells[ci].text = str(_noz_row[cn])

                # Save to buffer
                _ds_buf = io.BytesIO()
                _ds_doc.save(_ds_buf)
                _ds_buf.seek(0)
                st.download_button(
                    "⬇️ Download Datasheet DOCX",
                    data=_ds_buf,
                    file_name=f"{doc_number or 'MMF_datasheet'}.docx",
                    mime="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document",
                    key="ds_docx_dl",
                )
        else:
            st.info("Install python-docx to enable DOCX export.")

    # ─────────────────────────────────────────────────────────────────────
    # TAB 12 · REPORT
    # ─────────────────────────────────────────────────────────────────────
    with tab_report:

        w_total_rep = (wt_body["weight_body_kg"] + w_noz
                       + wt_np["weight_total_kg"]
                       + wt_sup["weight_all_supports_kg"]
                       + wt_int["weight_internals_kg"])

        # ══════════════════════════════════════════════════════════════════
        # SECTION SELECTOR
        # ══════════════════════════════════════════════════════════════════
        st.markdown("### Report builder")
        st.caption(
            "Select the sections to include. "
            "Identification (cover, project info, sign-off) is always present."
        )

        _has_lining = lining_result["protection_type"] != "None"

        sel_c1, sel_c2, sel_c3, sel_c4 = st.columns(4)

        with sel_c1:
            st.markdown("**B · Process Design**")
            s_process  = st.checkbox("Process basis",           value=True, key="rs_proc")
            s_water    = st.checkbox("Water properties",        value=True, key="rs_water")
            s_media    = st.checkbox("Media configuration",     value=True, key="rs_media")
            s_dp       = st.checkbox("Filtration ΔP",          value=True, key="rs_dp")
            s_cycle    = st.checkbox("Filtration cycle & BW feasibility",
                                     value=True, key="rs_cycle")

        with sel_c2:
            st.markdown("**C · Mechanical & Structural**")
            s_vessel   = st.checkbox("Vessel dimensions & ASME thickness",
                                     value=True, key="rs_vessel")
            s_nzpl     = st.checkbox("Nozzle plate design",    value=True, key="rs_nzpl")
            s_wt_empty = st.checkbox("Empty weight breakdown", value=True, key="rs_wtempty")
            s_wt_oper  = st.checkbox("Operating weight & support loads",
                                     value=True, key="rs_wtoper")
            s_saddle   = st.checkbox("Saddle design",          value=True, key="rs_saddle")

        with sel_c3:
            st.markdown("**D · Backwash System**")
            s_bw_hyd   = st.checkbox("BW hydraulics & collector",
                                     value=True, key="rs_bwhyd")
            s_bw_equip = st.checkbox("BW equipment sizing\n(pump / blower / tank)",
                                     value=True, key="rs_bweq")
            st.markdown("**E · Internal Protection**")
            _lining_label = (f"Lining / coating  ({lining_result['protection_type']})"
                             if _has_lining else "Lining / coating  *(None — skipped)*")
            s_lining   = st.checkbox(_lining_label, value=_has_lining,
                                     disabled=not _has_lining, key="rs_lining")

        with sel_c4:
            st.markdown("**F · Energy & Economics**")
            s_hyd_prof = st.checkbox("Hydraulic head profile", value=True, key="rs_hyd")
            s_energy   = st.checkbox("Energy & OPEX",          value=True, key="rs_energy")
            st.markdown("**G · Post-treatment**")
            s_cart     = st.checkbox("Cartridge filter",       value=True, key="rs_cart")

        st.divider()

        # ══════════════════════════════════════════════════════════════════
        # WORD EXPORT
        # ══════════════════════════════════════════════════════════════════
        def _build_docx() -> bytes:
            doc = _DocxDocument()

            def _tbl(rows_data, cols=("Parameter", "Value")):
                tbl = doc.add_table(rows=len(rows_data), cols=len(cols))
                tbl.style = "Table Grid"
                for i, row_vals in enumerate(rows_data):
                    for j, v in enumerate(row_vals):
                        tbl.rows[i].cells[j].text = str(v)
                doc.add_paragraph("")
                return tbl

            def _h(text, lvl=2):
                doc.add_heading(text, lvl)

            # ── A. Cover (always) ───────────────────────────────────────
            t = doc.add_heading("AQUASIGHT™  Horizontal Multi-Media Filter", 0)
            t.alignment = _WD_ALIGN.CENTER
            t2 = doc.add_heading("Calculation Report", 1)
            t2.alignment = _WD_ALIGN.CENTER
            doc.add_paragraph("")
            _h("Project Information")
            _tbl([
                ("Project",       project_name),
                ("Document No.",  f"{doc_number}  ·  Rev {revision}"),
                ("Client",        client or "—"),
                ("Prepared by",   engineer),
                ("Date",          str(__import__("datetime").date.today())),
            ])

            # Track section number dynamically
            _sn = [1]
            def _sec(title):
                _sn[0] += 1
                _h(f"{_sn[0]}. {title}")

            # ── B. Process Design ───────────────────────────────────────
            if s_process:
                _sec("Process Basis")
                _tbl([
                    ("Total plant flow",          f"{total_flow:,.0f} m³/h"),
                    ("Streams",                   str(streams)),
                    ("Filters / stream",          str(n_filters)),
                    ("Redundancy",                f"N-{redundancy} per stream"),
                    ("Flow / filter (N)",         f"{q_per_filter:.1f} m³/h"),
                    ("Filtration rate (N)",       f"{q_per_filter/avg_area:.2f} m/h"),
                    ("Cross-sectional area",      f"{avg_area:.3f} m²"),
                ])

            if s_water:
                _sec("Water Properties")
                _tbl([
                    ("",            "Feed",                        "Backwash"),
                    ("Salinity",    f"{feed_sal:.2f} ppt",        f"{bw_sal:.2f} ppt"),
                    ("Temperature", f"{feed_temp:.1f} °C",        f"{bw_temp:.1f} °C"),
                    ("Density",     f"{rho_feed:.3f} kg/m³",      f"{rho_bw:.3f} kg/m³"),
                    ("Viscosity",   f"{mu_feed*1000:.4f} cP",     f"{mu_bw*1000:.4f} cP"),
                ], cols=("Property", "Feed", "Backwash"))

            if s_media:
                _sec("Media Configuration")
                _m_hdr = ("Media", "Support", "Depth (m)", "d10 (mm)", "CU",
                           "ε₀", "ρp,eff (kg/m³)", "ψ", "Vol (m³)")
                _m_rows = [_m_hdr] + [
                    (b["Type"],
                     "✓" if b.get("is_support") else "",
                     f"{b['Depth']:.3f}",
                     f"{b['d10']:.2f}",
                     f"{b['cu']:.2f}",
                     f"{b.get('epsilon0', '—'):.3f}",
                     f"{b['rho_p_eff']:.0f}",
                     f"{b.get('psi', '—')}",
                     f"{b['Vol']:.4f}",
                    ) for b in base]
                _tbl(_m_rows, cols=_m_hdr)

            if s_dp:
                _sec("Filtration Pressure Drop")
                _tbl([
                    ("Clean-bed ΔP (Ergun)",    f"{bw_dp['dp_clean_bar']:.4f} bar"),
                    ("50 % loaded ΔP (½ solid loading)", f"{(bw_dp['dp_clean_bar'] + bw_dp['dp_dirty_bar']) / 2:.4f} bar"),
                    ("Dirty ΔP (M_max)",        f"{bw_dp['dp_dirty_bar']:.4f} bar"),
                    ("BW trigger setpoint",     f"{dp_trigger_bar:.2f} bar"),
                    ("Superficial LV (N)",      f"{bw_dp['u_m_h']:.2f} m/h"),
                    ("Specific resistance α",
                     f"{filt_cycles['N']['alpha_used_m_kg']/1e9:.1f} × 10⁹ m/kg  "
                     f"({filt_cycles['N']['alpha_source']})"),
                ])

            if s_cycle:
                _sec("Filtration Cycle & BW Feasibility")
                _tbl([
                    ("Max solid loading",        f"{solid_loading:.2f} kg/m²"),
                    ("BW total duration",        f"{bw_total_min} min"),
                    ("BW sequence",
                     f"Drain {bw_s_drain}' · Air {bw_s_air}' · Air+W {bw_s_airw}' · "
                     f"HW {bw_s_hw}' · Settle {bw_s_settle}' · Fill {bw_s_fill}'"),
                ])
                doc.add_paragraph(
                    "Cycle duration (h) — N scenario, design temperature:")
                _cyc_n = cycle_matrix.get("N", {}).get(
                    f"Design ({feed_temp:.0f}°C)", {})
                if _cyc_n:
                    _ch = ("TSS (mg/L)", "Cycle (h)", "Avail. (%)",
                           "BW/day", "BW systems (plant-wide)")
                    _cr = [_ch]
                    for _tl, _tv in zip(_tss_labels, _tss_vals):
                        _tr = next((r for r in _cyc_n["tss_results"]
                                    if r["TSS (mg/L)"] == _tv), None)
                        _fk = feasibility_matrix.get("N", {}).get(
                            f"Design ({feed_temp:.0f}°C)", {}).get(_tl, {})
                        _cr.append((
                            f"{_tv:.0f}",
                            f"{_tr['Cycle duration (h)']:.1f}" if _tr else "—",
                            f"{_fk.get('avail_pct', 0):.1f}",
                            f"{_fk.get('bw_per_day', 0):.1f}",
                            str(_fk.get("bw_trains", "—")),
                        ))
                    _tbl(_cr, cols=_ch)

            # ── C. Mechanical & Structural ──────────────────────────────
            if s_vessel:
                _sec("Vessel Dimensions & ASME Thickness")
                _tbl([
                    ("Nominal ID",              f"{nominal_id:.3f} m"),
                    ("Real hydraulic ID",       f"{real_id:.4f} m"),
                    ("Outside diameter",        f"{mech['od_m']:.4f} m"),
                    ("Total length T/T",        f"{total_length:.3f} m"),
                    ("Cylindrical shell length", f"{cyl_len:.3f} m"),
                    ("End geometry",            end_geometry),
                    ("Material",                material_name),
                    ("Design pressure",         f"{design_pressure:.2f} bar g"),
                    ("Corrosion allowance",     f"{corrosion:.1f} mm"),
                    ("Shell t_required",        f"{mech['t_shell_min_mm']:.2f} mm"),
                    ("Shell t_design",          f"{mech['t_shell_design_mm']} mm"),
                    ("Head t_required",         f"{mech['t_head_min_mm']:.2f} mm"),
                    ("Head t_design",           f"{mech['t_head_design_mm']} mm"),
                    ("Shell joint efficiency E", f"{mech['shell_E']:.2f}"),
                    ("Head joint efficiency E",  f"{mech['head_E']:.2f}"),
                    ("Shell radiography",       f"{shell_radio}  (E={mech['shell_E']:.2f})"),
                    ("Head radiography",        f"{head_radio}  (E={mech['head_E']:.2f})"),
                ])

            if s_nzpl:
                _sec("Nozzle Plate Design")
                _tbl([
                    ("Plate height",            f"{nozzle_plate_h:.3f} m"),
                    ("Plate t_min / t_design",  f"{wt_np['t_min_mm']:.2f} / {wt_np['t_design_mm']} mm"),
                    ("Nozzle density",          f"{np_density:.0f} /m²"),
                    ("Nozzle bore",             f"{np_bore_dia:.1f} mm"),
                    ("Total nozzles",           f"{wt_np['n_bores']}"),
                    ("Plate area",              f"{wt_np['area_total_m2']:.4f} m²"),
                    ("Open area ratio",         f"{wt_np['open_ratio_pct']:.2f} %"),
                    ("Design ΔP (nozzle)",      f"{wt_np['q_dp_kpa']:.2f} kPa"),
                    ("Beam section",            f"{wt_np['beam_section']}  ({wt_np['n_beams']} beams)"),
                    ("Plate + beam weight",     f"{wt_np['weight_total_kg']:,.1f} kg"),
                ])

            if s_wt_empty:
                _sec("Empty Weight Breakdown")
                _tbl([
                    ("Cylindrical shell",       f"{wt_body['weight_shell_kg']:,.1f} kg"),
                    ("2 × Dish ends",           f"{wt_body['weight_two_heads_kg']:,.1f} kg"),
                    ("Nozzles",                 f"{w_noz:,.1f} kg"),
                    ("Nozzle plate assembly",   f"{wt_np['weight_total_kg']:,.1f} kg"),
                    (f"Supports ({wt_sup['support_type']})",
                                               f"{wt_sup['weight_all_supports_kg']:,.1f} kg"),
                    ("Strainer nozzles",        f"{wt_int['weight_strainers_kg']:,.1f} kg"),
                    ("Air scour header",        f"{wt_int['weight_air_header_kg']:,.1f} kg"),
                    ("Manholes",               f"{wt_int['weight_manholes_kg']:,.1f} kg"),
                    ("TOTAL EMPTY WEIGHT",
                     f"{w_total_rep:,.1f} kg  =  {w_total_rep/1000:.3f} t"),
                ])

            if s_wt_oper:
                _sec("Operating Weight & Support Loads")
                _tbl([
                    ("Empty vessel",            f"{wt_oper['w_empty_kg']:,.1f} kg"),
                    (f"Internal {lining_result['protection_type'].lower() or 'lining'}",
                                               f"{wt_oper['w_lining_kg']:,.1f} kg"),
                    ("Media — dry solid",       f"{wt_oper['w_media_kg']:,.1f} kg"),
                    ("Process water (full)",    f"{wt_oper['w_water_kg']:,.1f} kg"),
                    ("OPERATING WEIGHT",
                     f"{wt_oper['w_operating_kg']:,.1f} kg  =  {wt_oper['w_operating_t']:.3f} t"),
                    ("Number of supports",      str(wt_oper['n_supports'])),
                    ("Load per support",
                     f"{wt_oper['load_per_support_kg']:,.1f} kg  "
                     f"= {wt_oper['load_per_support_t']:.3f} t  "
                     f"= {wt_oper['load_per_support_kN']:.1f} kN"),
                ])

            if s_saddle:
                _sec("Saddle Design (Zick Method)")
                _sd = wt_saddle
                _tbl([
                    ("Saddle spacing factor α",   f"{_sd['alpha']:.2f}"),
                    ("Saddle 1 position",         f"{_sd['saddle_1_from_left_m']:.3f} m from left T/L"),
                    ("Saddle 2 position",         f"{_sd['saddle_2_from_left_m']:.3f} m from left T/L"),
                    ("Reaction per saddle",       f"{_sd['reaction_t']:.3f} t"),
                    ("Contact angle",             f"{_sd['contact_angle_deg']:.0f}°"),
                    ("Section selected",          str(_sd.get('section', '—'))),
                    ("Section capacity",          f"{_sd.get('capacity_t', '—')} t"),
                    ("Structural weight / saddle",f"{_sd.get('w_one_saddle_kg', 0):,.0f} kg"),
                    ("Status",                    "OVERSTRESSED" if _sd.get('overstressed') else "OK"),
                ])

            # ── D. Backwash System ──────────────────────────────────────
            if s_bw_hyd:
                _sec("Backwash Hydraulics & Collector Check")
                _tbl([
                    ("BW velocity (proposed)",   f"{bw_velocity:.1f} m/h"),
                    ("Max safe BW velocity",     f"{bw_col['max_safe_bw_m_h']:.1f} m/h"),
                    ("Air scour rate",           f"{air_scour_rate:.1f} m/h"),
                    ("Min. freeboard required",  f"{freeboard_mm:.0f} mm"),
                    ("Actual freeboard",
                     f"{bw_col['freeboard_m']:.3f} m  ({bw_col['freeboard_pct']:.1f}%)"),
                    ("Collector status",         bw_col["status"]),
                ])
                doc.add_paragraph("Bed expansion per layer:")
                _bw_rows = [("Layer", "Fluidised", "ε_fluid.", "Settled (m)",
                              "Expanded (m)", "u_mf (m/h)", "u_t (m/h)")]
                for _i, _bl in enumerate(bw_exp["layers"]):
                    _lname = layers[_i].get("Type", f"Layer {_i+1}") if _i < len(layers) else f"Layer {_i+1}"
                    _bw_rows.append((
                        _lname,
                        "Yes" if _bl.get("fluidised") else "No",
                        f"{_bl.get('eps_f', 0):.3f}",
                        f"{_bl.get('depth_settled_m', 0):.3f}",
                        f"{_bl.get('depth_expanded_m', 0):.3f}",
                        f"{_bl.get('u_mf_m_h', 0):.2f}",
                        f"{_bl.get('u_t_m_h', 0):.2f}",
                    ))
                _tbl(_bw_rows, cols=("Layer", "Fluidised", "ε_fluid.", "Settled (m)",
                                     "Expanded (m)", "u_mf (m/h)", "u_t (m/h)"))

            if s_bw_equip:
                _sec("BW Equipment Sizing")
                _bws = bw_sizing
                _tbl([
                    ("BW design flow",           f"{_bws['q_bw_design_m3h']:,.1f} m³/h"),
                    ("BW pump head",             f"{_bws['bw_head_mwc']:.2f} mWC"),
                    ("BW pump shaft power",      f"{_bws['p_pump_shaft_kw']:.1f} kW"),
                    ("BW pump motor power",      f"{_bws['p_pump_motor_kw']:.1f} kW"),
                    ("Air flow (design)",        f"{_bws['q_air_design_m3h']:,.1f} Am³/h"),
                    ("Blower back-pressure",     f"{_bws['P2_pa']/1e5:.3f} bar abs"),
                    ("Blower shaft power",       f"{_bws['p_blower_shaft_kw']:.1f} kW"),
                    ("Blower motor power",       f"{_bws['p_blower_motor_kw']:.1f} kW"),
                    ("BW water tank volume",     f"{_bws['v_tank_m3']:.1f} m³"),
                    ("BW systems (plant-wide)",  str(_bws.get('n_bw_systems', '—'))),
                ])

            # ── E. Internal Protection ──────────────────────────────────
            if s_lining and _has_lining:
                _sec(f"Internal Protection — {lining_result['protection_type']}")
                _tbl([
                    ("Protection type",          lining_result['protection_type']),
                    ("Total area protected",     f"{lining_result['a_total_m2']:.2f} m²"),
                    ("  — Shell (cyl + heads)",  f"{lining_result['a_shell_m2']:.2f} m²"),
                    ("  — Nozzle plate",         f"{lining_result['a_plate_m2']:.2f} m²"),
                    ("Lining / coating weight",  f"{lining_result['weight_kg']:,.1f} kg"),
                    ("Material cost",            f"USD {lining_result['material_cost_usd']:,.0f}"),
                    ("Application labour cost",  f"USD {lining_result['labor_cost_usd']:,.0f}"),
                    ("TOTAL COATING COST",       f"USD {lining_result['total_cost_usd']:,.0f}"),
                ] + [(k, v) for k, v in lining_result["detail"].items()
                     if k not in {"Material cost", "Labour cost", "Total cost"}])

            # ── F. Energy & Economics ───────────────────────────────────
            if s_hyd_prof:
                _sec("Hydraulic Head Profile")
                _hp_rows = [("Item", "Clean bed", "Dirty bed")]
                for k in hyd_prof["clean"]["items_bar"]:
                    _hp_rows.append((
                        k,
                        f"{hyd_prof['clean']['items_bar'][k]:.4f} bar"
                        f" / {hyd_prof['clean']['items_mwc'][k]:.2f} mWC",
                        f"{hyd_prof['dirty']['items_bar'][k]:.4f} bar"
                        f" / {hyd_prof['dirty']['items_mwc'][k]:.2f} mWC",
                    ))
                _hp_rows.append((
                    "TOTAL PUMP DUTY",
                    f"{hyd_prof['clean']['total_bar']:.4f} bar"
                    f" / {hyd_prof['clean']['total_mwc']:.2f} mWC",
                    f"{hyd_prof['dirty']['total_bar']:.4f} bar"
                    f" / {hyd_prof['dirty']['total_mwc']:.2f} mWC",
                ))
                _tbl(_hp_rows, cols=("Item", "Clean bed", "Dirty bed"))

            if s_energy:
                _sec("Energy Consumption & OPEX")
                _tbl([
                    ("Filtration pump power (dirty, per filter)",
                     f"{energy['p_filt_dirty_kw']:.1f} kW"),
                    ("BW pump power",            f"{energy['p_bw_kw']:.1f} kW"),
                    ("Air blower power (elec.)", f"{energy['p_blower_elec_kw']:.1f} kW"),
                    ("Annual filtration energy", f"{energy['e_filt_kwh_yr']/1e3:,.0f} MWh/yr"),
                    ("Annual BW energy (pump + blower)",
                     f"{(energy['e_bw_pump_kwh_yr']+energy['e_blower_kwh_yr'])/1e3:,.1f} MWh/yr"),
                    ("TOTAL annual energy",      f"{energy['e_total_kwh_yr']/1e3:,.0f} MWh/yr"),
                    ("Specific energy",          f"{energy['kwh_per_m3']:.4f} kWh/m³"),
                    ("Annual energy OPEX",       f"USD {energy['cost_usd_yr']:,.0f}/yr"),
                    ("Electricity tariff",       f"USD {elec_tariff:.3f}/kWh"),
                    ("Operating hours / year",   f"{op_hours_yr:,} h"),
                ])

            # ── G. Post-treatment ───────────────────────────────────────
            if s_cart:
                _sec("Cartridge Filter")
                _tbl([
                    ("Design flow",              f"{cart_result['design_flow_m3h']:,.1f} m³/h"),
                    ("Element material",         cart_result["element_material"]),
                    ("Element size",             cart_result["element_size"]),
                    ("Rating",                   f"{cart_result['rating_um']} µm absolute"),
                    ("Safety factor",            f"{cart_result['safety_factor']}×"),
                    ("Feed viscosity",           f"{cart_result['mu_cP']:.2f} cP"),
                    ("Elements required",        str(cart_result["n_elements"])),
                    ("Housings required",        str(cart_result["n_housings"])),
                    ("Flow / element",
                     f"{cart_result['actual_flow_m3h_element']:.3f} m³/h"
                     f"  ({cart_result['q_lpm_element']:.1f} lpm)"),
                    ("ΔP clean (BOL)",           f"{cart_result['dp_clean_bar']:.4f} bar"),
                    ("ΔP EOL",                   f"{cart_result['dp_eol_bar']:.4f} bar"),
                    ("DHC / element",            f"{cart_result['dhc_g_element']:.0f} g"),
                    ("Annual element cost",      f"USD {cart_result['annual_cost_usd']:,.0f}"),
                ])

            # ── A. Sign-off (always) ────────────────────────────────────
            doc.add_page_break()
            _h("Sign-off & Revision Record")
            _tbl([
                ("Prepared by",   engineer),
                ("Role",          "Process Expert — AQUASIGHT™"),
                ("Document No.",  f"{doc_number}  ·  Rev {revision}"),
                ("Project",       project_name),
                ("Client",        client or "—"),
                ("Date",          str(__import__("datetime").date.today())),
                ("Software",      "AQUASIGHT™ MMF Calculator"),
                ("Checked by",    ""),
                ("Approved by",   ""),
            ])
            doc.add_paragraph(
                "\nThis document was generated automatically. "
                "All calculations are the responsibility of the engineer of record."
            )

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.getvalue()

        # ── Download ───────────────────────────────────────────────────────
        _n_sections = sum([
            s_process, s_water, s_media, s_dp, s_cycle,
            s_vessel, s_nzpl, s_wt_empty, s_wt_oper, s_saddle,
            s_bw_hyd, s_bw_equip, s_lining and _has_lining,
            s_hyd_prof, s_energy, s_cart,
        ])
        st.caption(f"{_n_sections} section(s) selected + Identification & Sign-off (always included)")

        dl_col, _ = st.columns([2, 3])
        with dl_col:
            if _DOCX_OK:
                st.download_button(
                    label="⬇️  Download Word report (.docx)",
                    data=_build_docx(),
                    file_name=f"{doc_number}_Rev{revision}.docx",
                    mime=("application/vnd.openxmlformats-officedocument"
                          ".wordprocessingml.document"),
                )
            else:
                st.warning("Install python-docx:  `pip install python-docx`")

        st.divider()

        # ── Inline preview (mirrors section selections) ────────────────────
        st.markdown(f"**{project_name}** · {doc_number} · Rev {revision} · {engineer}")
        st.markdown("")

        if s_process:
            st.markdown("### B1 · Process Basis")
            st.markdown(f"""
| Parameter | Value |
|---|---|
| Total plant flow | {total_flow:,.0f} m³/h |
| Streams × filters / stream | {streams} × {n_filters} |
| Redundancy | N-{redundancy} per stream |
| Flow / filter (N) | {q_per_filter:.1f} m³/h |
| Filtration rate (N) | {q_per_filter/avg_area:.2f} m/h |
| Cross-sectional area | {avg_area:.3f} m² |
""")

        if s_water:
            st.markdown("### B2 · Water Properties")
            st.markdown(f"""
| Property | Feed | Backwash |
|---|---|---|
| Salinity | {feed_sal:.2f} ppt | {bw_sal:.2f} ppt |
| Temperature | {feed_temp:.1f} °C | {bw_temp:.1f} °C |
| Density | {rho_feed:.3f} kg/m³ | {rho_bw:.3f} kg/m³ |
| Viscosity | {mu_feed*1000:.4f} cP | {mu_bw*1000:.4f} cP |
""")

        if s_media:
            st.markdown("### B3 · Media Configuration")
            _med_hdr = "| Media | Sup. | Depth (m) | d10 (mm) | CU | ε₀ | ρp,eff | ψ |"
            _med_sep = "|---|---|---|---|---|---|---|---|"
            _med_rows = "\n".join(
                f"| {b['Type']} | {'✓' if b.get('is_support') else ''} "
                f"| {b['Depth']:.3f} | {b['d10']:.2f} | {b['cu']:.2f} "
                f"| {b.get('epsilon0', 0):.3f} | {b['rho_p_eff']:.0f} "
                f"| {b.get('psi', '—')} |"
                for b in base
            )
            st.markdown(f"{_med_hdr}\n{_med_sep}\n{_med_rows}")

        if s_dp:
            st.markdown("### B4 · Filtration ΔP")
            st.markdown(f"""
| | Value |
|---|---|
| Clean-bed ΔP (Ergun) | {bw_dp['dp_clean_bar']:.4f} bar |
| 50 % loaded ΔP | {(bw_dp['dp_clean_bar'] + bw_dp['dp_dirty_bar']) / 2:.4f} bar |
| Dirty ΔP (M_max) | {bw_dp['dp_dirty_bar']:.4f} bar |
| BW trigger setpoint | {dp_trigger_bar:.2f} bar |
| Superficial LV (N) | {bw_dp['u_m_h']:.2f} m/h |
""")

        if s_cycle:
            st.markdown("### B5 · Filtration Cycle & BW Feasibility")
            _cyc_n2 = cycle_matrix.get("N", {}).get(f"Design ({feed_temp:.0f}°C)", {})
            if _cyc_n2:
                _md_ch = "| TSS (mg/L) | Cycle (h) | Avail. (%) | BW/day | BW systems |"
                _md_cs = "|---|---|---|---|---|"
                _md_cr = []
                for _tl2, _tv2 in zip(_tss_labels, _tss_vals):
                    _tr3 = next((r for r in _cyc_n2["tss_results"]
                                 if r["TSS (mg/L)"] == _tv2), None)
                    _fk2 = feasibility_matrix.get("N", {}).get(
                        f"Design ({feed_temp:.0f}°C)", {}).get(_tl2, {})
                    _md_cr.append(
                        f"| {_tv2:.0f} "
                        f"| {_tr3['Cycle duration (h)']:.1f} " if _tr3 else "| — "
                        f"| {_fk2.get('avail_pct', 0):.1f} "
                        f"| {_fk2.get('bw_per_day', 0):.1f} "
                        f"| {_fk2.get('bw_trains', '—')} |"
                    )
                st.markdown(f"{_md_ch}\n{_md_cs}\n" + "\n".join(_md_cr))

        if s_vessel:
            st.markdown("### C1 · Vessel Dimensions & ASME Thickness")
            st.markdown(f"""
| Parameter | Value |
|---|---|
| Nominal ID / Real hyd. ID | {nominal_id:.3f} m / {real_id:.4f} m |
| Outside diameter | {mech['od_m']:.4f} m |
| Total length T/T | {total_length:.3f} m |
| End geometry | {end_geometry} |
| Material | {material_name} |
| Design pressure | {design_pressure:.2f} bar g |
| Shell t_req / t_design | {mech['t_shell_min_mm']:.2f} / {mech['t_shell_design_mm']} mm |
| Head t_req / t_design | {mech['t_head_min_mm']:.2f} / {mech['t_head_design_mm']} mm |
""")

        if s_wt_empty:
            st.markdown("### C3 · Empty Weight")
            st.markdown(f"""
| Component | Weight |
|---|---|
| Shell + 2 heads | {wt_body['weight_body_kg']:,.0f} kg |
| Nozzles | {w_noz:,.0f} kg |
| Nozzle plate assembly | {wt_np['weight_total_kg']:,.0f} kg |
| Supports | {wt_sup['weight_all_supports_kg']:,.0f} kg |
| Internals | {wt_int['weight_internals_kg']:,.0f} kg |
| **TOTAL EMPTY** | **{w_total_rep:,.0f} kg = {w_total_rep/1000:.3f} t** |
""")

        if s_wt_oper:
            st.markdown("### C4 · Operating Weight & Support Loads")
            st.markdown(f"""
| Component | Weight |
|---|---|
| Empty vessel | {wt_oper['w_empty_kg']:,.0f} kg |
| Lining / coating | {wt_oper['w_lining_kg']:,.0f} kg |
| Media (dry) | {wt_oper['w_media_kg']:,.0f} kg |
| Process water | {wt_oper['w_water_kg']:,.0f} kg |
| **OPERATING WEIGHT** | **{wt_oper['w_operating_kg']:,.0f} kg = {wt_oper['w_operating_t']:.3f} t** |
| Load / support ({wt_oper['n_supports']} supports) | {wt_oper['load_per_support_t']:.3f} t = {wt_oper['load_per_support_kN']:.1f} kN |
""")

        if s_saddle:
            _sd2 = wt_saddle
            st.markdown("### C5 · Saddle Design")
            st.markdown(f"""
| Parameter | Value |
|---|---|
| Spacing factor α | {_sd2['alpha']:.2f} |
| Saddle 1 / Saddle 2 position | {_sd2['saddle_1_from_left_m']:.3f} m / {_sd2['saddle_2_from_left_m']:.3f} m from left T/L |
| Reaction per saddle | {_sd2['reaction_t']:.3f} t |
| Section selected | {_sd2.get('section', '—')} — capacity {_sd2.get('capacity_t', '—')} t |
| Structural weight / saddle | {_sd2.get('w_one_saddle_kg', 0):,.0f} kg |
| Status | {"OVERSTRESSED" if _sd2.get('overstressed') else "OK"} |
""")

        if s_bw_hyd:
            st.markdown("### D1 · Backwash Hydraulics & Collector")
            st.markdown(f"""
| Parameter | Value |
|---|---|
| BW velocity (proposed / max safe) | {bw_velocity:.1f} / {bw_col['max_safe_bw_m_h']:.1f} m/h |
| Freeboard (min. req. / actual) | {freeboard_mm:.0f} mm / {bw_col['freeboard_m']*1000:.0f} mm |
| Collector status | {bw_col['status']} |
""")

        if s_bw_equip:
            _bws2 = bw_sizing
            st.markdown("### D2 · BW Equipment Sizing")
            st.markdown(f"""
| Equipment | Key Parameter |
|---|---|
| BW pump flow | {_bws2['q_bw_design_m3h']:,.1f} m³/h |
| BW pump head | {_bws2['bw_head_mwc']:.2f} mWC |
| BW pump motor | {_bws2['p_pump_motor_kw']:.1f} kW |
| Air blower flow | {_bws2['q_air_design_m3h']:,.1f} Am³/h |
| Blower back-pressure | {_bws2['P2_pa']/1e5:.3f} bar abs |
| Blower motor | {_bws2['p_blower_motor_kw']:.1f} kW |
| BW water tank | {_bws2['v_tank_m3']:.1f} m³ |
""")

        if s_lining and _has_lining:
            st.markdown(f"### E · Internal Protection — {lining_result['protection_type']}")
            st.markdown(f"""
| Parameter | Value |
|---|---|
| Total area | {lining_result['a_total_m2']:.2f} m² |
| Weight | {lining_result['weight_kg']:,.0f} kg |
| Material cost | USD {lining_result['material_cost_usd']:,.0f} |
| Labour cost | USD {lining_result['labor_cost_usd']:,.0f} |
| **Total coating cost** | **USD {lining_result['total_cost_usd']:,.0f}** |
""")

        if s_hyd_prof:
            st.markdown("### F1 · Hydraulic Head Profile")
            _hd_hdr = "| Item | Clean bed | Dirty bed |"
            _hd_sep = "|---|---|---|"
            _hd_rows = "\n".join(
                f"| {k} | {hyd_prof['clean']['items_bar'][k]:.4f} bar / "
                f"{hyd_prof['clean']['items_mwc'][k]:.2f} mWC | "
                f"{hyd_prof['dirty']['items_bar'][k]:.4f} bar / "
                f"{hyd_prof['dirty']['items_mwc'][k]:.2f} mWC |"
                for k in hyd_prof["clean"]["items_bar"]
            )
            _hd_tot = (f"| **Total** | **{hyd_prof['clean']['total_bar']:.4f} bar / "
                       f"{hyd_prof['clean']['total_mwc']:.2f} mWC** | "
                       f"**{hyd_prof['dirty']['total_bar']:.4f} bar / "
                       f"{hyd_prof['dirty']['total_mwc']:.2f} mWC** |")
            st.markdown(f"{_hd_hdr}\n{_hd_sep}\n{_hd_rows}\n{_hd_tot}")

        if s_energy:
            st.markdown("### F2 · Energy & OPEX")
            st.markdown(f"""
| KPI | Value |
|---|---|
| Filtration pump power (dirty, per filter) | {energy['p_filt_dirty_kw']:.1f} kW |
| BW pump power | {energy['p_bw_kw']:.1f} kW |
| Air blower power (elec.) | {energy['p_blower_elec_kw']:.1f} kW |
| Annual total energy | {energy['e_total_kwh_yr']/1e3:,.0f} MWh/yr |
| Specific energy | {energy['kwh_per_m3']:.4f} kWh/m³ |
| Annual energy OPEX | USD {energy['cost_usd_yr']:,.0f}/yr |
""")

        if s_cart:
            st.markdown("### G · Cartridge Filter")
            st.markdown(f"""
| Parameter | Value |
|---|---|
| Design flow | {cart_result['design_flow_m3h']:,.1f} m³/h |
| Element | {cart_result['element_size']} · {cart_result['rating_um']} µm · {cart_result['element_material']} |
| Elements / Housings | {cart_result['n_elements']} / {cart_result['n_housings']} |
| Flow / element | {cart_result['actual_flow_m3h_element']:.3f} m³/h ({cart_result['q_lpm_element']:.1f} lpm) |
| ΔP BOL / EOL | {cart_result['dp_clean_bar']:.4f} / {cart_result['dp_eol_bar']:.4f} bar |
| Annual element cost | USD {cart_result['annual_cost_usd']:,.0f} |
""")

        st.divider()
        col_sign, _ = st.columns([1, 2])
        with col_sign:
            st.info(f"**{engineer}**  \nProcess Expert — AQUASIGHT™  \n\n"
                    f"{doc_number} · Rev {revision}")

    # ─────────────────────────────────────────────────────────────────────
    # TAB 13 · ECONOMICS
    # ─────────────────────────────────────────────────────────────────────
    with tab_econ:
        st.subheader("Economics — CAPEX · OPEX · Carbon · Benchmarks")

        # ── Key headline metrics ──────────────────────────────────────────
        em1, em2, em3, em4 = st.columns(4)
        em1.metric("Total CAPEX",
                   f"USD {econ_capex['total_capex_usd']:,.0f}",
                   delta=f"{econ_bench['capex_per_m3d']:.1f} USD/m³/d  {econ_bench['capex_status']}",
                   delta_color="off")
        em2.metric("Annual OPEX",
                   f"USD {econ_opex['total_opex_usd_yr']:,.0f}/yr",
                   delta=f"{econ_bench['opex_per_m3']:.4f} USD/m³  {econ_bench['opex_status']}",
                   delta_color="off")
        em3.metric("LCOW",
                   f"{econ_bench['lcow']:.4f} USD/m³",
                   delta=econ_bench["lcow_status"],
                   delta_color="off")
        em4.metric("CO₂ operational",
                   f"{econ_carbon['co2_per_m3_operational']:.4f} kgCO₂/m³",
                   delta=econ_bench["carbon_status"],
                   delta_color="off")

        # ── CAPEX breakdown ───────────────────────────────────────────────
        with st.expander("1 · CAPEX breakdown", expanded=True):
            _capex_items = {
                "Steel (structure)":      econ_capex["steel_cost_usd"],
                "Erection":               econ_capex["erection_usd"],
                "Piping":                 econ_capex["piping_usd"],
                "Instrumentation":        econ_capex["instrumentation_usd"],
                "Civil works":            econ_capex["civil_usd"],
                "Engineering":            econ_capex["engineering_usd"],
                "Contingency":            econ_capex["contingency_usd"],
            }
            c_left, c_right = st.columns([1, 1])
            with c_left:
                st.table(pd.DataFrame([
                    [k, f"USD {v:,.0f}", f"{v/max(econ_capex['total_capex_usd'],1)*100:.1f} %"]
                    for k, v in _capex_items.items()
                ] + [["**TOTAL**",
                       f"**USD {econ_capex['total_capex_usd']:,.0f}**", "**100 %**"]],
                    columns=["Item", "Cost (USD)", "Share"]))
                st.caption(
                    f"{_n_total_vessels} vessels · steel {steel_cost_usd_kg:.2f} USD/kg · "
                    f"engineering {engineering_pct:.0f} % · contingency {contingency_pct:.0f} %"
                )
            with c_right:
                if _PLOTLY_OK:
                    _fig_cap = _go.Figure(_go.Pie(
                        labels=list(_capex_items.keys()),
                        values=list(_capex_items.values()),
                        hole=0.35,
                        textinfo="label+percent",
                        textfont_size=11,
                    ))
                    _fig_cap.update_layout(
                        title="CAPEX split",
                        showlegend=False,
                        margin=dict(t=40, b=10, l=10, r=10),
                        height=340,
                    )
                    st.plotly_chart(_fig_cap, use_container_width=True)
                else:
                    st.info("Install plotly for pie charts.")

        # ── OPEX breakdown ────────────────────────────────────────────────
        with st.expander("2 · Annual OPEX breakdown", expanded=True):
            _opex_items = {
                "Energy":             econ_opex["energy_cost_usd_yr"],
                "Media replacement":  econ_opex["media_cost_usd_yr"],
                "Nozzle replacement": econ_opex["nozzle_cost_usd_yr"],
                "Labour":             econ_opex["labour_cost_usd_yr"],
                "Chemicals":          econ_opex["chemical_cost_usd_yr"],
            }
            o_left, o_right = st.columns([1, 1])
            with o_left:
                st.table(pd.DataFrame([
                    [k, f"USD {v:,.0f}/yr", f"{v/max(econ_opex['total_opex_usd_yr'],1)*100:.1f} %"]
                    for k, v in _opex_items.items()
                ] + [["**TOTAL**",
                       f"**USD {econ_opex['total_opex_usd_yr']:,.0f}/yr**", "**100 %**"]],
                    columns=["Item", "Cost (USD/yr)", "Share"]))
                st.caption(
                    f"Specific OPEX: **{econ_opex['opex_per_m3_usd']:.4f} USD/m³**  ·  "
                    f"Annual flow: {econ_opex['annual_flow_m3']/1e6:.2f} Mm³/yr  ·  "
                    f"Media interval: {media_replace_years:.0f} yr"
                )
            with o_right:
                if _PLOTLY_OK:
                    _fig_op = _go.Figure(_go.Pie(
                        labels=list(_opex_items.keys()),
                        values=list(_opex_items.values()),
                        hole=0.35,
                        textinfo="label+percent",
                        textfont_size=11,
                    ))
                    _fig_op.update_layout(
                        title="OPEX split",
                        showlegend=False,
                        margin=dict(t=40, b=10, l=10, r=10),
                        height=340,
                    )
                    st.plotly_chart(_fig_op, use_container_width=True)

        # ── Carbon footprint ──────────────────────────────────────────────
        with st.expander("3 · Carbon footprint", expanded=True):
            cf1, cf2, cf3, cf4 = st.columns(4)
            cf1.metric("Operational CO₂/yr", f"{econ_carbon['co2_operational_kg_yr']/1000:,.1f} t/yr")
            cf2.metric("Construction CO₂",   f"{econ_carbon['co2_construction_kg']/1000:,.1f} t")
            cf3.metric("Lifecycle CO₂",      f"{econ_carbon['co2_lifecycle_kg']/1000:,.1f} t",
                       delta=f"over {econ_carbon['design_life_years']} yr", delta_color="off")
            cf4.metric("Specific operational", f"{econ_carbon['co2_per_m3_operational']:.4f} kgCO₂/m³",
                       delta=econ_bench["carbon_status"], delta_color="off")

            st.table(pd.DataFrame([
                ["Operational CO₂ / year",
                 f"{econ_carbon['co2_operational_kg_yr']:,.0f} kg/yr",
                 f"Grid: {grid_intensity:.3f} kgCO₂/kWh"],
                ["Construction — steel",
                 f"{econ_carbon['co2_steel_kg']:,.0f} kg",
                 f"{steel_carbon_kg:.2f} kgCO₂/kg steel"],
                ["Construction — media",
                 f"{econ_carbon['co2_media_kg']:,.0f} kg",
                 "Weighted by mass"],
                ["Construction — concrete",
                 f"{econ_carbon['co2_concrete_kg']:,.0f} kg",
                 f"{concrete_carbon_kg:.2f} kgCO₂/kg"],
                ["Lifecycle total",
                 f"{econ_carbon['co2_lifecycle_kg']:,.0f} kg",
                 f"= {econ_carbon['co2_lifecycle_kg']/1000:.1f} t over {econ_carbon['design_life_years']} yr"],
                ["Specific — operational",
                 f"{econ_carbon['co2_per_m3_operational']:.4f} kgCO₂/m³",
                 econ_bench["carbon_status"]],
                ["Specific — lifecycle",
                 f"{econ_carbon['co2_per_m3_lifecycle']:.4f} kgCO₂/m³",
                 "Incl. construction, amortised"],
            ], columns=["Item", "Value", "Basis"]))

        # ── Global benchmark comparison ───────────────────────────────────
        with st.expander("4 · Global benchmark comparison", expanded=True):
            st.caption(
                "Benchmarks: horizontal MMF for SWRO / brackish pre-treatment "
                "(Middle East / Mediterranean, 2024 basis). "
                "🟢 = within range · 🟡 = borderline · 🔴 = outside range."
            )
            st.table(pd.DataFrame([
                ["CAPEX",
                 f"{econ_bench['capex_per_m3d']:.2f} USD/m³/d",
                 econ_bench["capex_benchmark"],
                 econ_bench["capex_status"]],
                ["OPEX",
                 f"{econ_bench['opex_per_m3']:.4f} USD/m³",
                 econ_bench["opex_benchmark"],
                 econ_bench["opex_status"]],
                ["Operational carbon",
                 f"{econ_bench['co2_per_m3']:.4f} kgCO₂/m³",
                 econ_bench["carbon_benchmark"],
                 econ_bench["carbon_status"]],
                ["LCOW",
                 f"{econ_bench['lcow']:.4f} USD/m³",
                 econ_bench["lcow_benchmark"],
                 econ_bench["lcow_status"]],
            ], columns=["Metric", "Project", "Benchmark range", "Status"]))

            st.caption(
                f"Daily capacity: {econ_bench['daily_flow_m3d']:,.0f} m³/d  ·  "
                f"Annual flow: {econ_bench['annual_flow_m3']/1e6:.2f} Mm³/yr  ·  "
                f"LCOW basis: CRF = 8 % (≈ 12-yr payback at 5 %)."
            )