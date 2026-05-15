"""ui/tab_report.py — Report tab for AQUASIGHT™ MMF."""
import streamlit as st
from ui.helpers import fmt, ulbl, dv, fmt_bar_mwc
from engine.design_basis_report import (
    collector_summary_rows,
    design_basis_meta_rows,
    plain_text,
    traceability_table_rows,
)
from engine.pdf_report import build_pdf, PDF_OK as _PDF_OK
from ui.scroll_markers import inject_anchor

try:
    from docx import Document as _DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


def render_tab_report(inputs: dict, computed: dict):
    inject_anchor("mmf-anchor-main-report")
    wt_body       = computed["wt_body"]
    w_noz         = computed["w_noz"]
    wt_np         = computed["wt_np"]
    wt_sup        = computed["wt_sup"]
    wt_int        = computed["wt_int"]
    lining_result = computed["lining_result"]
    q_per_filter  = computed["q_per_filter"]
    avg_area      = computed["avg_area"]
    rho_feed      = computed["rho_feed"]
    mu_feed       = computed["mu_feed"]
    rho_bw        = computed["rho_bw"]
    mu_bw         = computed["mu_bw"]
    base          = computed["base"]
    bw_dp         = computed["bw_dp"]
    filt_cycles   = computed["filt_cycles"]
    bw_col        = computed["bw_col"]
    bw_sizing     = computed["bw_sizing"]
    hyd_prof      = computed["hyd_prof"]
    energy        = computed["energy"]
    econ_financial = computed.get("econ_financial") or {}
    cart_result   = computed["cart_result"]
    nominal_id    = computed["nominal_id"]
    real_id       = computed["real_id"]
    mech          = computed["mech"]
    total_length  = computed["total_length"]
    cyl_len       = computed["cyl_len"]
    end_geometry  = computed["end_geometry"]
    material_name = computed["material_name"]
    wt_oper       = computed["wt_oper"]
    wt_saddle     = computed["wt_saddle"]
    w_total       = computed["w_total"]

    project_name    = inputs["project_name"]
    doc_number      = inputs["doc_number"]
    revision        = inputs["revision"]
    client          = inputs["client"]
    engineer        = inputs["engineer"]
    total_flow      = inputs["total_flow"]
    streams         = inputs["streams"]
    n_filters       = inputs["n_filters"]
    redundancy      = inputs["redundancy"]
    hydraulic_assist = int(inputs.get("hydraulic_assist", 0))
    feed_sal        = inputs["feed_sal"]
    bw_sal          = inputs["bw_sal"]
    feed_temp       = inputs["feed_temp"]
    bw_temp         = inputs["bw_temp"]
    dp_trigger_bar  = inputs["dp_trigger_bar"]
    solid_loading   = inputs["solid_loading"]
    bw_total_min    = inputs["bw_total_min"]
    bw_s_drain      = inputs["bw_s_drain"]
    bw_s_air        = inputs["bw_s_air"]
    bw_s_airw       = inputs["bw_s_airw"]
    bw_s_hw         = inputs["bw_s_hw"]
    bw_s_settle     = inputs["bw_s_settle"]
    bw_s_fill       = inputs["bw_s_fill"]
    design_pressure = inputs["design_pressure"]
    corrosion       = inputs["corrosion"]
    shell_radio     = inputs["shell_radio"]
    head_radio      = inputs["head_radio"]
    nozzle_plate_h  = inputs["nozzle_plate_h"]
    np_density      = inputs["np_density"]
    np_bore_dia     = inputs["np_bore_dia"]
    bw_velocity     = inputs["bw_velocity"]
    freeboard_mm    = inputs["freeboard_mm"]
    air_scour_rate  = inputs["air_scour_rate"]

    w_total_rep = (wt_body["weight_body_kg"] + w_noz
                   + wt_np["weight_total_kg"]
                   + wt_sup["weight_all_supports_kg"]
                   + wt_int["weight_internals_kg"])
    _has_lining = lining_result["protection_type"] != "None"

    st.caption(
        "Project **Save / Save as / Load** and **New project** sit **under the title**, above **Quick jump**, "
        "in the results stack (full width when the input column is hidden)."
    )
    st.divider()

    st.markdown("### Report builder")
    st.caption(
        "Select the sections to include. "
        "Identification (cover, project info, sign-off) is always present."
    )

    sel_c1, sel_c2, sel_c3, sel_c4 = st.columns(4)
    with sel_c1:
        st.markdown("**B · Process Design**")
        s_process = st.checkbox("Process basis",                   value=True, key="rs_proc")
        s_water   = st.checkbox("Water properties",                value=True, key="rs_water")
        s_media   = st.checkbox("Media configuration",             value=True, key="rs_media")
        s_dp      = st.checkbox("Filtration ΔP",                  value=True, key="rs_dp")
        s_cycle   = st.checkbox("Filtration cycle & BW feasibility",
                                 value=True, key="rs_cycle")
    with sel_c2:
        st.markdown("**C · Mechanical & Structural**")
        s_vessel   = st.checkbox("Vessel dimensions & ASME thickness",
                                  value=True, key="rs_vessel")
        s_nzpl     = st.checkbox("Nozzle plate design",    value=True, key="rs_nzpl")
        s_wt_empty = st.checkbox("Empty weight breakdown", value=True, key="rs_wtempty")
        s_wt_oper  = st.checkbox("Operating weight & support loads",
                                  value=True, key="rs_wtoper")
        s_saddle   = st.checkbox("Saddle design",          value=True, key="rs_saddle")
    with sel_c3:
        st.markdown("**D · Backwash System**")
        s_bw_hyd   = st.checkbox("BW hydraulics & collector",
                                  value=True, key="rs_bwhyd")
        s_bw_equip = st.checkbox("BW equipment sizing\n(pump / blower / tank)",
                                  value=True, key="rs_bweq")
        st.markdown("**E · Internal Protection**")
        _lining_label = (f"Lining / coating  ({lining_result['protection_type']})"
                         if _has_lining else "Lining / coating  *(None — skipped)*")
        s_lining = st.checkbox(_lining_label, value=_has_lining,
                               disabled=not _has_lining, key="rs_lining")
    with sel_c4:
        st.markdown("**F · Energy & Economics**")
        s_hyd_prof = st.checkbox("Hydraulic head profile", value=True, key="rs_hyd")
        s_energy   = st.checkbox("Energy & OPEX",          value=True, key="rs_energy")
        s_financial = st.checkbox("Lifecycle financial (NPV / cash flow)", value=True, key="rs_fin")
        s_design_basis = st.checkbox(
            "Design basis & traceability",
            value=True,
            key="rs_design_basis",
            help="Assumptions, limits, output traceability, collector summary (also in JSON export).",
        )
        st.markdown("**H · Post-treatment**")
        s_cart     = st.checkbox("Cartridge filter",       value=True, key="rs_cart")

    st.divider()

    _basis = computed.get("design_basis") or {}
    with st.expander("Design basis & traceability (JSON export)", expanded=False):
        st.caption(
            f"Assumptions, limits, and output traceability for enterprise review. "
            f"Reference: **{_basis.get('doc_reference', 'AQUASIGHT_MMF_MODELS_AND_STRATEGIES.md')}**"
        )
        if _basis.get("assumptions"):
            st.markdown("**Assumptions**")
            for _a in _basis["assumptions"]:
                st.markdown(f"- {_a}")
        if _basis.get("limits_and_criteria"):
            st.markdown("**Limits & criteria**")
            for _lim in _basis["limits_and_criteria"]:
                st.markdown(f"- {_lim}")
        _col = _basis.get("collector") or {}
        if _col:
            _dist = _col.get("distribution") or {}
            st.markdown(
                f"**Collector distribution:** {_dist.get('iterations', '—')} iterations · "
                f"{'converged' if _dist.get('converged') else 'not converged'} · "
                f"residual {_dist.get('residual_rel', '—')}"
            )
        import json

        _basis_json = json.dumps(
            {"inputs": inputs, "design_basis": _basis},
            indent=2,
            default=str,
        )
        st.download_button(
            "⬇️ Download design basis JSON",
            data=_basis_json,
            file_name=f"{doc_number}_design_basis.json",
            mime="application/json",
            key="dl_design_basis_json",
        )

    def _build_docx() -> bytes:
        import io as _io
        doc = _DocxDocument()

        def _tbl(rows_data, cols=("Parameter", "Value")):
            tbl = doc.add_table(rows=len(rows_data), cols=len(cols))
            tbl.style = "Table Grid"
            for i, row_vals in enumerate(rows_data):
                for j, v in enumerate(row_vals):
                    tbl.rows[i].cells[j].text = str(v)
            doc.add_paragraph("")
            return tbl

        def _h(text, lvl=2):
            doc.add_heading(text, lvl)

        t = doc.add_heading("AQUASIGHT™  Horizontal Multi-Media Filter", 0)
        t.alignment = _WD_ALIGN.CENTER
        t2 = doc.add_heading("Calculation Report", 1)
        t2.alignment = _WD_ALIGN.CENTER
        doc.add_paragraph("")
        _h("Project Information")
        _tbl([
            ("Project",       project_name),
            ("Document No.",  f"{doc_number}  ·  Rev {revision}"),
            ("Client",        client or "—"),
            ("Prepared by",   engineer),
            ("Date",          str(__import__("datetime").date.today())),
        ])

        _sn = [1]
        def _sec(title):
            _sn[0] += 1
            _h(f"{_sn[0]}. {title}")

        if s_process:
            _sec("Process Basis")
            _tbl([
                ("Total plant flow",     fmt(total_flow, "flow_m3h", 0)),
                ("Streams",              str(streams)),
                ("Total physical number of filters / stream", str(n_filters)),
                ("Hydraulic assist",     f"{hydraulic_assist} spare(s) / stream (design N = installed − spare)"),
                ("Outage depth modelled", f"N-{redundancy} per stream"),
                ("Flow / filter (N)",    fmt(q_per_filter, "flow_m3h", 1)),
                ("Filtration rate (N)",  fmt(q_per_filter / avg_area, "velocity_m_h", 2)),
                ("Cross-sectional area", fmt(avg_area, "area_m2", 3)),
            ])

        if s_water:
            _sec("Water Properties")
            _tbl([
                ("",            "Feed",                    "Backwash"),
                ("Salinity",    f"{feed_sal:.2f} ppt",    f"{bw_sal:.2f} ppt"),
                ("Temperature", fmt(feed_temp, "temperature_c", 1), fmt(bw_temp, "temperature_c", 1)),
                ("Density",     fmt(rho_feed, "density_kg_m3", 3), fmt(rho_bw, "density_kg_m3", 3)),
                ("Viscosity",   fmt(mu_feed * 1000.0, "viscosity_cp", 4),
                 fmt(mu_bw * 1000.0, "viscosity_cp", 4)),
            ], cols=("Property", "Feed", "Backwash"))

        if s_media:
            _sec("Media Configuration")
            _m_hdr = (
                "Media", "Support",
                f"Depth ({ulbl('length_m')})", f"d10 ({ulbl('length_mm')})", "CU", "ε₀",
                f"ρp,eff ({ulbl('density_kg_m3')})", "ψ",
                f"Vol ({ulbl('volume_m3')})",
            )
            _m_rows = [_m_hdr] + [
                (b["Type"],
                 "✓" if b.get("is_support") else "",
                 f"{dv(b['Depth'], 'length_m'):.3f}", fmt(float(b["d10"]), "length_mm", 2), f"{b['cu']:.2f}",
                 f"{b.get('epsilon0', 0):.3f}", f"{dv(b['rho_p_eff'], 'density_kg_m3'):.0f}",
                 f"{b.get('psi', '—')}", f"{dv(b['Vol'], 'volume_m3'):.4f}",
                ) for b in base]
            _tbl(_m_rows, cols=_m_hdr)

        if s_dp:
            _sec("Filtration Pressure Drop")
            _tbl([
                ("Clean-bed ΔP (Ergun)",  fmt(bw_dp["dp_clean_bar"], "pressure_bar", 4)),
                ("50 % loaded ΔP",
                 fmt((bw_dp["dp_clean_bar"] + bw_dp["dp_dirty_bar"]) / 2.0, "pressure_bar", 4)),
                ("Dirty ΔP (M_max)",      fmt(bw_dp["dp_dirty_bar"], "pressure_bar", 4)),
                ("BW trigger setpoint",   fmt(dp_trigger_bar, "pressure_bar", 2)),
                ("Superficial LV (N)",    fmt(bw_dp["u_m_h"], "velocity_m_h", 2)),
                ("Specific resistance α",
                 f"{filt_cycles['N']['alpha_used_m_kg']/1e9:.1f} {ulbl('alpha_m_kg')} "
                 f"({filt_cycles['N']['alpha_source']})"),
            ])

        if s_cycle:
            _sec("Filtration Cycle & BW Feasibility")
            _tbl([
                ("Max solid loading", fmt(solid_loading, "loading_kg_m2", 2)),
                ("BW total duration", f"{bw_total_min} min"),
                ("BW sequence",
                 f"Drain {bw_s_drain}' · Air {bw_s_air}' · Air+W {bw_s_airw}' · "
                 f"HW {bw_s_hw}' · Settle {bw_s_settle}' · Fill {bw_s_fill}'"),
            ])

        if s_vessel:
            _sec("Vessel Dimensions & ASME Thickness")
            _tbl([
                ("Nominal ID",               fmt(nominal_id, "length_m", 3)),
                ("Real hydraulic ID",        fmt(real_id, "length_m", 4)),
                ("Outside diameter",         fmt(mech["od_m"], "length_m", 4)),
                ("Total length T/T",         fmt(total_length, "length_m", 3)),
                ("Cylindrical shell length", fmt(cyl_len, "length_m", 3)),
                ("End geometry",             end_geometry),
                ("Material",                 material_name),
                ("Design pressure",          fmt(design_pressure, "pressure_bar", 2) + " g"),
                ("Corrosion allowance",      fmt(corrosion, "length_mm", 1)),
                ("Shell t_required",         fmt(float(mech["t_shell_min_mm"]), "length_mm", 2)),
                ("Shell t_design",           fmt(float(mech["t_shell_design_mm"]), "length_mm", 2)),
                ("Head t_required",          fmt(float(mech["t_head_min_mm"]), "length_mm", 2)),
                ("Head t_design",            fmt(float(mech["t_head_design_mm"]), "length_mm", 2)),
                ("Shell radiography",        f"{shell_radio}  (E={mech['shell_E']:.2f})"),
                ("Head radiography",         f"{head_radio}  (E={mech['head_E']:.2f})"),
            ])

        if s_nzpl:
            _sec("Nozzle Plate Design")
            _tbl([
                ("Plate height",
                 fmt(nozzle_plate_h, "length_m", 3)),
                ("Plate t_min / t_design",
                 f"{dv(float(wt_np['t_min_mm']), 'length_mm'):.2f} / "
                 f"{dv(float(wt_np['t_design_mm']), 'length_mm'):.2f} {ulbl('length_mm')}"),
                ("Nozzle density",         f"{dv(np_density, 'quantity_per_m2'):.0f} {ulbl('quantity_per_m2')}"),
                ("Nozzle bore",            fmt(np_bore_dia, "length_mm", 1)),
                ("Total nozzles",          f"{wt_np['n_bores']}"),
                ("Plate area",             fmt(wt_np["area_total_m2"], "area_m2", 4)),
                ("Open area ratio",        f"{wt_np['open_ratio_pct']:.2f} %"),
                ("Design ΔP (nozzle)",     fmt(wt_np["q_dp_kpa"], "pressure_kpa", 2)),
                ("Beam section",
                 f"{wt_np['beam_section']}  ({wt_np['n_beams']} beams)"),
                ("Plate + beam weight",    fmt(wt_np["weight_total_kg"], "mass_kg", 1)),
            ])

        if s_wt_empty:
            _sec("Empty Weight Breakdown")
            _tbl([
                ("Cylindrical shell",     fmt(wt_body["weight_shell_kg"], "mass_kg", 1)),
                ("2 × Dish ends",         fmt(wt_body["weight_two_heads_kg"], "mass_kg", 1)),
                ("Nozzles",               fmt(w_noz, "mass_kg", 1)),
                ("Nozzle plate assembly", fmt(wt_np["weight_total_kg"], "mass_kg", 1)),
                (f"Supports ({wt_sup['support_type']})",
                 fmt(wt_sup["weight_all_supports_kg"], "mass_kg", 1)),
                ("Strainer nozzles",      fmt(wt_int["weight_strainers_kg"], "mass_kg", 1)),
                ("Air scour header",      fmt(wt_int["weight_air_header_kg"], "mass_kg", 1)),
                ("Manholes",              fmt(wt_int["weight_manholes_kg"], "mass_kg", 1)),
                ("TOTAL EMPTY WEIGHT",
                 f"{fmt(w_total_rep, 'mass_kg', 1)}  =  {fmt(w_total_rep / 1000.0, 'mass_t', 3)}"),
            ])

        if s_wt_oper:
            _sec("Operating Weight & Support Loads")
            _tbl([
                ("Empty vessel",
                 fmt(wt_oper["w_empty_kg"], "mass_kg", 1)),
                (f"Internal {lining_result['protection_type'].lower() or 'lining'}",
                 fmt(wt_oper["w_lining_kg"], "mass_kg", 1)),
                ("Media — dry solid",    fmt(wt_oper["w_media_kg"], "mass_kg", 1)),
                ("Process water (full)", fmt(wt_oper["w_water_kg"], "mass_kg", 1)),
                ("OPERATING WEIGHT",
                 f"{fmt(wt_oper['w_operating_kg'], 'mass_kg', 1)}  =  "
                 f"{fmt(wt_oper['w_operating_t'], 'mass_t', 3)}"),
                ("Number of supports",   str(wt_oper["n_supports"])),
                ("Load per support",
                 f"{fmt(wt_oper['load_per_support_kg'], 'mass_kg', 1)}  "
                 f"= {fmt(wt_oper['load_per_support_t'], 'mass_t', 3)}  "
                 f"= {fmt(wt_oper['load_per_support_kN'], 'force_kn', 1)}"),
            ])

        if s_saddle:
            _sec("Saddle Design (Zick Method)")
            _sd = wt_saddle
            _pos_sd = _sd.get("saddle_positions_m") or [
                _sd.get("saddle_1_from_left_m"),
                _sd.get("saddle_2_from_left_m"),
            ]
            _pos_sd = [p for p in _pos_sd if p is not None]
            _pos_txt = ", ".join(
                fmt(float(p), "length_m", 3) for p in _pos_sd
            ) or "—"
            _sp_sd = _sd.get("saddle_spacings_m") or []
            _sp_txt = (
                " / ".join(
                    f"{dv(float(s), 'length_m'):.2f} {ulbl('length_m')}"
                    for s in _sp_sd
                )
                if _sp_sd else "—"
            )
            _tbl([
                ("Saddle spacing factor α",
                 f"{_sd['alpha']:.2f}"),
                ("Saddle positions (from left T/L)",
                 _pos_txt),
                ("Centre-to-centre spacings",
                 _sp_txt),
                ("Reaction per saddle",       fmt(_sd["reaction_t"], "mass_t", 3)),
                ("Section selected",          str(_sd.get("section", "—"))),
                ("Section capacity",          (
                    fmt(float(_sd["capacity_t"]), "mass_t", 3)
                    if isinstance(_sd.get("capacity_t"), (int, float))
                    else str(_sd.get("capacity_t", "—")))),
                ("Structural weight / saddle",
                 fmt(_sd.get("w_one_saddle_kg", 0), "mass_kg", 0)),
                ("Status",
                 "OVERSTRESSED" if _sd.get("overstressed") else "OK"),
            ])

        if s_bw_hyd:
            _sec("Backwash Hydraulics & Collector Check")
            _tbl([
                ("BW velocity (proposed)",  fmt(bw_velocity, "velocity_m_h", 1)),
                ("Max safe BW velocity",    fmt(bw_col["max_safe_bw_m_h"], "velocity_m_h", 1)),
                ("Air scour rate",          fmt(air_scour_rate, "velocity_m_h", 1)),
                ("Actual freeboard",
                 f"{fmt(bw_col['freeboard_m'], 'length_m', 3)}  ({bw_col['freeboard_pct']:.1f}%)"),
                ("Collector status",        bw_col["status"]),
            ])

        if s_bw_equip:
            _sec("BW Equipment Sizing")
            _bws = bw_sizing
            _tbl([
                ("BW design flow",       fmt(_bws["q_bw_design_m3h"], "flow_m3h", 1)),
                ("BW pump head",         fmt(_bws["bw_head_mwc"], "pressure_mwc", 2)),
                ("BW pump shaft power",  fmt(_bws["p_pump_shaft_kw"], "power_kw", 1)),
                ("BW pump motor power",  fmt(_bws["p_pump_motor_kw"], "power_kw", 1)),
                ("Air flow (design)",    fmt(_bws["q_air_design_nm3h"], "air_flow_nm3h", 1)),
                ("Blower back-pressure", fmt(_bws["P2_pa"] / 1e5, "pressure_bar", 3) + " abs"),
                ("Blower shaft power",   fmt(_bws["p_blower_shaft_kw"], "power_kw", 1)),
                ("Blower motor power",   fmt(_bws["p_blower_motor_kw"], "power_kw", 1)),
                ("BW water tank volume", fmt(_bws["v_tank_m3"], "volume_m3", 1)),
            ])

        if s_lining and _has_lining:
            _sec(f"Internal Protection — {lining_result['protection_type']}")
            _tbl([
                ("Protection type",
                 lining_result["protection_type"]),
                ("Total area protected",
                 fmt(lining_result["a_total_m2"], "area_m2", 2)),
                ("Lining / coating weight",
                 fmt(lining_result["weight_kg"], "mass_kg", 1)),
                ("Material cost",
                 f"USD {lining_result['material_cost_usd']:,.0f}"),
                ("Application labour",
                 f"USD {lining_result['labor_cost_usd']:,.0f}"),
                ("TOTAL COATING COST",
                 f"USD {lining_result['total_cost_usd']:,.0f}"),
            ] + [(k, v) for k, v in lining_result["detail"].items()
                 if k not in {"Material cost", "Labour cost", "Total cost"}])

        if s_hyd_prof:
            _sec("Hydraulic Head Profile")
            _hp_rows = [("Item", "Clean bed", "Dirty bed")]
            for k in hyd_prof["clean"]["items_bar"]:
                _hp_rows.append((
                    k,
                    fmt_bar_mwc(
                        hyd_prof["clean"]["items_bar"][k],
                        hyd_prof["clean"]["items_mwc"][k],
                    ),
                    fmt_bar_mwc(
                        hyd_prof["dirty"]["items_bar"][k],
                        hyd_prof["dirty"]["items_mwc"][k],
                    ),
                ))
            _hp_rows.append((
                "TOTAL PUMP DUTY",
                fmt_bar_mwc(hyd_prof["clean"]["total_bar"], hyd_prof["clean"]["total_mwc"]),
                fmt_bar_mwc(hyd_prof["dirty"]["total_bar"], hyd_prof["dirty"]["total_mwc"]),
            ))
            _tbl(_hp_rows, cols=("Item", "Clean bed", "Dirty bed"))

        if s_energy:
            _sec("Energy Consumption & OPEX")
            _tbl([
                ("Filtration pump power (dirty, per filter)",
                 fmt(energy["p_filt_dirty_kw"], "power_kw", 1)),
                ("BW pump power",            fmt(energy["p_bw_kw"], "power_kw", 1)),
                ("Air blower power (elec.)", fmt(energy["p_blower_elec_kw"], "power_kw", 1)),
                ("Annual total energy",
                 f"{energy['e_total_kwh_yr']/1e3:,.0f} MWh/yr"),
                ("Specific energy",          fmt(energy["kwh_per_m3"], "energy_kwh_m3", 4)),
                ("Annual energy OPEX",       f"USD {energy['cost_usd_yr']:,.0f}/yr"),
            ])

        if s_financial and econ_financial:
            _sec("Lifecycle Financial Summary")
            _irr = econ_financial.get("irr_pct")
            _roi = econ_financial.get("roi_pct")
            _tbl([
                ("NPV (net cash flow)", f"USD {econ_financial.get('npv', 0):,.0f}"),
                ("IRR", f"{_irr:.2f} %" if _irr is not None else "—"),
                ("ROI (simple)", f"{_roi:.1f} %" if _roi is not None else "—"),
                ("Simple payback (yr)",
                 f"{econ_financial['simple_payback_years']:.2f}"
                 if econ_financial.get("simple_payback_years") is not None else "—"),
                ("Discounted payback (yr)",
                 f"{econ_financial['discounted_payback_years']:.2f}"
                 if econ_financial.get("discounted_payback_years") is not None else "—"),
                ("Lifecycle cost (undisc. sum of net CF)", f"USD {econ_financial.get('lifecycle_cost', 0):,.0f}"),
                ("Annualized cost (LCOW × flow)", f"USD {econ_financial.get('annualized_cost', 0):,.0f}/yr"
                 if econ_financial.get("annualized_cost") is not None else "—"),
            ])
            _cf = econ_financial.get("cashflow_table") or []
            if _cf:
                _sec("Cash flow table (excerpt)")
                _hdr = list(_cf[0].keys())
                _rows = [_hdr] + [[str(r.get(h, "")) for h in _hdr] for r in _cf[: min(12, len(_cf))]]
                _tbl(_rows, cols=tuple(_hdr))
            _rs = econ_financial.get("replacement_schedule") or []
            if _rs:
                _sec("Replacement schedule")
                _tbl([["Year", "Events", "Spend (USD)"]] + [
                    [str(r.get("year", "")), ",".join(r.get("events", [])), f"{r.get('replacement_spend_usd', 0):,.0f}"]
                    for r in _rs
                ], cols=("Year", "Events", "Spend (USD)"))

        if s_cart:
            _sec("Cartridge Filter")
            _tbl([
                ("Design flow",       fmt(cart_result["design_flow_m3h"], "flow_m3h", 1)),
                ("Element material",  cart_result["element_material"]),
                ("Element size",      cart_result["element_size"]),
                ("Rating",            f"{cart_result['rating_um']} µm absolute"),
                ("Elements required", str(cart_result["n_elements"])),
                ("Housings required", str(cart_result["n_housings"])),
                ("ΔP BOL / EOL",
                 f"{fmt(cart_result['dp_clean_bar'], 'pressure_bar', 4)} / "
                 f"{fmt(cart_result['dp_eol_bar'], 'pressure_bar', 4)}"),
                ("DHC basis",
                 "Vendor datasheet" if cart_result.get("dhc_basis") == "vendor_override"
                 else "Model (g/TIE × rating)"),
                ("Annual element cost",
                 f"USD {cart_result['annual_cost_usd']:,.0f}"),
            ])

        if s_design_basis and _basis:
            _sec("Design Basis & Traceability")
            _tbl(design_basis_meta_rows(_basis))
            if _basis.get("assumptions"):
                _h("Assumptions", 3)
                for _a in _basis["assumptions"]:
                    doc.add_paragraph(plain_text(_a), style="List Bullet")
            if _basis.get("limits_and_criteria"):
                _h("Limits & criteria", 3)
                for _lim in _basis["limits_and_criteria"]:
                    doc.add_paragraph(plain_text(_lim), style="List Bullet")
            _tr = traceability_table_rows(_basis)
            if len(_tr) > 1:
                _h("Output traceability", 3)
                _tbl(_tr, cols=tuple(_tr[0]))
            _col = _basis.get("collector") or {}
            if _col:
                _h("Collector model summary", 3)
                _tbl([(r[0], r[1]) for r in collector_summary_rows(_col)])
                _chk = list(_col.get("design_checklist") or [])
                if _chk:
                    _h("Collector design checklist", 3)
                    for _line in _chk[:12]:
                        doc.add_paragraph(plain_text(_line), style="List Bullet")
            if _basis.get("exclusions"):
                _h("Exclusions (not modelled)", 3)
                for _ex in _basis["exclusions"]:
                    doc.add_paragraph(plain_text(_ex), style="List Bullet")

        doc.add_page_break()
        _h("Sign-off & Revision Record")
        _tbl([
            ("Prepared by",  engineer),
            ("Role",         "Process Expert — AQUASIGHT™"),
            ("Document No.", f"{doc_number}  ·  Rev {revision}"),
            ("Project",      project_name),
            ("Client",       client or "—"),
            ("Date",         str(__import__("datetime").date.today())),
            ("Software",     "AQUASIGHT™ MMF Calculator"),
            ("Checked by",   ""),
            ("Approved by",  ""),
        ])
        doc.add_paragraph(
            "\nThis document was generated automatically. "
            "All calculations are the responsibility of the engineer of record."
        )

        buf = _io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    _n_sections = sum([
        s_process, s_water, s_media, s_dp, s_cycle,
        s_vessel, s_nzpl, s_wt_empty, s_wt_oper, s_saddle,
        s_bw_hyd, s_bw_equip, s_lining and _has_lining,
        s_hyd_prof, s_energy, s_financial, s_design_basis, s_cart,
    ])
    st.caption(
        f"{_n_sections} section(s) selected + Identification & Sign-off (always included)")

    dl_word, dl_pdf, _ = st.columns([2, 2, 1])
    with dl_word:
        if _DOCX_OK:
            st.download_button(
                label="⬇️  Download Word report (.docx)",
                data=_build_docx(),
                file_name=f"{doc_number}_Rev{revision}.docx",
                mime=("application/vnd.openxmlformats-officedocument"
                      ".wordprocessingml.document"),
            )
        else:
            st.warning("Install python-docx:  `pip install python-docx`")
    with dl_pdf:
        if _PDF_OK:
            _pdf_sections = {
                "process": s_process, "water": s_water, "media": s_media,
                "dp": s_dp, "vessel": s_vessel, "bw_hyd": s_bw_hyd,
                "bw_equip": s_bw_equip, "energy": s_energy, "financial": s_financial,
                "design_basis": s_design_basis,
            }
            st.download_button(
                label="⬇️  Download PDF report (.pdf)",
                data=build_pdf(
                    inputs, computed, _pdf_sections,
                    st.session_state.get("unit_system", "metric"),
                ),
                file_name=f"{doc_number}_Rev{revision}.pdf",
                mime="application/pdf",
            )
        else:
            st.warning("Install reportlab:  `pip install reportlab`")

    st.divider()
    st.markdown(f"**{project_name}** · {doc_number} · Rev {revision} · {engineer}")
    st.markdown("")

    if s_process:
        st.markdown("### B1 · Process Basis")
        st.markdown(f"""
| Parameter | Value |
|---|---|
| Total plant flow | {fmt(total_flow, "flow_m3h", 0)} |
| Streams | {streams} |
| Total physical number of filters / stream | {n_filters} |
| Total physical filters (plant-wide) | {streams * n_filters} |
| Hydraulic assist | {hydraulic_assist} spare(s)/stream (design N = installed − spare) |
| Outage depth modelled | N-{redundancy} per stream |
| Flow / filter (N) | {fmt(q_per_filter, "flow_m3h", 1)} |
| Filtration rate (N) | {fmt(q_per_filter/avg_area, "velocity_m_h", 2)} |
| Cross-sectional area | {fmt(avg_area, "area_m2", 3)} |
""")

    if s_water:
        st.markdown("### B2 · Water Properties")
        st.markdown(f"""
| Property | Feed | Backwash |
|---|---|---|
| Salinity | {feed_sal:.2f} ppt | {bw_sal:.2f} ppt |
| Temperature | {fmt(feed_temp, "temperature_c", 1)} | {fmt(bw_temp, "temperature_c", 1)} |
| Density | {fmt(rho_feed, "density_kg_m3", 3)} | {fmt(rho_bw, "density_kg_m3", 3)} |
| Viscosity | {fmt(mu_feed*1000.0, "viscosity_cp", 4)} | {fmt(mu_bw*1000.0, "viscosity_cp", 4)} |
""")

    if s_media:
        st.markdown("### B3 · Media Configuration")
        _med_hdr = (
            f"| Media | Sup. | Depth ({ulbl('length_m')}) | d10 ({ulbl('length_mm')}) | CU | ε₀ | "
            f"ρp,eff ({ulbl('density_kg_m3')}) | ψ |"
        )
        _med_sep = "|---|---|---|---|---|---|---|---|"
        _med_rows = "\n".join(
            f"| {b['Type']} | {'✓' if b.get('is_support') else ''} "
            f"| {dv(b['Depth'], 'length_m'):.3f} | {dv(float(b['d10']), 'length_mm'):.2f} | {b['cu']:.2f} "
            f"| {b.get('epsilon0', 0):.3f} | {dv(b['rho_p_eff'], 'density_kg_m3'):.0f} "
            f"| {b.get('psi', '—')} |"
            for b in base
        )
        st.markdown(f"{_med_hdr}\n{_med_sep}\n{_med_rows}")

    if s_dp:
        st.markdown("### B4 · Filtration ΔP")
        st.markdown(f"""
| | Value |
|---|---|
| Clean-bed ΔP (Ergun) | {fmt(bw_dp['dp_clean_bar'], 'pressure_bar', 4)} |
| 50 % loaded ΔP | {fmt((bw_dp['dp_clean_bar']+bw_dp['dp_dirty_bar'])/2, 'pressure_bar', 4)} |
| Dirty ΔP (M_max) | {fmt(bw_dp['dp_dirty_bar'], 'pressure_bar', 4)} |
| BW trigger setpoint | {fmt(dp_trigger_bar, 'pressure_bar', 2)} |
| Superficial LV (N) | {fmt(bw_dp['u_m_h'], 'velocity_m_h', 2)} |
""")

    if s_vessel:
        st.markdown("### C1 · Vessel Dimensions & ASME Thickness")
        st.markdown(f"""
| Parameter | Value |
|---|---|
| Nominal ID / Real hyd. ID | {fmt(nominal_id, 'length_m', 3)} / {fmt(real_id, 'length_m', 4)} |
| Outside diameter | {fmt(mech['od_m'], 'length_m', 4)} |
| Total length T/T | {fmt(total_length, 'length_m', 3)} |
| End geometry | {end_geometry} |
| Material | {material_name} |
| Design pressure | {fmt(design_pressure, 'pressure_bar', 2)} gauge |
| Shell t_req / t_design | {dv(float(mech['t_shell_min_mm']), 'length_mm'):.2f} / {dv(float(mech['t_shell_design_mm']), 'length_mm'):.2f} {ulbl('length_mm')} |
| Head t_req / t_design | {dv(float(mech['t_head_min_mm']), 'length_mm'):.2f} / {dv(float(mech['t_head_design_mm']), 'length_mm'):.2f} {ulbl('length_mm')} |
""")

    if s_wt_empty:
        st.markdown("### C3 · Empty Weight")
        st.markdown(f"""
| Component | Weight |
|---|---|
| Shell + 2 heads | {fmt(wt_body['weight_body_kg'], 'mass_kg', 0)} |
| Nozzles | {fmt(w_noz, 'mass_kg', 0)} |
| Nozzle plate assembly | {fmt(wt_np['weight_total_kg'], 'mass_kg', 0)} |
| Supports | {fmt(wt_sup['weight_all_supports_kg'], 'mass_kg', 0)} |
| Internals | {fmt(wt_int['weight_internals_kg'], 'mass_kg', 0)} |
| **TOTAL EMPTY** | **{fmt(w_total_rep, 'mass_kg', 0)} = {fmt(w_total_rep/1000.0, 'mass_t', 3)}** |
""")

    if s_wt_oper:
        st.markdown("### C4 · Operating Weight & Support Loads")
        st.markdown(f"""
| Component | Weight |
|---|---|
| Empty vessel | {fmt(wt_oper['w_empty_kg'], 'mass_kg', 0)} |
| Lining / coating | {fmt(wt_oper['w_lining_kg'], 'mass_kg', 0)} |
| Media (dry) | {fmt(wt_oper['w_media_kg'], 'mass_kg', 0)} |
| Process water | {fmt(wt_oper['w_water_kg'], 'mass_kg', 0)} |
| **OPERATING WEIGHT** | **{fmt(wt_oper['w_operating_kg'], 'mass_kg', 0)} = {fmt(wt_oper['w_operating_t'], 'mass_t', 3)}** |
| Load / support ({wt_oper['n_supports']} supports) | {fmt(wt_oper['load_per_support_t'], 'mass_t', 3)} = {fmt(wt_oper['load_per_support_kN'], 'force_kn', 1)} |
""")

    if s_saddle:
        _sd2 = wt_saddle
        _pos2 = _sd2.get("saddle_positions_m") or [
            _sd2.get("saddle_1_from_left_m"),
            _sd2.get("saddle_2_from_left_m"),
        ]
        _pos2 = [p for p in _pos2 if p is not None]
        _pos2_txt = ", ".join(fmt(float(p), "length_m", 3) for p in _pos2) or "—"
        _sp2 = _sd2.get("saddle_spacings_m") or []
        _sp2_txt = (
            " / ".join(
                f"{dv(float(s), 'length_m'):.2f} {ulbl('length_m')}"
                for s in _sp2
            )
            if _sp2 else "—"
        )
        st.markdown("### C5 · Saddle Design")
        st.markdown(f"""
| Parameter | Value |
|---|---|
| Spacing factor α | {_sd2['alpha']:.2f} |
| Saddle positions (from left T/L, m) | {_pos2_txt} |
| Centre-to-centre spacings | {_sp2_txt} |
| Reaction per saddle | {fmt(_sd2['reaction_t'], 'mass_t', 3)} |
| Section selected | {_sd2.get('section', '—')} — capacity {fmt(float(_sd2['capacity_t']), 'mass_t', 3) if isinstance(_sd2.get('capacity_t'), (int, float)) else _sd2.get('capacity_t', '—')} |
| Status | {"OVERSTRESSED" if _sd2.get('overstressed') else "OK"} |
""")

    if s_bw_hyd:
        st.markdown("### D1 · Backwash Hydraulics & Collector")
        st.markdown(f"""
| Parameter | Value |
|---|---|
| BW velocity (proposed / max safe) | {fmt(bw_velocity, 'velocity_m_h', 1)} / {fmt(bw_col['max_safe_bw_m_h'], 'velocity_m_h', 1)} |
| Freeboard (min. req. / actual) | {fmt(freeboard_mm, 'length_mm', 0)} / {fmt(bw_col['freeboard_m'], 'length_m', 3)} |
| Collector status | {bw_col['status']} |
""")

    if s_bw_equip:
        _bws2 = bw_sizing
        st.markdown("### D2 · BW Equipment Sizing")
        st.markdown(f"""
| Equipment | Key Parameter |
|---|---|
| BW pump flow | {fmt(_bws2['q_bw_design_m3h'], 'flow_m3h', 1)} |
| BW pump head | {fmt(_bws2['bw_head_mwc'], 'pressure_mwc', 2)} |
| BW pump motor | {fmt(_bws2['p_pump_motor_kw'], 'power_kw', 1)} |
| Air blower flow | {fmt(_bws2['q_air_design_nm3h'], 'air_flow_nm3h', 1)} |
| Blower back-pressure | {fmt(_bws2['P2_pa']/1e5, 'pressure_bar', 3)} abs |
| Blower motor | {fmt(_bws2['p_blower_motor_kw'], 'power_kw', 1)} |
| BW water tank | {fmt(_bws2['v_tank_m3'], 'volume_m3', 1)} |
""")

    if s_lining and _has_lining:
        st.markdown(f"### E · Internal Protection — {lining_result['protection_type']}")
        st.markdown(f"""
| Parameter | Value |
|---|---|
| Total area | {fmt(lining_result['a_total_m2'], 'area_m2', 2)} |
| Weight | {fmt(lining_result['weight_kg'], 'mass_kg', 0)} |
| Material cost | USD {lining_result['material_cost_usd']:,.0f} |
| Labour cost | USD {lining_result['labor_cost_usd']:,.0f} |
| **Total coating cost** | **USD {lining_result['total_cost_usd']:,.0f}** |
""")

    if s_hyd_prof:
        st.markdown("### F1 · Hydraulic Head Profile")
        _hd_hdr = "| Item | Clean bed | Dirty bed |"
        _hd_sep = "|---|---|---|"
        _hd_rows = "\n".join(
            f"| {k} | {fmt_bar_mwc(hyd_prof['clean']['items_bar'][k], hyd_prof['clean']['items_mwc'][k])} | "
            f"{fmt_bar_mwc(hyd_prof['dirty']['items_bar'][k], hyd_prof['dirty']['items_mwc'][k])} |"
            for k in hyd_prof["clean"]["items_bar"]
        )
        _hd_tot = (
            f"| **Total** | **{fmt_bar_mwc(hyd_prof['clean']['total_bar'], hyd_prof['clean']['total_mwc'])}** | "
            f"**{fmt_bar_mwc(hyd_prof['dirty']['total_bar'], hyd_prof['dirty']['total_mwc'])}** |"
        )
        st.markdown(f"{_hd_hdr}\n{_hd_sep}\n{_hd_rows}\n{_hd_tot}")

    if s_energy:
        st.markdown("### F2 · Energy & OPEX")
        st.markdown(f"""
| KPI | Value |
|---|---|
| Filtration pump power (dirty, per filter) | {fmt(energy['p_filt_dirty_kw'], 'power_kw', 1)} |
| BW pump power | {fmt(energy['p_bw_kw'], 'power_kw', 1)} |
| Air blower power (elec.) | {fmt(energy['p_blower_elec_kw'], 'power_kw', 1)} |
| Annual total energy | {energy['e_total_kwh_yr']/1e3:,.0f} MWh/yr |
| Specific energy | {fmt(energy['kwh_per_m3'], 'energy_kwh_m3', 4)} |
| Annual energy OPEX | USD {energy['cost_usd_yr']:,.0f}/yr |
""")

    if s_financial and econ_financial:
        st.markdown("### F3 · Lifecycle financial")
        _fi = econ_financial
        _i2 = _fi.get("irr_pct")
        _r2 = _fi.get("roi_pct")
        _irr_txt = f"{_i2:.2f} %" if _i2 is not None else "—"
        _roi_txt = f"{_r2:.1f} %" if _r2 is not None else "—"
        st.markdown(f"""
| KPI | Value |
|---|---|
| NPV | USD {_fi.get('npv', 0):,.0f} |
| IRR | {_irr_txt} |
| ROI (simple) | {_roi_txt} |
| Summary | {_fi.get('economic_summary', '—')} |
""")
    if s_cart:
        st.markdown("### G · Cartridge Filter")
        st.markdown(f"""
| Parameter | Value |
|---|---|
| Design flow | {fmt(cart_result['design_flow_m3h'], 'flow_m3h', 1)} |
| Element | {cart_result['element_size']} · {cart_result['rating_um']} µm · {cart_result['element_material']} |
| Elements / Housings | {cart_result['n_elements']} / {cart_result['n_housings']} |
| Flow / element | {fmt(cart_result['actual_flow_m3h_element'], 'flow_m3h', 3)} ({fmt(cart_result['q_lpm_element'], 'flow_l_min', 1)}) |
| ΔP BOL / EOL | {fmt(cart_result['dp_clean_bar'], 'pressure_bar', 4)} / {fmt(cart_result['dp_eol_bar'], 'pressure_bar', 4)} |
| DHC basis | {"Vendor datasheet" if cart_result.get("dhc_basis") == "vendor_override" else "Model (g/TIE × rating)"} |
| Annual element cost | USD {cart_result['annual_cost_usd']:,.0f} |
""")

    st.divider()
    col_sign, _ = st.columns([1, 2])
    with col_sign:
        st.info(f"**{engineer}**  \nProcess Expert — AQUASIGHT™  \n\n"
                f"{doc_number} · Rev {revision}")
