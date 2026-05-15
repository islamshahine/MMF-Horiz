"""Requisition-style pump / blower datasheets from MMF hydraulics (not vendor submittals).

Exports are suitable for RFQ attachments, internal review, and JSON/database ingestion.
Dry-installed centrifugal liquid pumps + air blowers only — matches product scope in the UI.
"""
from __future__ import annotations

import io
import json
import re
from datetime import datetime, timezone
from typing import Any

try:
    from docx import Document as _DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN

    DOCX_OK = True
except ImportError:
    _DocxDocument = None  # type: ignore[misc, assignment]
    _WD_ALIGN = None  # type: ignore[misc, assignment]
    DOCX_OK = False

try:
    from reportlab.lib import colors as _rl_colors
    from reportlab.lib.pagesizes import A4 as _A4
    from reportlab.lib.styles import getSampleStyleSheet as _getSampleStyleSheet
    from reportlab.lib.units import mm as _mm
    from reportlab.platypus import (
        Paragraph as _RLParagraph,
        SimpleDocTemplate as _SimpleDocTemplate,
        Spacer as _RLSpacer,
        Table as _RLTable,
        TableStyle as _RLTableStyle,
    )

    PDF_OK = True
except ImportError:
    PDF_OK = False

SCHEMA_VERSION = "1.0"
GENERATOR = "AQUASIGHT™ MMF"

DISCLAIMER = (
    "This file is a **process duty and requisition basis** produced from AQUASIGHT™ MMF hydraulic "
    "and energy models. It is **not** a vendor-certified datasheet. **NPSHa**, full **HQ** curves, "
    "GA drawings, certified weights, nameplate data, and site piping are **TBA** by vendor / owner "
    "unless explicitly filled below."
)

DISCLAIMER_AIR_BLOWER = (
    "This file is a **process air duty and site-basis** for **air scour / combined air–water blowers** "
    "derived from AQUASIGHT™ MMF backwash hydraulics. It is **not** a vendor submittal. "
    "Guaranteed performance, surge control, acoustics, controls integration, anchor loads, "
    "and certified curves are **TBA** by the blower manufacturer unless attached separately."
)

DATASHEET_TYPE_LIQUID_PUMPS = "liquid_pumps"
DATASHEET_TYPE_AIR_BLOWER = "air_blower"


def default_blower_rfq_environment() -> dict[str, Any]:
    """Default site / ambient fields when session has no keys (tests, first open)."""
    return {
        "elevation_amsl_m": None,
        "ambient_temp_min_c": None,
        "ambient_temp_avg_c": None,
        "ambient_temp_max_c": None,
        "relative_humidity_min_pct": None,
        "relative_humidity_avg_pct": None,
        "relative_humidity_max_pct": None,
        "barometric_pressure_bara": None,
        "installation_class": None,
        "dust_salt_exposure_notes": None,
        "corrosive_atmosphere_notes": None,
        "noise_limit_dba": None,
        "electrical_area_classification": None,
        "site_location_notes": None,
    }


def _f(x: Any, default: float = 0.0) -> float:
    try:
        return float(x)
    except (TypeError, ValueError):
        return default


def _i(x: Any, default: int = 0) -> int:
    try:
        return int(x)
    except (TypeError, ValueError):
        return default


def hydraulic_fluid_power_kw(q_m3h: float, head_mwc: float, rho_kg_m3: float) -> float:
    """Ideal liquid hydraulic power (kW) = ρ g Q H with Q in m³/h."""
    g = 9.80665
    q_m3s = max(q_m3h, 0.0) / 3600.0
    return rho_kg_m3 * g * q_m3s * max(head_mwc, 0.0) / 1000.0


def default_ui_snapshot() -> dict[str, Any]:
    """Defaults when Streamlit session keys are absent (e.g. tests, API)."""
    return {
        "pp_n_feed_parallel": 1,
        "pp_n_bw_dol": 3,
        "pp_n_bw_vfd": 2,
        "pp_n_blowers": 1,
        "pp_econ_bw_phil": "DOL",
        "pp_blower_mode": "single_duty",
        "pp_feed_orient": "Horizontal",
        "pp_feed_std": "ISO 5199",
        "pp_feed_mat": "SS316",
        "pp_feed_seal": "Single mechanical seal",
        "pp_feed_vfd": False,
        "pp_bw_orient": "Horizontal",
        "pp_bw_std": "ISO 5199",
        "pp_bw_mat": "SS316",
        "pp_bw_seal": "Single mechanical seal",
        "pp_bw_vfd_allow": True,
        "pp_feed_iec": "IE3",
    }


def build_pump_datasheet_bundle(
    inputs: dict[str, Any],
    computed: dict[str, Any],
    ui: dict[str, Any] | None = None,
    *,
    export_timestamp_utc: str | None = None,
) -> dict[str, Any]:
    """Assemble JSON-serializable bundle: document control + **feed and backwash liquid pumps only**."""
    ui_m = {**default_ui_snapshot(), **(ui or {})}
    ts = export_timestamp_utc or datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    pp: dict[str, Any] = computed.get("pump_perf") or {}
    auto: dict[str, Any] = pp.get("auto") or {}
    feed_p: dict[str, Any] = pp.get("feed_pump") or {}
    bw_p: dict[str, Any] = pp.get("bw_pump") or {}

    feed_wp: dict[str, Any] = computed.get("feed_wp") or {}
    bw_wp: dict[str, Any] = computed.get("bw_wp") or {}
    env: dict[str, Any] = computed.get("env_structural") or {}

    streams = max(1, _i(auto.get("streams"), 1))
    total_flow = _f(auto.get("total_flow_m3h"))
    n_feed_par = max(1, _i(ui_m.get("pp_n_feed_parallel"), 1))
    q_stream = total_flow / streams
    q_each_feed = q_stream / n_feed_par

    h_dirty = _f(auto.get("head_dirty_mwc"))
    h_clean = _f(auto.get("head_clean_mwc"))
    rho_f = _f(auto.get("rho_feed_kg_m3"))
    rho_bw = _f(auto.get("rho_bw_kg_m3"))
    q_bw = _f(auto.get("q_bw_design_m3h"))
    bw_head = _f(auto.get("bw_head_mwc"))

    motor_iec_feed = _f(feed_p.get("motor_iec_kw"))
    p_elec_dirty = _f(feed_p.get("p_motor_elec_dirty_kw"))
    eta_p_est = _f(feed_p.get("eta_pump_est"))

    residual_bar = _f(auto.get("dp_residual_bar"))

    doc = {
        "schema_version": SCHEMA_VERSION,
        "generator": GENERATOR,
        "disclaimer": DISCLAIMER,
        "export_timestamp_utc": ts,
        "document_control": {
            "project_name": str(inputs.get("project_name") or "").strip() or None,
            "document_number": str(inputs.get("doc_number") or "").strip() or None,
            "unit_system": str(inputs.get("unit_system") or "metric"),
            "equipment_family": "Horizontal MMF — filtration / backwash",
            "revision": "A",
            "external_environment": str(inputs.get("external_environment") or ""),
        },
        "application": {
            "installation_notes": str(env.get("external_environment") or inputs.get("external_environment") or ""),
            "noise_limit_dba": None,
            "hazardous_area": "TBA (project electrical / area classification)",
            "scope_note": "This export is **liquid pumps only** (feed + backwash). Air blowers use the separate **Air blower RFQ** datasheet in the Pumps & power tab.",
        },
        "datasheet_type": DATASHEET_TYPE_LIQUID_PUMPS,
        "ui_selections": {k: v for k, v in ui_m.items() if k.startswith("pp_")},
        "equipment": [],
    }

    sg_f = round(rho_f / 1000.0, 4) if rho_f else None
    feed_fluid = {
        "phase": "Liquid",
        "density_kg_m3": round(rho_f, 2),
        "specific_gravity": sg_f,
        "dynamic_viscosity_cp": feed_wp.get("viscosity_cp"),
        "temperature_operating_c": _f(inputs.get("feed_temp")),
        "temperature_design_c": _f(inputs.get("design_temp") or inputs.get("feed_temp")),
        "salinity_ppt": _f(inputs.get("feed_sal")),
        "tds_mg_l_approx": feed_wp.get("tds_mg_l"),
        "vapor_pressure_bara": None,
        "ph": None,
        "chloride_mg_l": None,
        "tss_mg_l": None,
    }

    p_hyd_dirty = round(hydraulic_fluid_power_kw(q_each_feed, h_dirty, rho_f), 3)
    p_hyd_clean = round(hydraulic_fluid_power_kw(q_each_feed, h_clean, rho_f), 3)

    feed_block: dict[str, Any] = {
        "role": "filtration_feed_pump",
        "equipment_type": "Centrifugal pump (dry-installed; horizontal or vertical dry-pit per UI)",
        "service": "MMF / header filtration feed",
        "configuration": {
            "orientation": str(ui_m.get("pp_feed_orient")),
            "quality_standard": str(ui_m.get("pp_feed_std")),
            "wetted_material_preference": str(ui_m.get("pp_feed_mat")),
            "seal_type_preference": str(ui_m.get("pp_feed_seal")),
            "vfd_on_feed": bool(ui_m.get("pp_feed_vfd")),
            "motor_efficiency_class": str(ui_m.get("pp_feed_iec") or inputs.get("motor_iec_class") or "IE3"),
            "streams": streams,
            "parallel_pumps_per_stream": n_feed_par,
            "installed_pumps_total": streams * n_feed_par,
        },
        "fluid_properties": feed_fluid,
        "operating_conditions": {
            "model_residual_pressure_target_barg": residual_bar,
            "duty_points": [
                {
                    "case": "Clean bed (reference)",
                    "flow_m3h": round(q_each_feed, 3),
                    "differential_head_mwc": round(h_clean, 3),
                    "npsha_m": None,
                    "hydraulic_fluid_power_kw": p_hyd_clean,
                    "pump_efficiency_assumed": round(eta_p_est, 3),
                    "absorbed_power_electric_model_kw": round(_f(feed_p.get("p_motor_elec_clean_kw")), 3),
                },
                {
                    "case": "Dirty bed (design basis for motor sizing)",
                    "flow_m3h": round(q_each_feed, 3),
                    "differential_head_mwc": round(h_dirty, 3),
                    "npsha_m": None,
                    "hydraulic_fluid_power_kw": p_hyd_dirty,
                    "pump_efficiency_assumed": round(eta_p_est, 3),
                    "absorbed_power_electric_model_kw": round(p_elec_dirty, 3),
                },
            ],
            "iec_snap_motor_kw_per_pump": round(motor_iec_feed, 2),
            "notes": [
                "Flow is **per running pump** on the stream header (parallel split).",
                "NPSHa requires owner suction line, static lift, and strainer loss — not evaluated in MMF.",
            ],
        },
    }

    bw_fluid = {
        "phase": "Liquid",
        "density_kg_m3": round(rho_bw, 2),
        "specific_gravity": round(rho_bw / 1000.0, 4) if rho_bw else None,
        "dynamic_viscosity_cp": bw_wp.get("viscosity_cp"),
        "temperature_operating_c": _f(inputs.get("bw_temp")),
        "temperature_design_c": _f(inputs.get("design_temp") or inputs.get("bw_temp")),
        "salinity_ppt": _f(inputs.get("bw_sal")),
        "tds_mg_l_approx": bw_wp.get("tds_mg_l"),
        "vapor_pressure_bara": None,
    }
    p_bw_hyd = round(hydraulic_fluid_power_kw(q_bw, bw_head, rho_bw), 3)
    p_bw_rated = _f(bw_p.get("p_rated_elec_kw"))
    bw_block: dict[str, Any] = {
        "role": "backwash_liquid_pump",
        "equipment_type": "Centrifugal pump (dry-installed)",
        "service": "MMF backwash water (sequence staging per BW model)",
        "configuration": {
            "orientation": str(ui_m.get("pp_bw_orient")),
            "quality_standard": str(ui_m.get("pp_bw_std")),
            "wetted_material_preference": str(ui_m.get("pp_bw_mat")),
            "seal_type_preference": str(ui_m.get("pp_bw_seal")),
            "economics_bw_philosophy": str(ui_m.get("pp_econ_bw_phil")),
            "vfd_allowance_screening": bool(ui_m.get("pp_bw_vfd_allow")),
            "installed_trains_dol_philosophy": _i(ui_m.get("pp_n_bw_dol"), 3),
            "installed_trains_vfd_philosophy": _i(ui_m.get("pp_n_bw_vfd"), 2),
        },
        "fluid_properties": bw_fluid,
        "operating_conditions": {
            "duty_points": [
                {
                    "case": "Rated BW water (model design point)",
                    "flow_m3h": round(q_bw, 3),
                    "differential_head_mwc": round(bw_head, 3),
                    "npsha_m": None,
                    "hydraulic_fluid_power_kw": p_bw_hyd,
                    "absorbed_power_electric_model_kw": round(p_bw_rated, 3),
                },
            ],
            "iec_snap_motor_kw_dol_half_train": _f(bw_p.get("motor_iec_kw_dol_half")),
            "iec_snap_motor_kw_vfd_full_train": _f(bw_p.get("motor_iec_kw_vfd_full")),
            "notes": [
                "IEC motor snaps are screening values from the energy model — vendor curves govern.",
                "Installed train counts are **CAPEX / redundancy** inputs; hydraulics follow BW sequence table.",
            ],
        },
    }

    doc["equipment"] = [feed_block, bw_block]
    doc["model_warnings"] = list(pp.get("warnings") or [])
    return doc


def build_air_blower_datasheet_bundle(
    inputs: dict[str, Any],
    computed: dict[str, Any],
    ui: dict[str, Any] | None,
    blower_env: dict[str, Any] | None,
    *,
    export_timestamp_utc: str | None = None,
) -> dict[str, Any]:
    """RFQ bundle for **air scour blowers only**, including site / ambient fields for manufacturers."""
    ui_m = {**default_ui_snapshot(), **(ui or {})}
    env_m = {**default_blower_rfq_environment(), **(blower_env or {})}
    ts = export_timestamp_utc or datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    pp: dict[str, Any] = computed.get("pump_perf") or {}
    auto: dict[str, Any] = pp.get("auto") or {}
    bl: dict[str, Any] = pp.get("blower") or {}
    bl_detail: dict[str, Any] = (bl.get("detail") or {}) if isinstance(bl.get("detail"), dict) else {}
    str_env: dict[str, Any] = computed.get("env_structural") or {}

    q_nm3h = _f(auto.get("q_air_design_nm3h"))
    q_m3h_air = _f(auto.get("q_air_design_m3h"))
    blower_block: dict[str, Any] = {
        "role": "backwash_air_blower",
        "equipment_type": "Air blower (PD typical at moderate duty; centrifugal possible at high duty — see hint)",
        "service": "MMF air scour / combined air-water steps",
        "configuration": {
            "installed_units": _i(ui_m.get("pp_n_blowers"), 1),
            "operating_mode": str(ui_m.get("pp_blower_mode")),
        },
        "gas_duty": {
            "design_flow_nm3h_0C_1atm": round(q_nm3h, 2) if q_nm3h else None,
            "design_flow_m3h_inlet": round(q_m3h_air, 2) if q_m3h_air else None,
            "pressure_ratio": _f(auto.get("blower_pressure_ratio")),
            "motor_shaft_power_model_kw": _f(bl.get("p_motor_kw")),
            "technology_hint": str(bl.get("technology_hint") or ""),
            "detail": {k: bl_detail.get(k) for k in (
                "blower_air_delta_p_bar", "dp_airside_bar", "vessel_pressure_bar",
                "p_blower_ideal_kw", "p_blower_shaft_kw", "blower_eta",
            )},
        },
    }

    _warnings = list(pp.get("warnings") or [])
    bw_warn = []
    for w in _warnings:
        ws = str(w).lower()
        if "blower" in ws or "air" in ws or "npsh" in ws or "velocity" in ws:
            bw_warn.append(w)
    doc = {
        "schema_version": SCHEMA_VERSION,
        "generator": GENERATOR,
        "disclaimer": DISCLAIMER_AIR_BLOWER,
        "export_timestamp_utc": ts,
        "datasheet_type": DATASHEET_TYPE_AIR_BLOWER,
        "document_control": {
            "project_name": str(inputs.get("project_name") or "").strip() or None,
            "document_number": str(inputs.get("doc_number") or "").strip() or None,
            "unit_system": str(inputs.get("unit_system") or "metric"),
            "equipment_family": "Horizontal MMF — air scour / backwash blower",
            "revision": "A",
            "external_environment": str(inputs.get("external_environment") or ""),
        },
        "blower_environment": env_m,
        "application": {
            "service": "Air scour and/or combined air–water backwash",
            "duty": "Intermittent / cyclic (per BW sequence model)",
            "blower_inlet_temp_model_c": _f(inputs.get("blower_inlet_temp_c")),
            "vessel_gauge_pressure_input_bar": _f(inputs.get("vessel_pressure_bar")),
            "installation_context_structural": str(str_env.get("external_environment") or ""),
        },
        "ui_selections": {k: v for k, v in ui_m.items() if k.startswith("pp_")},
        "equipment": [blower_block],
        "model_warnings": bw_warn if bw_warn else _warnings,
    }
    return doc


def bundle_to_json(bundle: dict[str, Any]) -> str:
    return json.dumps(bundle, indent=2, ensure_ascii=False)


def collect_datasheet_parts(bundle: dict[str, Any], *, full_template: bool) -> list[dict[str, Any]]:
    """Structured blocks for **liquid pumps** datasheet (Markdown / Word / PDF)."""
    parts: list[dict[str, Any]] = []
    dc = bundle.get("document_control") or {}
    parts.append({"type": "heading", "level": 1, "text": "Liquid pumps requisition datasheet"})
    parts.append(
        {
            "type": "paragraph",
            "text": (
                f"{bundle.get('generator')} · schema {bundle.get('schema_version')} · "
                f"exported {bundle.get('export_timestamp_utc')}"
            ),
        }
    )
    parts.append({"type": "paragraph", "text": str(bundle.get("disclaimer") or "")})

    parts.append({"type": "heading", "level": 2, "text": "1. Document control"})
    parts.append(
        {
            "type": "table",
            "headers": ["Field", "Value"],
            "rows": [
                ["Project", dc.get("project_name") or "—"],
                ["Document #", dc.get("document_number") or "—"],
                ["Unit system", dc.get("unit_system") or "metric"],
                ["Equipment family", dc.get("equipment_family") or "—"],
                ["Revision", dc.get("revision") or "A"],
                ["Environment (inputs)", dc.get("external_environment") or "—"],
            ],
        }
    )

    app = bundle.get("application") or {}
    parts.append({"type": "heading", "level": 2, "text": "2. Application"})
    parts.append(
        {
            "type": "table",
            "headers": ["Parameter", "Value"],
            "rows": [
                ["External / installation context", app.get("installation_notes") or "—"],
                ["Noise limit dB(A)", app.get("noise_limit_dba") or ("TBA" if full_template else "—")],
                ["Hazardous area", app.get("hazardous_area") or "—"],
                ["Export scope", app.get("scope_note") or "—"],
            ],
        }
    )

    for eq in bundle.get("equipment") or []:
        role = str(eq.get("role") or "equipment")
        title = {
            "filtration_feed_pump": "3. Filtration feed pump",
            "backwash_liquid_pump": "4. Backwash liquid pump",
        }.get(role, role)
        parts.append({"type": "heading", "level": 2, "text": title})
        parts.append({"type": "paragraph", "text": f"Type: {eq.get('equipment_type')}"})
        parts.append({"type": "paragraph", "text": f"Service: {eq.get('service')}"})

        conf = eq.get("configuration") or {}
        if conf:
            parts.append({"type": "heading", "level": 3, "text": "Configuration"})
            parts.append(
                {
                    "type": "table",
                    "headers": ["Parameter", "Value"],
                    "rows": [[str(k), str(v)] for k, v in conf.items()],
                }
            )

        fp = eq.get("fluid_properties")
        if fp:
            parts.append({"type": "heading", "level": 3, "text": "Fluid properties"})
            rows = [
                [str(k), ("TBA" if full_template else "—") if v is None else str(v)]
                for k, v in fp.items()
            ]
            parts.append({"type": "table", "headers": ["Parameter", "Value"], "rows": rows})

        oc = eq.get("operating_conditions") or {}
        gd = eq.get("gas_duty") or {}
        if oc:
            parts.append({"type": "heading", "level": 3, "text": "Operating conditions"})
            if oc.get("model_residual_pressure_target_barg") is not None:
                parts.append(
                    {
                        "type": "paragraph",
                        "text": (
                            "Model downstream residual target (hydraulic profile): "
                            f"{oc['model_residual_pressure_target_barg']} barg"
                        ),
                    }
                )
            dps = oc.get("duty_points") or []
            if dps:
                hdr = [str(h).replace("_", " ").title() for h in dps[0].keys()]
                data_rows: list[list[str]] = []
                for row in dps:
                    data_rows.append(["" if row.get(h) is None else str(row.get(h)) for h in dps[0].keys()])
                parts.append({"type": "table", "headers": hdr, "rows": data_rows})
            snap_lines: list[str] = []
            for k in (
                "iec_snap_motor_kw_per_pump",
                "iec_snap_motor_kw_dol_half_train",
                "iec_snap_motor_kw_vfd_full_train",
            ):
                if k in oc and oc[k] is not None:
                    snap_lines.append(f"{k.replace('_', ' ')}: {oc[k]}")
            for line in snap_lines:
                parts.append({"type": "paragraph", "text": line})
            for note in oc.get("notes") or []:
                parts.append({"type": "paragraph", "text": f"Note: {note}"})
        if gd:
            parts.append({"type": "heading", "level": 3, "text": "Gas / blower duty"})
            rows_g = [[str(k), str(v)] for k, v in gd.items() if k != "detail"]
            parts.append({"type": "table", "headers": ["Parameter", "Value"], "rows": rows_g})
            det = gd.get("detail") or {}
            if any(det.values()):
                parts.append({"type": "paragraph", "text": "Sizing detail (model)"})
                parts.append(
                    {
                        "type": "table",
                        "headers": ["Key", "Value"],
                        "rows": [[str(k), str(det[k])] for k in det],
                    }
                )

    parts.append({"type": "heading", "level": 2, "text": "6. Vendor documentation (typical RFQ)"})
    parts.append(
        {
            "type": "bullets",
            "items": [
                "GA drawing, certified performance curves, cross-section",
                "Motor datasheet, complete weights, foundation loads",
                "NPSH test / witness requirements per contract",
                "O&M manual, commissioning procedure, spare parts list",
            ],
        }
    )

    if full_template:
        parts.append({"type": "heading", "level": 2, "text": "7. Vendor / site data (TBA)"})
        parts.append(
            {
                "type": "table",
                "headers": ["Item", "Status"],
                "rows": [
                    ["Nozzle sizes & ratings", "TBA"],
                    ["Mechanical seal plan & auxiliaries", "TBA"],
                    ["Bearings & lubrication", "TBA"],
                    ["Baseplate / coupling / guard", "TBA"],
                    ["Weights (pump, motor, skid)", "TBA"],
                    ["NPSHa (all cases)", "TBA"],
                    ["Witness testing & FAT", "TBA per contract"],
                    ["Painting / coating system", "TBA"],
                ],
            }
        )

    mw = bundle.get("model_warnings") or []
    if mw:
        parts.append({"type": "heading", "level": 2, "text": "Model warnings"})
        parts.append({"type": "bullets", "items": [str(w) for w in mw]})

    return parts


def collect_air_blower_datasheet_parts(bundle: dict[str, Any], *, full_template: bool) -> list[dict[str, Any]]:
    """Structured blocks for **air blower RFQ** datasheet."""
    parts: list[dict[str, Any]] = []
    dc = bundle.get("document_control") or {}
    benv: dict[str, Any] = bundle.get("blower_environment") or {}

    def _cell(label: str, key: str) -> list[str]:
        v = benv.get(key)
        return [label, "—" if v is None or v == "" else str(v)]

    parts.append({"type": "heading", "level": 1, "text": "Air scour blower requisition datasheet"})
    parts.append(
        {
            "type": "paragraph",
            "text": (
                f"{bundle.get('generator')} · schema {bundle.get('schema_version')} · "
                f"exported {bundle.get('export_timestamp_utc')}"
            ),
        }
    )
    parts.append({"type": "paragraph", "text": str(bundle.get("disclaimer") or "")})

    parts.append({"type": "heading", "level": 2, "text": "1. Document control"})
    parts.append(
        {
            "type": "table",
            "headers": ["Field", "Value"],
            "rows": [
                ["Project", dc.get("project_name") or "—"],
                ["Document #", dc.get("document_number") or "—"],
                ["Unit system", dc.get("unit_system") or "metric"],
                ["Equipment family", dc.get("equipment_family") or "—"],
                ["Revision", dc.get("revision") or "A"],
                ["Plant environment (MMF inputs)", dc.get("external_environment") or "—"],
            ],
        }
    )

    parts.append({"type": "heading", "level": 2, "text": "2. Site & ambient environment (for manufacturer)"})
    parts.append(
        {
            "type": "table",
            "headers": ["Parameter", "Value"],
            "rows": [
                _cell("Site elevation above mean sea level (m AMSL)", "elevation_amsl_m"),
                _cell("Ambient temperature — minimum (°C)", "ambient_temp_min_c"),
                _cell("Ambient temperature — average (°C)", "ambient_temp_avg_c"),
                _cell("Ambient temperature — maximum (°C)", "ambient_temp_max_c"),
                _cell("Relative humidity — minimum (%)", "relative_humidity_min_pct"),
                _cell("Relative humidity — average (%)", "relative_humidity_avg_pct"),
                _cell("Relative humidity — maximum (%)", "relative_humidity_max_pct"),
                _cell("Barometric pressure — site / design (bara)", "barometric_pressure_bara"),
                _cell("Installation class (indoor / outdoor / shelter)", "installation_class"),
                _cell("Dust / salt / sand exposure notes", "dust_salt_exposure_notes"),
                _cell("Corrosive / chemical atmosphere notes", "corrosive_atmosphere_notes"),
                _cell("Purchaser noise limit dB(A) @ reference", "noise_limit_dba"),
                _cell("Electrical hazardous area classification", "electrical_area_classification"),
                _cell("Site location / address notes", "site_location_notes"),
            ],
        }
    )

    appb = bundle.get("application") or {}

    def _av(key: str) -> str:
        v = appb.get(key)
        return "—" if v is None or v == "" else str(v)

    parts.append({"type": "heading", "level": 2, "text": "3. Service & process basis (from MMF model)"})
    parts.append(
        {
            "type": "table",
            "headers": ["Parameter", "Value"],
            "rows": [
                ["Service", str(appb.get("service") or "—")],
                ["Duty", str(appb.get("duty") or "—")],
                ["Blower inlet temperature used in model (°C)", _av("blower_inlet_temp_model_c")],
                ["Vessel gauge pressure input (bar)", _av("vessel_gauge_pressure_input_bar")],
                ["Structural / external context (MMF)", str(appb.get("installation_context_structural") or "—")],
            ],
        }
    )

    eq = (bundle.get("equipment") or [{}])[0]
    parts.append(
        {"type": "heading", "level": 2, "text": "4. Process air duty & sizing (from MMF hydraulic model)"}
    )
    parts.append({"type": "paragraph", "text": f"Type: {eq.get('equipment_type')}"})
    parts.append({"type": "paragraph", "text": f"Service: {eq.get('service')}"})
    conf = eq.get("configuration") or {}
    if conf:
        parts.append({"type": "heading", "level": 3, "text": "Train configuration (from Pumps tab)"})
        parts.append(
            {
                "type": "table",
                "headers": ["Parameter", "Value"],
                "rows": [[str(k), str(v)] for k, v in conf.items()],
            }
        )
    gd = eq.get("gas_duty") or {}
    if gd:
        parts.append({"type": "heading", "level": 3, "text": "Gas duty & model power"})
        rows_g = [[str(k), str(v)] for k, v in gd.items() if k != "detail"]
        parts.append({"type": "table", "headers": ["Parameter", "Value"], "rows": rows_g})
        det = gd.get("detail") or {}
        if any(det.values()):
            parts.append({"type": "paragraph", "text": "Sizing detail (MMF backwash equipment model)"})
            parts.append(
                {
                    "type": "table",
                    "headers": ["Key", "Value"],
                    "rows": [[str(k), str(det[k])] for k in det],
                }
            )

    parts.append({"type": "heading", "level": 2, "text": "5. Vendor documentation (typical blower RFQ)"})
    parts.append(
        {
            "type": "bullets",
            "items": [
                "Certified air performance curves (including surge / choke margins if centrifugal)",
                "Motor and VFD datasheets, harmonics / filtering philosophy if applicable",
                "Complete weights, anchor loads, inertia block / skid requirements",
                "Noise breakout (inlet / discharge / casing), silencer options, compliance to purchaser noise criteria",
                "Anti-surge / recycle strategy and control narrative (DCS interface)",
                "Spares list, shop and field testing, FAT witness matrix",
            ],
        }
    )

    if full_template:
        parts.append({"type": "heading", "level": 2, "text": "6. Vendor / site data (TBA)"})
        parts.append(
            {
                "type": "table",
                "headers": ["Item", "Status"],
                "rows": [
                    ["Foundation bolts and grouting loads", "TBA"],
                    ["Inlet filtration / silencers / discharge plenum", "TBA"],
                    ["Package scope (common base, local panel, vibration probes)", "TBA"],
                    ["Ambient / altitude derating confirmation vs §2", "TBA"],
                    ["Coastal / C5-M coating if marine site", "TBA"],
                    ["Witness performance test at reference conditions", "TBA per contract"],
                ],
            }
        )

    mw = bundle.get("model_warnings") or []
    if mw:
        parts.append({"type": "heading", "level": 2, "text": "Model warnings (air / hydraulics context)"})
        parts.append({"type": "bullets", "items": [str(w) for w in mw]})

    return parts


def _md_table(headers: list[str], rows: list[list[Any]]) -> str:
    esc = lambda s: str(s).replace("|", "\\|")
    lines = ["| " + " | ".join(esc(h) for h in headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(esc(c) for c in row) + " |")
    return "\n".join(lines) + "\n"


def parts_to_markdown(parts: list[dict[str, Any]]) -> str:
    lines: list[str] = []
    for p in parts:
        t = p["type"]
        if t == "heading":
            lvl = int(p["level"])
            lines.append(f"{'#' * lvl} {p['text']}\n")
        elif t == "paragraph":
            lines.append(f"{p['text']}\n")
        elif t == "table":
            lines.append(
                _md_table(
                    list(p["headers"]),
                    [list(r) for r in p["rows"]],
                )
            )
        elif t == "bullets":
            for item in p["items"]:
                lines.append(f"- {item}\n")
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def _bundle_parts_meta(
    bundle: dict[str, Any], *, full_template: bool
) -> tuple[list[dict[str, Any]], str, str, frozenset[str]]:
    """Return (parts, pdf_banner_line, docx_subtitle, skip_duplicate_h1_texts)."""
    if bundle.get("datasheet_type") == DATASHEET_TYPE_AIR_BLOWER:
        return (
            collect_air_blower_datasheet_parts(bundle, full_template=full_template),
            "AQUASIGHT™ MMF — Air scour blower requisition datasheet",
            "Air scour blower requisition datasheet",
            frozenset({"Air scour blower requisition datasheet"}),
        )
    return (
        collect_datasheet_parts(bundle, full_template=full_template),
        "AQUASIGHT™ MMF — Liquid pumps requisition datasheet",
        "Liquid pumps requisition datasheet",
        frozenset({"Liquid pumps requisition datasheet"}),
    )


def bundle_to_markdown(bundle: dict[str, Any], *, full_template: bool) -> str:
    """Markdown datasheet: **duty** = model-filled only; **full** = duty + vendor TBA tables."""
    parts, _, _, _ = _bundle_parts_meta(bundle, full_template=full_template)
    return parts_to_markdown(parts)


def bundle_to_docx_bytes(bundle: dict[str, Any], *, full_template: bool) -> bytes:
    """Microsoft Word (.docx) for non-technical stakeholders."""
    if not DOCX_OK or _DocxDocument is None or _WD_ALIGN is None:
        raise RuntimeError("Install python-docx: pip install python-docx")
    parts, _pdf_banner, docx_sub, skip_h1 = _bundle_parts_meta(bundle, full_template=full_template)
    doc = _DocxDocument()
    t0 = doc.add_heading("AQUASIGHT™ MMF", 0)
    t0.alignment = _WD_ALIGN.CENTER
    t1 = doc.add_heading(docx_sub, 1)
    t1.alignment = _WD_ALIGN.CENTER
    doc.add_paragraph("")

    for p in parts:
        t = p["type"]
        if t == "heading":
            lvl = min(max(int(p["level"]), 1), 3)
            txt = str(p["text"])
            if lvl == 1 and txt.strip() in skip_h1:
                continue
            doc.add_heading(txt, lvl)
        elif t == "paragraph":
            doc.add_paragraph(str(p["text"]))
        elif t == "table":
            hdrs = list(p["headers"])
            rows = [list(r) for r in p["rows"]]
            tbl = doc.add_table(rows=1 + len(rows), cols=len(hdrs))
            tbl.style = "Table Grid"
            for j, h in enumerate(hdrs):
                tbl.rows[0].cells[j].text = str(h)
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    tbl.rows[i + 1].cells[j].text = str(cell)
            doc.add_paragraph("")
        elif t == "bullets":
            for item in p["items"]:
                doc.add_paragraph(str(item), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def bundle_to_pdf_bytes(bundle: dict[str, Any], *, full_template: bool) -> bytes:
    """PDF (ReportLab) for printing and email attachment."""
    if not PDF_OK:
        raise RuntimeError("Install reportlab: pip install reportlab")
    from xml.sax.saxutils import escape

    parts, pdf_banner, _, skip_h1 = _bundle_parts_meta(bundle, full_template=full_template)
    buf = io.BytesIO()
    w_page, _h_page = _A4
    margin = 14 * _mm
    doc = _SimpleDocTemplate(
        buf,
        pagesize=_A4,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    )
    stl = _getSampleStyleSheet()
    story: list[Any] = []
    body = stl["Normal"]
    body.fontSize = 9
    body.leading = 11
    h1s = stl["Heading1"]
    h1s.spaceAfter = 6
    h2s = stl["Heading2"]
    h2s.spaceBefore = 8
    h2s.spaceAfter = 4
    h3s = stl["Heading3"]
    h3s.spaceBefore = 4
    h3s.spaceAfter = 2

    def _para(txt: str, style) -> _RLParagraph:
        return _RLParagraph(escape(str(txt)).replace("\n", "<br/>"), style)

    tbl_style = _RLTableStyle(
        [
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.25, _rl_colors.HexColor("#cccccc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [_rl_colors.white, _rl_colors.HexColor("#f5f5f5")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]
    )

    usable_w = w_page - 2 * margin

    story.append(_para(pdf_banner, h1s))
    story.append(_RLSpacer(1, 2 * _mm))

    skip_titles = skip_h1

    for p in parts:
        t = p["type"]
        if t == "heading":
            txt = str(p["text"])
            if txt in skip_titles:
                continue
            lvl = int(p["level"])
            style = h1s if lvl <= 1 else h2s if lvl == 2 else h3s
            story.append(_para(txt, style))
        elif t == "paragraph":
            story.append(_para(p["text"], body))
            story.append(_RLSpacer(1, 1 * _mm))
        elif t == "table":
            hdrs = [str(h) for h in p["headers"]]
            rows = [[str(c) for c in r] for r in p["rows"]]
            data = [[_para(c, body) for c in hdrs]] + [[_para(c, body) for c in r] for r in rows]
            nc = len(hdrs)
            col_w = usable_w / max(nc, 1)
            tbl = _RLTable(data, colWidths=[col_w] * nc, repeatRows=1)
            tbl.setStyle(tbl_style)
            story.append(tbl)
            story.append(_RLSpacer(1, 2 * _mm))
        elif t == "bullets":
            for item in p["items"]:
                story.append(_para(f"• {item}", body))
            story.append(_RLSpacer(1, 2 * _mm))

    doc.build(story)
    return buf.getvalue()


def list_datasheet_export_choices(
    *,
    docx_ok: bool | None = None,
    pdf_ok: bool | None = None,
) -> list[tuple[str, str]]:
    """(choice_id, label) for selectbox — omits Word/PDF when optional libs missing."""
    _docx = DOCX_OK if docx_ok is None else docx_ok
    _pdf = PDF_OK if pdf_ok is None else pdf_ok
    choices: list[tuple[str, str]] = [
        ("md_duty", "Duty — Markdown (.md)"),
        ("md_full", "Duty + TBA — Markdown (.md)"),
        ("json", "Bundle — JSON (.json)"),
    ]
    if _docx:
        choices += [
            ("docx_duty", "Duty — Word (.docx)"),
            ("docx_full", "Duty + TBA — Word (.docx)"),
        ]
    if _pdf:
        choices += [
            ("pdf_duty", "Duty — PDF (.pdf)"),
            ("pdf_full", "Duty + TBA — PDF (.pdf)"),
        ]
    return choices


def build_datasheet_export(
    bundle: dict[str, Any],
    choice_id: str,
    *,
    equipment: str,
    slug: str,
) -> tuple[bytes, str, str]:
    """Build export bytes, file name, and MIME type for UI download."""
    slug = re.sub(r"[^\w\-]+", "_", str(slug or "export")).strip("_") or "export"
    if equipment == "liquid":
        stem = "AQUASIGHT_liquid_pumps_datasheet"
    elif equipment == "air_blower":
        stem = "AQUASIGHT_air_blower_RFQ"
    else:
        raise ValueError(f"Unknown equipment type: {equipment!r}")

    if choice_id == "md_duty":
        return (
            bundle_to_markdown(bundle, full_template=False).encode("utf-8"),
            f"{stem}_duty_{slug}.md",
            "text/markdown",
        )
    if choice_id == "md_full":
        return (
            bundle_to_markdown(bundle, full_template=True).encode("utf-8"),
            f"{stem}_full_{slug}.md",
            "text/markdown",
        )
    if choice_id == "json":
        return (
            bundle_to_json(bundle).encode("utf-8"),
            f"{stem}_{slug}.json",
            "application/json",
        )
    if choice_id == "docx_duty":
        return (
            bundle_to_docx_bytes(bundle, full_template=False),
            f"{stem}_duty_{slug}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if choice_id == "docx_full":
        return (
            bundle_to_docx_bytes(bundle, full_template=True),
            f"{stem}_full_{slug}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if choice_id == "pdf_duty":
        return (
            bundle_to_pdf_bytes(bundle, full_template=False),
            f"{stem}_duty_{slug}.pdf",
            "application/pdf",
        )
    if choice_id == "pdf_full":
        return (
            bundle_to_pdf_bytes(bundle, full_template=True),
            f"{stem}_full_{slug}.pdf",
            "application/pdf",
        )
    raise ValueError(f"Unknown export choice: {choice_id!r}")
